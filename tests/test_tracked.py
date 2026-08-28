"""변경 내용 추적(track changes)이 걸린 워드 설문지 파싱 검증.

    python3 tests/test_tracked.py

워드에서 추적을 켜고 수정하면 삽입된 글자가 <w:ins> 안으로 들어간다.
python-docx 의 paragraph.text 는 그걸 놓쳐서 문항이 통째로 사라진다.
spss_labels.element_text 가 이를 처리한다 (삭제분은 w:delText 라 자동 제외).
"""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
import spss_labels as sl

FIXTURE = ROOT / "tests" / "fixture_tracked.docx"
fails = 0


def check(cond, msg):
    global fails
    print(("  OK   " if cond else "  실패 ") + msg)
    if not cond:
        fails += 1


d = {v.name: v for v in sl.parse_docx(str(FIXTURE))}
check("SQ1" in d, "추적 없는 문항 인식")
check("SQ2" in d, "추적으로 삽입된 문항 인식 (<w:ins> 안의 글자)")
check(d.get("SQ1") and sorted(d["SQ1"].values) == [1, 2], "탭으로 나열된 보기 분리 유지")
check(d.get("SQ2") and sorted(d["SQ2"].values) == [1, 2],
      "삽입된 보기 인식")
check(d.get("SQ2") and "추적으로 삽입된" in d["SQ2"].label, "삽입된 문항 라벨 유지")

# 표 셀에 문단이 여러 개일 때 공백이 붙어버리지 않아야 한다
from docx import Document
from docx.oxml.ns import qn
doc = Document()
t = doc.add_table(rows=1, cols=1)
cell = t.cell(0, 0)
cell.text = "전혀 그렇지"
cell.add_paragraph("않다")
check(sl.element_text(cell._tc).strip() == "전혀 그렇지 않다",
      f"셀 내 여러 문단은 공백으로 이어짐 (실제 {sl.element_text(cell._tc).strip()!r})")

print("\n" + ("모두 통과" if fails == 0 else f"실패 {fails}건"))
sys.exit(1 if fails else 0)
