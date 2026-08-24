# -*- coding: utf-8 -*-
"""hwp_survey 회귀 테스트.

.hwpx는 ZIP+XML이라 테스트용 파일을 코드로 만들 수 있다.
바이너리 픽스처를 저장소에 넣지 않아도 CI에서 전 과정을 돌려볼 수 있다.
"""

import xml.sax.saxutils as sax
import zipfile

import pytest

from hwp_survey import items_to_dsl, parse_dsl, read_survey, summarize
from hwp_survey.parser import scale_columns
from hwp_survey.reader import clean

NS = "http://www.hancom.co.kr/hwpml/2011/paragraph"


def para(text):
    # <제목> 처럼 꺾쇠가 들어간 본문도 있으므로 XML 이스케이프가 필요하다
    return f"<hp:p><hp:run><hp:t>{sax.escape(text)}</hp:t></hp:run></hp:p>"


def cell(text):
    return f"<hp:tc><hp:subList>{para(text)}</hp:subList></hp:tc>"


def table(rows):
    body = "".join("<hp:tr>" + "".join(cell(c) for c in r) + "</hp:tr>" for r in rows)
    return f"<hp:p><hp:run><hp:tbl>{body}</hp:tbl></hp:run></hp:p>"


@pytest.fixture
def sample_hwpx(tmp_path):
    body = [
        para("직무교육 만족도 조사"),
        para("안녕하십니까? 본 조사는 교육과정 개선을 목적으로 합니다."),
        para("Ⅰ. 응답자 일반사항"),
        para("1. 귀하의 소속은 어디입니까? ① 영업 ② 개발 ③ 관리"),
        para("2. 귀하의 직급을 선택해 주십시오."),
        para("① 사원"), para("② 대리"), para("③ 과장 이상"),
        para("3. 유익했던 내용을 모두 골라 주십시오."),
        para("① 실습 ② 사례 발표"),
        para("Ⅱ. 교육 만족도"),
        table([["◀", "①", "②", "③", "④", "⑤", "▶"],
               ["전혀 그렇지 않다", "매우 그렇다"]]),
        table([["‘내용’에 관한 문항입니다."],
               ["1", "강의 내용이 유익했다.", "①", "②", "③", "④", "⑤"],
               ["2", "난이도가 적절했다.", "①", "②", "③", "④", "⑤"]]),
        para("4. 향후 교육에 대한 의견을 자유롭게 적어 주십시오."),
    ]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "설문지.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("mimetype", "application/hwp+zip")
        z.writestr("Contents/section0.xml", xml)
    return str(path)


# ------------------------------------------------------------------ 추출
def test_reader_extracts_paragraphs_and_tables(sample_hwpx):
    items = read_survey(sample_hwpx)
    assert sum(1 for k, _ in items if k == "table") == 2
    assert ("p", "Ⅰ. 응답자 일반사항") in items


def test_clean_joins_run_fragments():
    # 한글은 글자 모양이 바뀌는 지점마다 텍스트를 조각내므로 붙여 읽어야 한다
    assert clean("‘ 자율성 ’") == "‘자율성’"
    assert clean("200 만원 미만") == "200만원 미만"
    assert clean("[ 설문지 종료 ]") == "[설문지 종료]"


def test_clean_tightens_linebreak_padding():
    assert clean("기분이 좋   다.") == "기분이 좋다."


# ------------------------------------------------------------------ 파싱
def test_inline_options_split_into_choices(sample_hwpx):
    dsl = items_to_dsl(read_survey(sample_hwpx))
    assert "1. 귀하의 소속은 어디입니까? [단일]" in dsl
    assert "- 영업" in dsl


def test_options_on_following_lines_set_single_type(sample_hwpx):
    dsl = items_to_dsl(read_survey(sample_hwpx))
    assert "2. 귀하의 직급을 선택해 주십시오. [단일]" in dsl   # [단답]이 아니어야 한다


def test_multi_and_open_types(sample_hwpx):
    dsl = items_to_dsl(read_survey(sample_hwpx))
    assert "[복수]" in dsl        # '모두' 골라 주십시오
    assert "[장문]" in dsl        # '자유롭게' 적어 주십시오


def test_scale_legend_becomes_matrix_header():
    cols = scale_columns([["◀", "①", "②", "③", "④", "⑤", "▶"],
                          ["전혀 그렇지 않다", "매우 그렇다"]])
    assert len(cols) == 5
    assert cols[0].startswith("①") and "전혀" in cols[0]
    assert "매우" in cols[-1]


def test_matrix_table_with_group_row(sample_hwpx):
    blocks = parse_dsl(items_to_dsl(read_survey(sample_hwpx)))
    matrix = [b for b in blocks if b["kind"] == "question" and b["type"] == "표"]
    assert len(matrix) == 1
    kinds = [o["type"] for o in matrix[0]["options"]]
    assert kinds.count("group") == 1
    assert kinds.count("row") == 2


def test_summary_counts(sample_hwpx):
    found = summarize(parse_dsl(items_to_dsl(read_survey(sample_hwpx))))
    assert found["문항"] == 5          # 4문항 + 매트릭스 1
    assert found["매트릭스 표"] == 1
    assert found["섹션"] == 2


def test_dsl_roundtrip_is_stable(sample_hwpx):
    dsl = items_to_dsl(read_survey(sample_hwpx))
    assert summarize(parse_dsl(dsl)) == summarize(parse_dsl(dsl))


def test_unsupported_extension_raises(tmp_path):
    path = tmp_path / "설문지.doc"
    path.write_bytes(b"x")
    with pytest.raises(ValueError):
        read_survey(str(path))


# ================================================================== DP 스크립트
from hwp_survey.dp import (DPWriter, items_to_dp_dsl, parse_dp,  # noqa: E402
                           summarize_dp)


@pytest.fixture
def dp_hwpx(tmp_path):
    """리서치 초안 형태(문단만 있는 설문지) 샘플."""
    body = [
        para("<IT 리뷰 유튜버 설문지>"),
        para("■ 조사 대상자: 전국 만 20~39세 남녀"),
        para("■ 샘플 수: 250샘플"),
        para("■ 쿼터:"),
        para("20대 30대 합계"),
        para("남 62 63 125"),
        para("여 63 62 125"),
        para("---"),
        para("SQ1. 귀하의 성별은 무엇입니까?"),
        para("① 남성"), para("② 여성"),
        para("SQ2. 귀하의 연령은 어떻게 되십니까?"),
        para("(출생년도 입력)"),
        para("[PROG: 만20~39세만 진행]"),
        para("SQ3. 거주하고 계신 지역은 어디입니까?"),
        para("(운동 설문과 동일)"),
        para("SQ4. 평소 온라인 동영상 플랫폼을 이용하십니까?"),
        para("① 유튜브"), para("② 넷플릭스"), para("③ 이용하지 않음"),
        para("Q1. 다음은 유튜버의 특성에 관한 질문입니다."),
        para("(5점 척도)"),
        para("① 매력적이다."), para("② 멋있다."),
    ]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "초안.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("mimetype", "application/hwp+zip")
        z.writestr("Contents/section0.xml", xml)
    return str(path)


def test_dp_header_fields_and_quota(dp_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(dp_hwpx)))
    assert doc["제목"] == "IT 리뷰 유튜버 설문지"
    assert doc["샘플수"] == "250샘플"
    assert doc["쿼터표"][0] == ["", "20대", "30대", "합계"]   # 좌상단 빈 칸
    assert doc["쿼터표"][1] == ["남", "62", "63", "125"]


def test_dp_labels_are_preserved(dp_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(dp_hwpx)))
    labels = [b["label"] for b in doc["blocks"] if b["kind"] == "question"]
    assert labels == ["SQ1", "SQ2", "SQ3", "SQ4", "Q1"]        # 번호를 다시 매기지 않는다


def test_dp_response_tags(dp_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(dp_hwpx)))
    tags = {b["label"]: b["tag"] for b in doc["blocks"] if b["kind"] == "question"}
    assert tags["SQ1"] == "1개선택"
    assert tags["SQ2"] == "출생년도 입력"
    assert tags["SQ3"] == "지도에서 선택"
    assert tags["SQ4"] == "모두선택"          # '이용하지 않음' 배타 보기가 있으므로
    assert tags["Q1"] == "행별 1개선택"


def test_dp_matrix_scale_and_rows(dp_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(dp_hwpx)))
    q1 = [b for b in doc["blocks"] if b.get("label") == "Q1"][0]
    assert q1["scale"][0] == "전혀 그렇지 않다" and len(q1["scale"]) == 5
    assert [o["text"] for o in q1["options"]] == ["매력적이다.", "멋있다."]
    assert "체크해 주세요" in q1["text"]      # 안내 문장 자동 추가


def test_dp_alone_prog_added(dp_hwpx):
    dsl = items_to_dp_dsl(read_survey(dp_hwpx))
    assert "%PROG: 3번 보기는 단독선택만 가능" in dsl


def test_dp_alone_prog_can_be_disabled(dp_hwpx):
    dsl = items_to_dp_dsl(read_survey(dp_hwpx), add_alone_prog=False)
    assert "단독선택만 가능" not in dsl


def test_dp_prog_lines_kept(dp_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(dp_hwpx)))
    progs = [b["text"] for b in doc["blocks"] if b["kind"] == "prog"]
    assert "만20~39세만 진행" in progs
    assert "17개시도 지도제시" in progs        # 지도형 기본 지시문


def test_dp_summary(dp_hwpx):
    found = summarize_dp(parse_dp(items_to_dp_dsl(read_survey(dp_hwpx))))
    assert found["문항"] == 5
    assert found["선정문항(SQ)"] == 4
    assert found["행별 표"] == 1


def test_dp_docx_structure(dp_hwpx):
    import io

    from docx import Document

    doc = parse_dp(items_to_dp_dsl(read_survey(dp_hwpx)))
    doc["제외"] = "2026060452 참여자 제외"
    doc["blocks"].append({"kind": "verify", "text": "Q1 일자찍기 제외"})
    data = DPWriter().write(doc).to_bytes()

    d = Document(io.BytesIO(data))
    spec = d.tables[0]
    assert [c.text for c in spec.rows[0].cells][0] == "대상자"
    assert len(spec.rows[2].cells[1].tables) == 1          # 쿼터 격자는 중첩 표
    matrix = d.tables[-1]
    assert len(matrix.columns) == 6 and len(matrix.rows) == 3
    assert matrix.rows[0].cells[1].text.endswith("1")      # 라벨 + 척도 번호

    xml = d.element.xml
    assert 'w:val="yellow"' in xml                          # 제외 줄 형광
    assert 'w:color w:val="0000FF"' in xml                  # PROG 파란색
    assert 'w:color w:val="FF0000"' in xml                  # 검증 빨간색
    assert "나눔고딕" in xml


def test_dp_handles_table_based_survey(sample_hwpx):
    """표로 리커트를 짠 학술 설문지도 DP 스크립트로 변환된다."""
    doc = parse_dp(items_to_dp_dsl(read_survey(sample_hwpx)))
    qs = [b for b in doc["blocks"] if b["kind"] == "question"]
    labels = [q["label"] for q in qs]

    assert labels[:3] == ["SQ1", "SQ2", "SQ3"]        # 행별 문항 앞은 선정 문항
    assert any(l.startswith("Q") for l in labels)     # 행별 문항부터 본 문항
    matrix = [q for q in qs if q["tag"] == "행별 1개선택"]
    assert len(matrix) == 1
    assert [o["type"] for o in matrix[0]["options"]].count("group") == 1
    assert doc["쿼터표"] == []                        # 표를 쿼터로 오인하지 않는다


def test_dp_terminate_option_becomes_prog(tmp_path):
    body = [para("1. 이용 경험이 있으십니까?"), para("① 예   ② 아니오[설문지 종료]")]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "종료.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)

    dsl = items_to_dp_dsl(read_survey(str(path)))
    assert "- 아니오" in dsl and "설문지 종료]" not in dsl
    assert "%PROG: 2번 선택자 설문 종료" in dsl


def test_dp_scale_middle_labels_filled(tmp_path):
    """양 끝만 라벨이 있는 척도 안내표의 가운데 라벨을 채운다."""
    from hwp_survey.dp import fill_scale
    cols = ["① 전혀 그렇지 않다", "②", "③", "④", "⑤ 매우 그렇다"]
    assert fill_scale(cols) == ["전혀 그렇지 않다", "그렇지 않다", "보통이다",
                                "그렇다", "매우 그렇다"]


# ================================================ 면접조사 양식('문1.' + 코드 보기)
from hwp_survey.parser import detect_label_style, is_banner  # noqa: E402


def cell_multi(*texts):
    inner = "".join(para(t) for t in texts)
    return f"<hp:tc><hp:subList>{inner}</hp:subList></hp:tc>"


@pytest.fixture
def interview_hwpx(tmp_path):
    """'문N.' 번호 + '1. 보기' 코드 + 영역 배너 + 빈도 표를 쓰는 면접조사 설문지."""
    banner = table([["신문 이용"]])
    screening = ("<hp:p><hp:run><hp:tbl>"
                 "<hp:tr>" + cell("SQ1. 거주") + cell("1. 서울 2. 부산 3. 대구") + "</hp:tr>"
                 "</hp:tbl></hp:run></hp:p>")
    freq = table([["평일", "주말"],
                  ["이용 안함", "1일이용", "2일이용", "이용 안함", "1일이용", "2일이용"],
                  ["0", "1", "2", "0", "1", "2"]])
    likert = table([["[언론에 대한 인식]", "전혀 그렇지 않다", "별로그렇지않다",
                     "보통이다", "약간그렇다", "매우 그렇다"],
                    ["1. 우리나라 언론은 공정하다", "1", "2", "3", "4", "5"],
                    ["2. 우리나라 언론은 정확하다", "1", "2", "3", "4", "5"]])
    body = [
        screening, banner,
        para("문1. 지난 1주일 동안 종이신문을 읽으신 적이 있습니까?*"),
        para("* 홈페이지, 앱을 통한 이용은 제외"),
        para("1. 읽었다"),
        para("2. 읽지 않았다 → 문5로 이동"),
        para("문2. 얼마나 자주 읽으셨습니까?"),
        freq,
        para("문4. 읽으신 신문의 이름은 무엇입니까? [복수 응답]"),
        para("1. 조선일보"), para("9997. 기타(적을 것: ______)"),
        para("▷ 조사원: 문1의 1을 선택한 경우만 진행"),
        para("문40. 우리나라 언론에 대해 어떻게 생각하십니까?"),
        likert,
    ]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "면접조사.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)
    return str(path)


def test_label_style_detection(interview_hwpx, sample_hwpx):
    assert detect_label_style(read_survey(interview_hwpx)) == "prefixed"
    assert detect_label_style(read_survey(sample_hwpx)) == "bare"


def test_banner_table_becomes_section(interview_hwpx):
    assert is_banner([["신문 이용"]])
    dsl = items_to_dp_dsl(read_survey(interview_hwpx))
    assert "## 신문 이용" in dsl


def test_numbered_lines_are_options_not_questions(interview_hwpx):
    """'문N.' 표기 문서에서 '1. 읽었다'는 문항이 아니라 보기다."""
    doc = parse_dp(items_to_dp_dsl(read_survey(interview_hwpx)))
    q1 = [b for b in doc["blocks"] if b.get("label") == "문1"][0]
    assert [o["text"] for o in q1["options"]] == ["1. 읽었다",
                                                 "2. 읽지 않았다 → 문5로 이동"]
    assert any("홈페이지" in b["text"] for b in doc["blocks"]
               if b["kind"] == "note")             # 각주는 보기가 아니다


def test_original_labels_preserved(interview_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(interview_hwpx)))
    labels = [b["label"] for b in doc["blocks"] if b["kind"] == "question"]
    assert labels == ["SQ1", "문1", "문2", "문4", "문40"]   # 번호를 다시 매기지 않는다


def test_response_code_kept(interview_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(interview_hwpx)))
    q4 = [b for b in doc["blocks"] if b.get("label") == "문4"][0]
    assert q4["tag"] == "모두선택"                        # [복수 응답] 꼬리표
    assert any(o["text"].startswith("9997.") for o in q4["options"])


def test_screening_table_inline_options(interview_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(interview_hwpx)))
    sq1 = [b for b in doc["blocks"] if b.get("label") == "SQ1"][0]
    assert sq1["text"] == "거주"
    assert [o["text"] for o in sq1["options"]] == ["1. 서울", "2. 부산", "3. 대구"]


def test_fieldwork_instruction_becomes_prog(interview_hwpx):
    dsl = items_to_dp_dsl(read_survey(interview_hwpx))
    assert "%PROG: 조사원 - 문1의 1을 선택한 경우만 진행" in dsl


def test_unclassified_table_kept_as_grid(interview_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(interview_hwpx)))
    grids = [b for b in doc["blocks"] if b["kind"] == "grid"]
    assert len(grids) == 1                                # 빈도 표
    assert grids[0]["rows"][2] == ["0", "1", "2", "0", "1", "2"]
    q2 = [b for b in doc["blocks"] if b.get("label") == "문2"][0]
    assert q2["tag"] == "표 응답"                          # 응답 칸은 표에 있다


def test_numeric_likert_table_becomes_matrix(interview_hwpx):
    doc = parse_dp(items_to_dp_dsl(read_survey(interview_hwpx)))
    q40 = [b for b in doc["blocks"] if b.get("label") == "문40"][0]
    assert q40["tag"] == "행별 1개선택"
    assert q40["scale"][0] == "전혀 그렇지 않다"
    assert [o["text"] for o in q40["options"]] == ["우리나라 언론은 공정하다",
                                                  "우리나라 언론은 정확하다"]


def test_interview_docx_has_grid_and_banner(interview_hwpx):
    import io

    from docx import Document

    doc = parse_dp(items_to_dp_dsl(read_survey(interview_hwpx)))
    d = Document(io.BytesIO(DPWriter().write(doc).to_bytes()))
    texts = [t.rows[0].cells[0].text for t in d.tables]
    assert "신문 이용" in texts                             # 배너 표
    assert any(len(t.columns) == 6 for t in d.tables)      # 빈도/행별 표


def test_shared_parser_handles_interview_form(interview_hwpx):
    """중간 표현(parse_dsl)은 DP/ISAS가 함께 쓰는 공용 층이라 계속 검사한다."""
    blocks = parse_dsl(items_to_dsl(read_survey(interview_hwpx)))
    found = summarize(blocks)
    assert found["문항"] == 5 and found["일반 표"] == 1
    labels = [b.get("label") for b in blocks if b["kind"] == "question"]
    assert "문40" in labels


# ============================================================== ISAS 표준 서식
from hwp_survey.isas import (ISASWriter, items_to_isas_dsl,  # noqa: E402
                             parse_isas, summarize_isas)


@pytest.fixture
def isas_hwpx(tmp_path):
    """PART 구역 + ※ 척도 안내 + ○ 표 + (□ 예, □ 아니오) 동의 문항."""
    likert = table([["문항", "1", "2", "3", "4", "5"],
                    ["P2-1. 반려견은 가족이다.", "○", "○", "○", "○", "○"],
                    ["P2-2. 정서적 유대가 있다.", "○", "○", "○", "○", "○"]])
    body = [
        para("Chiengora 인식 조사"),
        para("안녕하십니까? 본 설문은 견모에 대한 인식을 조사합니다."),
        para("(□ 예, □ 아니오) 나는 설명문을 읽었습니다."),
        para("(□ 예, □ 아니오) 나는 참여에 동의합니다."),
        para("선별 문항"),
        para("S1. 귀하는 현재 대한민국에 거주하고 있습니까? □ 예 □ 아니오(설문 종료)"),
        para("S2. 귀하의 성별은 무엇입니까? □ 여성 □ 그 외(설문 종료)"),
        para("PART 1. 일반 특성"),
        para("1. 최종 학력 □ 고등학교 이하 □ 대학교 □ 대학원 이상"),
        para("2. 거주 지역 □ 수도권 □ 광역시 □ 그 외"),
        para("PART 2. 반려견 친화성"),
        para("다음 문항에 대해 귀하의 생각과 가장 가까운 정도를 표시해 주십시오."),
        para("※ 1=전혀 그렇지 않다, 3=보통이다, 5=매우 그렇다"),
        likert,
        para("응답자 배경"),
        para("배문1. 귀하의 직업은 무엇입니까? □ 학생 □ 직장인 □ 기타"),
    ]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "isas.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)
    return str(path)


def test_isas_numbering_scheme(isas_hwpx):
    """renumber=True: 선별 SQ / 동의 AQ / 본문 Qk-n / 배경 DQ."""
    doc = parse_isas(items_to_isas_dsl(read_survey(isas_hwpx), renumber=True))
    labels = [b["label"] for b in doc["blocks"] if b["kind"] == "question"]
    assert labels == ["AQ", "SQ1", "SQ2", "Q1-1", "Q1-2", "Q2", "DQ1"]


def test_isas_tag_vocabulary(isas_hwpx):
    dsl = items_to_isas_dsl(read_survey(isas_hwpx))
    assert "[1개 선택]" in dsl                      # 숫자와 '개' 사이를 띄운다
    assert "[행별 1개 선택]" in dsl
    assert "[1개선택]" not in dsl


def test_isas_box_options_and_stop_prog(isas_hwpx):
    doc = parse_isas(items_to_isas_dsl(read_survey(isas_hwpx), renumber=True))
    sq1 = [b for b in doc["blocks"] if b.get("label") == "SQ1"][0]
    assert [o["text"] for o in sq1["options"]] == ["예", "아니오"]   # □ 보기 분리
    progs = [b["text"] for b in doc["blocks"] if b["kind"] == "prog"]
    assert "2번 선택자 설문 종료" in progs                          # (설문 종료)


def test_isas_consent_lines_merge_into_matrix(isas_hwpx):
    doc = parse_isas(items_to_isas_dsl(read_survey(isas_hwpx), renumber=True))
    aq = [b for b in doc["blocks"] if b.get("label") == "AQ"][0]
    assert aq["scale"] == ["예", "아니오"]
    assert len(aq["options"]) == 2


def test_isas_scale_note_becomes_labels(isas_hwpx):
    doc = parse_isas(items_to_isas_dsl(read_survey(isas_hwpx), renumber=True))
    q2 = [b for b in doc["blocks"] if b.get("label") == "Q2"][0]
    assert q2["scale"] == ["전혀 그렇지 않다", "그렇지 않다", "보통이다",
                          "그렇다", "매우 그렇다"]
    assert [o["text"] for o in q2["options"]] == ["반려견은 가족이다.",
                                                 "정서적 유대가 있다."]


def test_isas_no_text_added_to_stem(isas_hwpx):
    """규칙 1) 원문 텍스트를 바꾸지 않는다 -> 안내 문장을 덧붙이지 않는다."""
    assert "체크해 주세요" not in items_to_isas_dsl(read_survey(isas_hwpx))


def test_isas_docx_layout(isas_hwpx):
    import io

    from docx import Document

    doc = parse_isas(items_to_isas_dsl(read_survey(isas_hwpx)))
    doc["대상자"] = "만 20~69세 여성"
    doc["샘플수"] = "200"
    doc["쿼터표"] = [["", "20-29세", "합계"], ["여성", "40", "40"]]
    d = Document(io.BytesIO(ISASWriter().write(doc).to_bytes()))

    spec = d.tables[1]
    assert spec.rows[0].cells[0].text == "조사대상"
    assert len(spec.rows[2].cells[1].tables) == 1          # 할당 격자는 중첩 표

    matrix = [t for t in d.tables if len(t.columns) == 6][0]
    assert matrix.rows[0].cells[0].text == ""              # 첫 칸은 비운다
    assert matrix.rows[0].cells[1].text.endswith("1")      # 라벨 + 보기 코드
    assert matrix.rows[1].cells[3].text == "3"             # 표 안은 코드로 채움

    xml = d.element.xml
    assert 'w:color w:val="0000FF"' in xml                 # PROG 파란색
    assert "나눔고딕" in xml


def test_isas_summary(isas_hwpx):
    found = summarize_isas(parse_isas(items_to_isas_dsl(read_survey(isas_hwpx))))
    assert found["문항"] == 7 and found["행별 표"] == 2


# ========================================================== 변환 누락 검증
from hwp_survey.verify import compare, docx_text, hwp_text, sentences  # noqa: E402


def test_sentences_strip_markers_and_numbers():
    keys = [k for _, k in sentences("문1. 귀하의 최종 학력은 □ 고등학교 이하 □ 대학원")]
    assert "고등학교이하" in keys                        # 기호를 지우고 보기로 분리
    assert any("귀하의최종학력은" in k for k in keys)     # 문항 번호도 지운다


def test_sentences_ignore_scale_notes():
    assert sentences("※ 1=전혀 그렇지 않다, 3=보통이다, 5=매우 그렇다") == []


def test_compare_detects_missing_sentence():
    src = "1. 최종 학력\n1) 고등학교 이하\n2) 대학원 이상은 제외합니다"
    target = "Q1-1. 최종 학력 [1개 선택]\n1) 고등학교 이하"
    result = compare(src, target)
    assert result["누락"] == 1
    assert "대학원" in result["누락 목록"][0]


def test_compare_ignores_restructuring():
    """번호·기호가 바뀌고 표로 재배치되어도 누락으로 보지 않는다."""
    src = "(□ 예, □ 아니오) 나는 참여에 동의합니다."
    target = "AQ. 아래 항목을 읽고 응답해 주십시오.\n예\t아니오\n나는 참여에 동의합니다."
    assert compare(src, target)["누락"] == 0


def test_compare_partial_match_bucket():
    src = "귀하의 연령은 어떻게 되십니까 출생연도를 입력해 주십시오"
    target = "SQ3. 귀하의 연령은 어떻게 되십니까 [출생연도로 응답]"
    result = compare(src, target)
    assert result["누락"] == 0 and result["부분 일치"] == 1


def test_verify_end_to_end_on_generated_docx(isas_hwpx):
    """파서 텍스트 기준 검증: 자기 출력물에서는 누락이 없어야 한다."""
    import io

    from hwp_survey.isas import ISASWriter, items_to_isas_dsl, parse_isas

    doc = parse_isas(items_to_isas_dsl(read_survey(isas_hwpx)))
    data = ISASWriter().write(doc).to_bytes()
    result = compare(hwp_text(isas_hwpx), docx_text(io.BytesIO(data)))
    assert result["누락"] == 0, result["누락 목록"]


def test_matrix_group_row_is_kept(sample_hwpx):
    """검증기가 찾아낸 버그: 대분류 행이 표에서 빠지지 않아야 한다."""
    import io

    from docx import Document

    from hwp_survey.isas import ISASWriter, items_to_isas_dsl, parse_isas

    doc = parse_isas(items_to_isas_dsl(read_survey(sample_hwpx)))
    d = Document(io.BytesIO(ISASWriter().write(doc).to_bytes()))
    texts = [c.text for t in d.tables for r in t.rows for c in r.cells]
    assert any("문항입니다" in t for t in texts)


# ================================== 문 N. 표기 + 보기가 표에 든 설문지(SSK 계열)
@pytest.fixture
def spaced_label_hwpx(tmp_path):
    """'문 1.' 띄어쓰기, 'Q1】' 기호, '5-1.' 하위 번호, 보기만 담긴 한 칸 표."""
    likert = table([["문항", "전혀 없음", "", "거의 없음", "", "가끔", "",
                     "자주", "", "매우 자주"],
                    ["1.", "과제 아이디어를 얻는 데 사용하였다.", "①", "", "②", "",
                     "③", "", "④", "", "⑤"],
                    ["2.", "자료 요약에 사용하였다.", "①", "", "②", "", "③", "",
                     "④", "", "⑤"]])
    option_box = table([["① 30분 미만 ② 30분 이상~1시간 미만 ③ 1시간 이상"]])
    intro_box = table([["안녕하십니까? 저희 연구진은 생성형 AI의 학습 활용에 관한 대학생 "
                        "인식 연구를 위한 설문조사를 시행하고자 합니다. 모든 응답은 "
                        "철저히 익명으로 처리되며 어떠한 불이익도 발생하지 않습니다."]])
    body = [
        intro_box,
        para("귀하는 본 연구에 참여하는 것에 동의하십니까?"),
        para("① 동의함 ② 동의하지 않음"),
        para("(② 동의하지 않음 응답 시 설문 종료)"),
        para("문 1. 다음 문항은 활용 경험에 관해 묻는 문항입니다."),
        para("Q1】 귀하는 생성형 AI를 사용한 경험이 있습니까? ① 예 ② 아니오"),
        para("5-1. 다음과 같은 경험이 있는 친구는 어느 정도입니까?"),
        likert,
        para("10-1. 하루 평균 몇 시간 사용했습니까?"),
        option_box,
    ]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "ssk.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)
    return str(path)


def test_spaced_and_bracket_labels(spaced_label_hwpx):
    """'문 1.'(띄어쓰기)과 'Q1】'(기호)도 문항 번호로 인식한다."""
    doc = parse_isas(items_to_isas_dsl(read_survey(spaced_label_hwpx)))
    stems = [b["text"] for b in doc["blocks"] if b["kind"] == "question"]
    assert any("사용한 경험이 있습니까" in s for s in stems)
    assert not any(s.startswith("Q1") or s.startswith("문 1") for s in stems)


def test_group_heading_becomes_section(spaced_label_hwpx):
    """보기도 표도 없이 뒤 문항을 묶기만 하는 '문 1.'은 구역 제목이다."""
    doc = parse_isas(items_to_isas_dsl(read_survey(spaced_label_hwpx)))
    sections = [b["text"] for b in doc["blocks"] if b["kind"] == "section"]
    assert any("활용 경험에 관해 묻는 문항입니다" in s for s in sections)


def test_options_inside_single_cell_table(spaced_label_hwpx):
    """한 칸짜리 표에 든 보기는 상자글이 아니라 앞 문항의 보기다."""
    doc = parse_isas(items_to_isas_dsl(read_survey(spaced_label_hwpx)))
    q = [b for b in doc["blocks"] if b["kind"] == "question"
         and "하루 평균" in b["text"]][0]
    assert [o["text"] for o in q["options"]] == ["30분 미만",
                                                "30분 이상~1시간 미만",
                                                "1시간 이상"]


def test_sub_numbered_question_is_matrix_stem(spaced_label_hwpx):
    """'5-1.' 하위 번호 문항은 바로 뒤 표의 문항 문장이 된다."""
    doc = parse_isas(items_to_isas_dsl(read_survey(spaced_label_hwpx)))
    matrix = [b for b in doc["blocks"] if b["kind"] == "question"
              and b["tag"].startswith("행별")]
    assert len(matrix) == 1
    assert "친구는 어느 정도입니까" in matrix[0]["text"]
    assert matrix[0]["scale"] == ["전혀 없음", "거의 없음", "가끔", "자주", "매우 자주"]


def test_question_sentence_without_number(spaced_label_hwpx):
    """번호가 없어도 물음표로 끝나고 보기가 따르면 문항이다."""
    doc = parse_isas(items_to_isas_dsl(read_survey(spaced_label_hwpx)))
    q = [b for b in doc["blocks"] if b["kind"] == "question"
         and "참여하는 것에 동의" in b["text"]]
    assert len(q) == 1
    assert [o["text"] for o in q[0]["options"]] == ["동의함", "동의하지 않음"]


def test_skip_note_becomes_prog(spaced_label_hwpx):
    dsl = items_to_isas_dsl(read_survey(spaced_label_hwpx))
    assert "%PROG: ② 동의하지 않음 응답 시 설문 종료" in dsl


def test_prose_table_stays_box(spaced_label_hwpx):
    """긴 안내문이 담긴 표는 격자가 아니라 상자로 남는다."""
    doc = parse_isas(items_to_isas_dsl(read_survey(spaced_label_hwpx)))
    assert any(b["kind"] == "box" and "안녕하십니까" in b["text"]
               for b in doc["blocks"])
    assert not any(b["kind"] == "grid" for b in doc["blocks"])


def test_scale_legend_not_confused_with_matrix():
    """문항이 들어 있는 3행 표를 척도 안내표로 오인하지 않는다."""
    from hwp_survey.parser import scale_columns

    rows = [["문항", "매우 낮음", "", "보통", "", "매우 높음"],
            ["", "발각될 가능성이 얼마나 있다고 봅니까?", "①", "②", "③", "⑤"],
            ["2.", "징계로 이어질 가능성은?", "①", "②", "③", "⑤"]]
    assert scale_columns(rows) is None            # 기호가 두 행 -> 매트릭스다

    legend = [["◀", "①", "②", "③", "④", "⑤", "▶"],
              ["전혀 그렇지 않다", "매우 그렇다"]]
    assert scale_columns(legend) is not None      # 기호가 한 행 -> 안내 표


def test_two_point_yes_no_matrix():
    """'예 / 아니오' 두 칸짜리 표도 매트릭스다(5점으로 오인하거나 격자로 흘리지 않는다)."""
    from hwp_survey.parser import matrix_rows, scale_header

    rows = [["문항", "예", "", "아니오", ""],
            ["1.", "남이 잘 된 것이 부러웠던 적이 있다.", "①", "", "②", ""],
            ["2.", "가끔씩 나는 분하게 느낀다.", "①", "", "②", ""]]
    assert scale_header(rows) == ["예", "아니오"]
    assert matrix_rows(rows[1:], min_marks=2) == [
        "- 1. 남이 잘 된 것이 부러웠던 적이 있다.",
        "- 2. 가끔씩 나는 분하게 느낀다.",
    ]
    assert matrix_rows(rows[1:]) is None          # 5점 기준으로는 잡히지 않는다


def test_two_point_matrix_end_to_end(tmp_path):
    yes_no = table([["문항", "예", "", "아니오", ""],
                    ["1.", "남이 잘 된 것이 부러웠던 적이 있다.", "①", "", "②", ""],
                    ["2.", "가끔씩 나는 분하게 느낀다.", "①", "", "②", ""]])
    body = [para("문 15. 다음 각 문항에 응답해 주십시오."), yes_no]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "yesno.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)

    doc = parse_isas(items_to_isas_dsl(read_survey(str(path))))
    q = [b for b in doc["blocks"] if b["kind"] == "question"][0]
    assert q["tag"] == "행별 1개 선택"
    assert q["scale"] == ["예", "아니오"]
    assert len(q["options"]) == 2
    assert not any(b["kind"] == "grid" for b in doc["blocks"])   # 격자로 새지 않는다


def test_question_block_inside_single_cell_is_expanded(tmp_path):
    """한 칸짜리 표 안에 문항+보기가 통째로 든 경우, 문항이 사라지지 않아야 한다."""
    # 실제 한글 파일처럼 한 셀 안에 문단이 여러 개인 표
    block = ("<hp:p><hp:run><hp:tbl><hp:tr>"
             + cell_multi("① 제공함 ☞ 10-5로 이동",
                          "② 제공하지 않음 ☞ 문 11로 이동",
                          "10-5. 구독 서비스는 어떠한 형태입니까?",
                          "① 크레딧 한도 내에서 사용",
                          "② 추가 비용 없이 사용")
             + "</hp:tr></hp:tbl></hp:run></hp:p>")
    body = [para("10-4. 귀하의 대학에서 구독 서비스를 제공하고 있습니까?"), block]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "cell.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)

    doc = parse_isas(items_to_isas_dsl(read_survey(str(path))))
    stems = [b["text"] for b in doc["blocks"] if b["kind"] == "question"]
    assert any("제공하고 있습니까" in s for s in stems)
    assert any("어떠한 형태입니까" in s for s in stems)      # 상자에 묻히지 않는다
    first = [b for b in doc["blocks"] if b["kind"] == "question"][0]
    assert len(first["options"]) == 2                        # 앞 문항의 보기로 붙는다


def test_model_recovery_helper_exists():
    """XHTML 변환기가 버리는 글상자 문단을 이진 구조에서 되찾는 보조 경로."""
    from hwp_survey.reader import _recover_dropped

    items = [("p", "문 10. 활용 현황 문항입니다."), ("p", "1. 보기")]
    # .hwp 가 아니면 조용히 원본을 그대로 돌려준다(되찾기는 부가 기능)
    assert _recover_dropped("/존재하지/않는/파일.hwp", list(items)) == items


# ============================ 보기 표 · [DP:] · 수치 기입 · 양 끝 라벨 척도
@pytest.fixture
def option_grid_hwpx(tmp_path):
    """보기를 여러 칸에 나눈 표, [DP:] 지시문, '( )명' 기입란."""
    stadiums = table([["① 대구 삼성라이온즈파크", "② 부산 사직야구장"],
                      ["③ 서울 잠실야구장", "④ 대전 한화생명볼파크"],
                      ["⑤ 광주 기아챔피언스필드", "기타 구장"]])
    scale = table([["전혀 없음", "", "", "", "", "", "", "매우 빈번함"],
                   ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧"]])
    likert = table([["번호", "문항", "전혀그렇지않다", "", "매우그렇다"],
                    ["1", "AI 음원은 듣기 좋다.", "①", "②", "③", "④", "⑤",
                     "⑥", "⑦"],
                    ["2", "AI 음원은 자연스럽다.", "①", "②", "③", "④", "⑤",
                     "⑥", "⑦"]])
    body = [
        para("SQ2. 귀하의 출생연도는 어떻게 되십니까?"),
        para("[DP: 20대, 30대, 40대 구간별 결과 분석]"),
        para("SQ4. 방문한 구장을 모두 선택해 주십시오."),
        stadiums,
        para("A2-1. 평균 동반 인원은 몇 명입니까?"),
        para("()명"),
        para("Q2. 생성형 AI 음악을 들어본 적이 있습니까?"),
        scale,
        para("Q3. 다음 문항에 응답해 주십시오."),
        likert,
    ]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "grid.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)
    return str(path)


def test_options_spread_across_table_cells(option_grid_hwpx):
    """보기를 여러 칸에 나눠 담은 표는 격자가 아니라 보기 목록이다."""
    doc = parse_isas(items_to_isas_dsl(read_survey(option_grid_hwpx)))
    q = [b for b in doc["blocks"] if b["kind"] == "question"
         and "구장을 모두" in b["text"]][0]
    texts = [o["text"] for o in q["options"]]
    assert texts[:2] == ["대구 삼성라이온즈파크", "부산 사직야구장"]   # 행 우선
    assert texts[-1] == "기타 구장"                                  # 기호 없는 보기도
    assert q["tag"] == "모두 선택"


def test_dp_instruction_becomes_data_note(option_grid_hwpx):
    dsl = items_to_isas_dsl(read_survey(option_grid_hwpx))
    assert "%검증: 20대, 30대, 40대 구간별 결과 분석" in dsl


def test_numeric_input_field(option_grid_hwpx):
    doc = parse_isas(items_to_isas_dsl(read_survey(option_grid_hwpx)))
    q = [b for b in doc["blocks"] if b["kind"] == "question"
         and "동반 인원" in b["text"]][0]
    assert q["tag"] == "수치형"


def test_single_item_scale_becomes_options(option_grid_hwpx):
    """'전혀 없음 ①…⑧ 매우 빈번함' 안내표는 앞 문항의 보기다(행별 표가 아니다)."""
    doc = parse_isas(items_to_isas_dsl(read_survey(option_grid_hwpx)))
    q = [b for b in doc["blocks"] if b["kind"] == "question"
         and "들어본 적이" in b["text"]][0]
    texts = [o["text"] for o in q["options"]]
    assert len(texts) == 8
    assert texts[0] == "전혀 없음" and texts[-1] == "매우 빈번함"


def test_scale_header_drops_index_columns_and_fills_points(option_grid_hwpx):
    """'번호 | 문항 | 전혀그렇지않다 | | 매우그렇다' -> 7점 라벨."""
    doc = parse_isas(items_to_isas_dsl(read_survey(option_grid_hwpx)))
    matrix = [b for b in doc["blocks"] if b["kind"] == "question"
              and b["tag"].startswith("행별")][0]
    assert matrix["scale"][0] == "전혀그렇지않다"
    assert matrix["scale"][-1] == "매우그렇다"
    assert len(matrix["scale"]) == 7           # 기호 개수에 맞춰 채운다
    assert "문항" not in matrix["scale"]        # 구분 열 이름은 라벨이 아니다


def test_options_after_grid_are_not_dropped():
    """표(@표:)가 끼어들어도 뒤따르는 보기를 잃지 않는다."""
    dsl = ("Q1. 지출 항목을 적어 주십시오. [표 응답]\n"
           "@표: 티켓 구매 비용,()원\n"
           "@표: 식음료 비용,()원\n"
           "- 가격이 비싸서\n"
           "- 원하는 메뉴가 없어서")
    doc = parse_isas(dsl)
    q = [b for b in doc["blocks"] if b["kind"] == "question"][0]
    assert [o["text"] for o in q["options"]] == ["가격이 비싸서", "원하는 메뉴가 없어서"]


def test_letter_part_labels_and_deep_subnumbers(tmp_path):
    """'B0.' 'A3-2-1.' 처럼 알파벳 파트 번호와 여러 단계 하위 번호를 인식한다."""
    body = [para("B0. 겪으신 피해 영역을 모두 선택해 주십시오."),
            para("① 온라인 티켓 예매 ② 현장판매 티켓 구매"),
            para("A3-2-1. 식음료를 구매하지 않으신 이유는 무엇입니까?"),
            para("① 가격이 비싸서 ② 원하는 메뉴가 없어서")]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "parts.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)

    doc = parse_isas(items_to_isas_dsl(read_survey(str(path))))
    qs = [b for b in doc["blocks"] if b["kind"] == "question"]
    assert len(qs) == 2                                  # 둘 다 문항으로 인식
    assert all(len(q["options"]) == 2 for q in qs)


def test_recovered_paragraph_placement_keeps_order():
    """되찾은 문단은 번호 차례에 맞는 자리에 들어간다."""
    from hwp_survey.reader import _label_key, _place

    items = [("p", "A3-2. 지출 항목별 비용을 기입해 주십시오."),
             ("table", [["① 가격이 비싸서"], ["② 원하는 메뉴가 없어서"]]),
             ("p", "A3-3. 총 비용은 어느 범위입니까?"),
             ("p", "A4. 온라인 예매 경험을 묻습니다.")]
    assert _label_key("A3-2-1. 왜 구매하지 않으셨습니까?") == (1, (3, 2, 1))

    # 읽던 자리(cursor)가 A3-3 을 지나쳐 버린 경우 -> 번호 차례로 되돌린다
    assert _place(items, "A3-2-1. 왜 구매하지 않으셨습니까?", len(items)) == 1
    # 아직 지나치지 않았으면 읽던 자리를 그대로 쓴다
    assert _place(items, "A3-2-1. 왜 구매하지 않으셨습니까?", 1) == 1
    # 번호가 앞선 문항(A3-3)을 넘어가지 않는다
    assert _place(items, "A3-4. 좌석은 어디였습니까?", len(items)) == 3


def test_isas_keeps_original_labels_by_default(isas_hwpx):
    """기본값은 원본 번호 유지. [PROG] 지시문이 원본 번호를 가리키기 때문."""
    doc = parse_isas(items_to_isas_dsl(read_survey(isas_hwpx)))
    labels = [b["label"] for b in doc["blocks"] if b["kind"] == "question"]
    assert "S1" in labels and "S2" in labels      # 다시 매기지 않는다
    assert "배문1" in labels
    assert not any(l.startswith("DQ") for l in labels)


# =============================== 표 첫 행이 구역 제목 / 한 행이 문항 하나인 표
@pytest.fixture
def titled_table_hwpx(tmp_path):
    """'Ⅵ. …' 제목을 표 안에 넣은 척도 표 + 인적사항(한 행 = 한 문항) 표."""
    titled = table([["Ⅵ. 다음은 제도적 지원에 관한 의견을 묻는 질문입니다."],
                    ["", "", "전혀 그렇지않다", "-", "보통이다", "-", "매우 그렇다"],
                    ["1", "창업자금 지원제도가 도움이 된다.", "①", "②", "③", "④",
                     "⑤", "⑥", "⑦"],
                    ["2", "정부의 자금지원이 도움이 된다.", "①", "②", "③", "④",
                     "⑤", "⑥", "⑦"]])
    demo = table([["귀하의 성별은?", "① 남성 ② 여성"],
                  ["귀하의 연령은?", "① 20대 ② 30대 ③ 40대"]])
    body = [titled, para("Ⅶ. 아래 질문을 읽어보시고 해당하는 답에 체크해 주십시오."),
            demo]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "titled.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)
    return str(path)


def test_table_title_row_becomes_section(titled_table_hwpx):
    """표 첫 행의 'Ⅵ. …' 제목을 문항 행으로 오인하지 않는다."""
    doc = parse_isas(items_to_isas_dsl(read_survey(titled_table_hwpx)))
    sections = [b["text"] for b in doc["blocks"] if b["kind"] == "section"]
    assert any("제도적 지원" in s for s in sections)

    matrix = [b for b in doc["blocks"] if b["kind"] == "question"
              and b["tag"].startswith("행별")][0]
    assert len(matrix["scale"]) == 7                      # 표 자체 라벨을 쓴다
    assert matrix["scale"][0] == "전혀 그렇지않다"
    rows = [o["text"] for o in matrix["options"]]
    assert rows == ["창업자금 지원제도가 도움이 된다.", "정부의 자금지원이 도움이 된다."]
    assert not any("Ⅵ." in r for r in rows)               # 제목이 행으로 들어가지 않는다


def test_one_question_per_row_table(titled_table_hwpx):
    """'귀하의 성별은? | ① 남성 ② 여성' 표는 행마다 문항 하나다."""
    doc = parse_isas(items_to_isas_dsl(read_survey(titled_table_hwpx)))
    qs = [b for b in doc["blocks"] if b["kind"] == "question"
          and b["text"].startswith("귀하의")]
    assert [q["text"] for q in qs] == ["귀하의 성별은?", "귀하의 연령은?"]
    assert [len(q["options"]) for q in qs] == [2, 3]
    assert all(q["tag"] == "1개 선택" for q in qs)          # 행별 표가 아니다


# ============================================== 행이 많은 행별 표 나누기
def _matrix_doc(rows: int):
    dsl = ("Q1. 다음 각 항목에 응답해 주십시오. [행별 1개 선택]\n"
           "@행별: 전혀 그렇지 않다,보통이다,매우 그렇다\n"
           + "\n".join(f"- 항목 {i}" for i in range(1, rows + 1)))
    return dsl


def test_matrix_not_split_under_limit():
    doc = parse_isas(_matrix_doc(18), split_matrix=20)
    qs = [b for b in doc["blocks"] if b["kind"] == "question"]
    assert len(qs) == 1 and len(qs[0]["options"]) == 18


def test_matrix_split_into_balanced_halves():
    """24행 · 상한 20 -> 12 + 12 (상한을 채우지 않고 고르게 나눈다)."""
    doc = parse_isas(_matrix_doc(24), split_matrix=20)
    qs = [b for b in doc["blocks"] if b["kind"] == "question"]
    assert [len(q["options"]) for q in qs] == [12, 12]
    assert [q["label"] for q in qs] == ["Q1-1", "Q1-2"]
    assert all(q["scale"] == ["전혀 그렇지 않다", "보통이다", "매우 그렇다"] for q in qs)
    assert all(q["text"].startswith("다음 각 항목에") for q in qs)   # 문항 문장 반복


def test_matrix_split_into_three():
    doc = parse_isas(_matrix_doc(45), split_matrix=20)
    sizes = [len(b["options"]) for b in doc["blocks"] if b["kind"] == "question"]
    assert sizes == [15, 15, 15]


def test_matrix_split_keeps_group_header_with_its_rows():
    """소제목 행만 남기고 끊지 않는다."""
    dsl = ("Q1. 다음 각 항목에 응답해 주십시오. [행별 1개 선택]\n"
           "@행별: 예,아니오\n"
           + "\n".join(f"- 항목 {i}" for i in range(1, 4))
           + "\n-- 두 번째 묶음\n"
           + "\n".join(f"- 항목 {i}" for i in range(4, 7)))
    doc = parse_isas(dsl, split_matrix=3)
    qs = [b for b in doc["blocks"] if b["kind"] == "question"]
    assert len(qs) >= 2
    assert qs[1]["options"][0]["type"] == "group"     # 소제목이 뒤 조각의 머리로
    assert qs[0]["options"][-1]["type"] == "row"      # 소제목으로 끝나지 않는다


def test_split_disabled_by_default():
    doc = parse_isas(_matrix_doc(30))
    assert len([b for b in doc["blocks"] if b["kind"] == "question"]) == 1


def test_numeric_response_matrix():
    """응답 칸이 기호(①)가 아니라 숫자(1 2 3 4 5)인 리커트 표."""
    from hwp_survey.parser import numeric_matrix

    rows = [["번호", "문항", "전혀 그렇지 않다", "그렇지 않다", "보통 이다",
             "그렇다", "매우 그렇다"],
            ["1", "나는 눈치를 많이 살피는 편이다.", "1", "2", "3", "4", "5"],
            ["2", "비판받으면 굴욕감을 느낀다.", "1", "2", "3", "4", "5"]]
    labels, texts = numeric_matrix(rows)
    assert labels == ["전혀 그렇지 않다", "그렇지 않다", "보통 이다", "그렇다",
                      "매우 그렇다"]                       # '번호'·'문항' 열은 제외
    assert texts == ["- 나는 눈치를 많이 살피는 편이다.",
                     "- 비판받으면 굴욕감을 느낀다."]        # 번호가 아니라 문항 문장


def test_numeric_matrix_ignores_frequency_grid():
    """0부터 시작하는 빈도 표는 리커트가 아니다."""
    from hwp_survey.parser import numeric_matrix

    rows = [["평일", "주말"],
            ["이용 안함", "1일이용", "0", "1"],
            ["이용 안함", "2일이용", "0", "2"]]
    assert numeric_matrix(rows) is None


def test_numeric_matrix_end_to_end(tmp_path):
    likert = table([["번호", "문항", "전혀 그렇지 않다", "그렇지 않다", "보통 이다",
                     "그렇다", "전적으로 그렇다"],
                    ["1", "나는 눈치를 많이 살피는 편이다.", "1", "2", "3", "4", "5"],
                    ["2", "비판받으면 굴욕감을 느낀다.", "1", "2", "3", "4", "5"],
                    ["3", "인생의 목표를 못 잡고 있다.", "1", "2", "3", "4", "5"]])
    body = [para("부록 1. 내현적 자기애 척도"), likert]
    xml = (f'<?xml version="1.0" encoding="UTF-8"?><hp:sec xmlns:hp="{NS}">'
           + "".join(body) + "</hp:sec>")
    path = tmp_path / "numeric.hwpx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", xml)

    doc = parse_isas(items_to_isas_dsl(read_survey(str(path))))
    q = [b for b in doc["blocks"] if b["kind"] == "question"][0]
    assert q["tag"] == "행별 1개 선택"
    assert q["scale"][-1] == "전적으로 그렇다"            # 표마다 다른 라벨을 지킨다
    assert len(q["options"]) == 3
    assert q["options"][0]["text"].startswith("나는 눈치를")
