"""견본 설문지에 '변경 내용 추적'이 걸린 문항을 추가한다.

    python3 tests/add_tracked_questions.py

워드에서 추적을 켜고 수정하면 삽입된 글자는 <w:ins>, 삭제된 글자는
<w:del> + <w:delText> 로 저장된다. python-docx 로는 이 마크업을 직접 만들 수 없어
문단을 먼저 쓴 뒤 XML 을 후처리한다.

추가되는 패턴
  Q19 : 문항 머리글과 보기 전체가 삽입된 문항      (문항이 통째로 사라지던 경우)
  Q20 : 보기 중 하나가 삭제된 문항                (코드가 1,2,4 로 끊김)
  Q21 : 문단 중간 일부만 삽입된 문항              (라벨이 잘려 나오던 경우)
  Q22 : 표 안의 척도 라벨이 삽입된 격자 문항       (표 셀의 추적 처리)
"""

import re
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

TESTS = Path(__file__).resolve().parent
SRC = TESTS / "fixture_patterns.docx"

# 후처리로 감쌀 run 을 찾기 위한 표식 (문서에는 그대로 남는 실제 문구)
INSERT_MARKS = [
    "Q19. 추적으로 삽입된 문항입니다. [단수]",
    "1) 예\t\t2) 아니오",
    "해당 조직 구성원(들)",
    "매우 그렇다(삽입)",
]
DELETE_MARKS = [
    "3) 삭제된 보기",
]


def set_table_grid(table, widths):
    """표에 테두리(tblBorders)와 열 너비(tblGrid)를 명시한다."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tblPr.addnext(grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)


def build_paragraphs(doc: Document) -> None:
    """추적 대상 문구를 별도 run 으로 나눠 담는다 (후처리에서 찾기 쉽도록)."""
    def para(*runs, bold_first=False):
        p = doc.add_paragraph()
        for i, text in enumerate(runs):
            r = p.add_run(text)
            r.font.size = Pt(10)
            if bold_first and i == 0:
                r.bold = True
        return p

    # Q19 — 머리글과 보기가 모두 삽입
    para("Q19. 추적으로 삽입된 문항입니다. [단수]", bold_first=True)
    para("1) 예\t\t2) 아니오")

    # Q20 — 보기 하나가 삭제되어 코드가 끊김
    para("Q20. 보기 하나가 추적으로 삭제된 문항입니다. [단수]", bold_first=True)
    para("1) 첫째 보기", "\t\t", "2) 둘째 보기", "\t\t", "3) 삭제된 보기", "\t\t", "4) 넷째 보기")

    # Q21 — 문단 중간만 삽입
    para("Q21. 나와 ", "해당 조직 구성원(들)", "은 어떤 관계입니까? [단수]", bold_first=True)
    para("1) 가깝다\t\t2) 보통이다\t\t3) 멀다")

    # Q22 — 표 안의 척도 라벨이 삽입
    para("Q22. 다음 각 문장에 얼마나 동의하십니까? [행별 1개 선택]", bold_first=True)
    table = doc.add_table(rows=3, cols=4)
    # 테두리와 열 너비를 명시한다. 이게 없으면 .doc 로 변환할 때
    # LibreOffice 가 표를 문단으로 뭉개서 격자 문항이 사라진다.
    set_table_grid(table, [2400, 2200, 2200, 2200])
    header = ["속성", "전혀 그렇지 않다", "보통", "매우 그렇다(삽입)"]
    for i, text in enumerate(header):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for r, item in enumerate(["1. 첫째 항목", "2. 둘째 항목"], start=1):
        table.cell(r, 0).text = item
        for c, code in enumerate(["1", "2", "3"], start=1):
            table.cell(r, c).text = code

    # 표가 문서 맨 끝에 오면 .doc 변환 때 문단으로 뭉개진다. 뒤에 문단을 하나 둔다.
    para("[PROG: Q22 는 행별 1개 선택]")


def wrap_revisions(xml: str) -> str:
    """표식이 든 run 을 <w:ins> / <w:del> 로 감싼다."""
    rid = 9100
    for mark in INSERT_MARKS:
        pattern = re.compile(
            r"<w:r>(?:(?!</w:r>).)*?" + re.escape(mark.replace("\t", "")) + r".*?</w:r>", re.S)
        m = pattern.search(xml)
        if not m:
            # 탭이 <w:tab/> 로 쪼개진 경우: 표식의 앞부분만으로 다시 찾는다
            head = mark.split("\t")[0]
            pattern = re.compile(
                r"<w:r>(?:(?!</w:r>).)*?" + re.escape(head) + r".*?</w:r>", re.S)
            m = pattern.search(xml)
        if not m:
            print(f"  경고: 삽입 표식을 찾지 못함 -> {mark[:24]!r}")
            continue
        rid += 1
        xml = xml.replace(
            m.group(0),
            f'<w:ins w:id="{rid}" w:author="검수" w:date="2026-01-01T00:00:00Z">'
            f"{m.group(0)}</w:ins>", 1)

    for mark in DELETE_MARKS:
        pattern = re.compile(
            r"<w:r>(?:(?!</w:r>).)*?" + re.escape(mark) + r".*?</w:r>", re.S)
        m = pattern.search(xml)
        if not m:
            print(f"  경고: 삭제 표식을 찾지 못함 -> {mark!r}")
            continue
        rid += 1
        # 삭제된 글자는 w:t 가 아니라 w:delText 로 저장된다
        deleted = m.group(0).replace("<w:t>", "<w:delText>").replace(
            "<w:t ", "<w:delText ").replace("</w:t>", "</w:delText>")
        xml = xml.replace(
            m.group(0),
            f'<w:del w:id="{rid}" w:author="검수" w:date="2026-01-01T00:00:00Z">'
            f"{deleted}</w:del>", 1)
    return xml


def main() -> int:
    if not SRC.exists():
        print(f"견본이 없습니다: {SRC}")
        return 1

    tmp = TESTS / "_with_tracked.docx"
    shutil.copy(SRC, tmp)
    doc = Document(str(tmp))
    build_paragraphs(doc)
    doc.save(str(tmp))

    with zipfile.ZipFile(tmp) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    xml = wrap_revisions(xml)

    with zipfile.ZipFile(tmp) as zin, zipfile.ZipFile(SRC, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = xml.encode("utf-8") if item.filename == "word/document.xml" \
                else zin.read(item.filename)
            zout.writestr(item, data)
    tmp.unlink()

    counts = {tag: len(re.findall(f"<{tag}[ >]", xml)) for tag in ("w:ins", "w:del", "w:delText")}
    print(f"{SRC.name} 갱신 — {counts}")
    print("이어서 .doc 도 만드세요:")
    print(f"  soffice --headless --convert-to doc --outdir {TESTS} {SRC}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
