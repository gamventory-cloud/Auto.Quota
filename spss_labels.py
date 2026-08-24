"""
╔══════════════════════════════════════════════════════════════════════════╗
║  파일명 : spss_labels.py                                                  ║
║  위치   : 리포지토리 최상단  (utils.py 와 같은 폴더, pages/ 안이 아님!)      ║
║                                                                          ║
║  이 파일은 SPSS 라벨링 공용 모듈입니다. 화면(UI) 코드가 전혀 없습니다.        ║
║  화면 코드는 pages/3___SPSS_라벨링.py 에 있습니다.                          ║
╚══════════════════════════════════════════════════════════════════════════╝

spss_labels.py — 워드 설문지에서 SPSS 라벨 만들기

  설문지.docx  ->  코드북(DataFrame/xlsx)  ->  .sps 구문 / .sav 데이터

코드북을 중간에 두는 이유
------------------------
워드 서식은 항상 예외가 있어서 자동 판정이 100% 가 될 수 없다. 한 번 눈으로 훑고
넘기는 편이, SPSS 에서 라벨이 어긋난 걸 나중에 발견하는 것보다 훨씬 싸다.
설문지가 개정되면 파싱만 다시 돌린다.

변수명 표기
----------
기본은 대문자입니다 (`Q13_3_1`, `A4_R1`). 소문자로 바꾸려면 이 파일 상단의
`VAR_NAME_CASE = "upper"` 를 `"lower"` 로 바꾸면 전체에 적용됩니다.
SPSS 는 변수명 대소문자를 구분하지 않으므로 데이터 매칭에는 영향이 없습니다.

utils.py 와의 관계
-----------------
  - 변수명 정리는 utils.sanitize_var_name 을 재사용한다 (SSOT 유지)
  - 원자료(csv/xlsx) 로드는 화면 쪽에서 utils.load_df 를 쓴다.
    .sav 는 utils 가 다루지 않으므로 이 파일의 read_data_upload 가 처리한다.

인식하는 문항 머리글
------------------
접두사는 문서마다 다르므로 대문자 조합을 폭넓게 받는다.
  숫자가 붙는 형태 : Q13-3. / A5-1. / DQ2. / CS1. / IN2. / AC1. / SQ3.
  숫자가 없는 형태 : EQD. / MV. / MS. / IR.   (블록 전체가 하나의 격자 문항)
  문자 접미        : Q14-a. / Q14-b.
  대소문자 혼용    : Com1. / Com1-2.
  마침표 누락       : B6-2-2 (하이픈 번호에 한해 허용)
보기는 `1.` `1)`, 동그라미 숫자 `①…㉚`, 표 안의 보기, 워드 자동 불릿을 모두 인식한다.
(불릿 보기는 번호가 없으므로 순서대로 코드를 부여하고 '확인필요'로 표시한다)
문단 안의 소프트 리턴(shift+enter)은 공백으로 합쳐 한 문장으로 본다.
구형 .doc(Word 97-2003) 은 LibreOffice 가 있으면 자동으로 .docx 로 변환한다.
격자 항목이 자기 변수명을 갖고 있으면(`1) EQD1. 우리 사회에서는 …`) 그 이름을 쓴다.

주요 규칙
--------
1. 보기        : `1. 남성` / `2) 아니오`, 한 단락에 탭으로 여러 개도 인식
2. 척도표      : 머리행의 `(1) … (7)`, `7 매우 신뢰한다`, `매우 진보↵(0)`,
                 `전혀 그렇지 않다 1` 모두 해석
3. 격자/순위/복수응답/슬라이더/주관식/입력란 자동 판정 (detect_type)
   입력란 표(`(      )`, `[Range: …]`, 체크박스 `□`)는 척도표가 아니므로
   열 제목을 코드로 오인하지 않고 숫자·문자 변수로 만든다
4. 1x1 표는 입력 안내일 때만 주관식. 개념 정의 박스는 무시한다
5. SPSS 한도는 바이트 기준 : 값라벨 120B, 변수라벨 256B (한글 1자 = 3B)
   라벨이 한도를 넘으면 문항 서두를 버리고 항목 텍스트를 살린다 (compose_label)
6. .sps 는 UTF-8 BOM 으로 내보낸다 (SPSS 유니코드 모드에서 한글 보존)
7. 라벨 검수 : validate() 로 코드북 정합성과 실제 데이터 대조를 검사한다
     - 데이터에 있는 값이 값라벨에 없으면 '오류' (코드 밀림·척도 시작값 오류 신호)
     - 복수응답 저장 방식이 데이터와 다르면 '오류'
     - 생성된 .sav 를 되읽어 라벨이 실제로 들어갔는지 왕복 검증
8. 값라벨 표기(value_style) — 기본값 numbered
     numbered : `1'  1) 남성'`  (기본. 조사기관 Label 시트 표기와 동일)
     plain    : `1 "남성"`
   척도 문항은 라벨 없는 중간 지점도 코드로 남긴다. 양 끝에만 라벨이 붙은
   7점·9점 척도에서 `2'  2) '` 처럼 번호만 출력된다 (plain 표기에서는 제외).
9. 복수응답 저장 방식(multi_style)
     category  : 열마다 선택한 보기 코드, 미선택은 공백 (기본) -> MRSETS MCGROUP
     position  : 보기별 열이며 그 열은 자기 코드만 (1열=1, 2열=2 …) -> MCGROUP
     dichotomy : 0=비선택 / 1=선택 더미 -> MRSETS MDGROUP
"""

import io
import re
import tempfile
import zipfile
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd
import pyreadstat
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

import utils

# 이 파일이 진짜 spss_labels.py 인지 호출부에서 확인하는 표식.
MODULE_ROLE = "spss_labels"
__version__ = "1.10.0"




# ==============================================================================
# 1. 워드 설문지 파싱
# ==============================================================================

# ---------------------------------------------------------------- 정규식 규칙

# 문항 머리글. 접두사는 문서마다 다르므로 대문자 조합을 폭넓게 받는다.
#   숫자가 붙는 형태 : Q13-3. / A5-1. / DQ2. / CS1. / IN2. / AC1.
#   숫자가 없는 형태 : EQD. / MV. / MS. / IR.   (블록 전체가 하나의 격자 문항)
# 단, 한 글자 + 숫자 없음(`A.`)은 오탐이 많아 제외한다.
RE_QHEAD = re.compile(
    r"^\s*("
    r"[A-Z][A-Za-z]{0,5}\s*\d+(?:\s*[-‐–_]\s*[0-9a-zA-Z]+)*"   # 문자+숫자 (Q14-a, Com1-2, Q4-1-1-1)
    r"|[A-Z][A-Za-z]{1,5}"                      # 문자만 (2글자 이상: EQD, MV)
    r")\s*[.)]\s*(.*)$"
)

# 번호 뒤 마침표가 누락된 머리글. `B6-2-2 (B6-1-1 응답자만) 만약 …`
# 하이픈이 들어간 번호에만 적용해 오탐을 줄인다.
RE_QHEAD_NODOT = re.compile(
    r"^\s*([A-Z][A-Za-z]{0,5}\s*\d+(?:\s*[-‐–_]\s*[0-9a-zA-Z]+)+)\s+(?=[(\[가-힣])(.*)$")

# 격자 항목이 자기 변수명을 갖고 있는 경우: `1) EQD1. 우리 사회에서는 …`
RE_ITEM_VAR = re.compile(r"^\s*(?:\d{1,3}\s*[.)]\s*)?([A-Z]{1,6}\d{1,3})\s*[.)]\s*(.+)$")

# 동그라미 숫자 -> 정수. DP 문서가 ①②③ 로 보기를 적는 경우가 많다.
CIRCLED = {c: i for i, c in enumerate("①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳", start=1)}
CIRCLED.update({c: i for i, c in enumerate("㉑㉒㉓㉔㉕㉖㉗㉘㉙㉚", start=21)})
CIRCLED_CHARS = "".join(CIRCLED)
CIRCLED_CHARS_RE = re.compile(f"[{CIRCLED_CHARS}]")

# 보기 한 개. `1. 남성` / `2) 아니오` / `99. 모름` / `① 남성`
RE_OPTION = re.compile(
    rf"^\s*(?:(\d{{1,3}})\s*[.)]|([{CIRCLED_CHARS}])\s*[.)]?)\s*(.+)$"
)


def option_code(match: re.Match) -> int:
    """RE_OPTION 매치에서 보기 코드를 꺼낸다 (숫자형·동그라미형 공통)."""
    return int(match.group(1)) if match.group(1) else CIRCLED[match.group(2)]


def option_text(match: re.Match) -> str:
    return match.group(3)

# 한 단락 안에 여러 보기가 탭/다중공백으로 붙어있는 경우 분리용
RE_OPT_SPLIT_LOOSE = re.compile(r"\s+(?=\d{1,2}\s*[.)]\s*\S)")
RE_OPT_SPLIT = re.compile(
    rf"(?:\t[ \t]*|\s{{3,}})(?=(?:\d{{1,3}}\s*[.)]|[{CIRCLED_CHARS}])\s*\S)")

# 라벨에서 제거할 지시문.
#   - `[PROG: …]` 및 실무에서 흔한 오타 `[RROG:`, `[PROR:`, `[DP:`
#   - `[Range: 1~100]`, `[1개 선택]`, `[모두 선택]`, `[직접 기입]` 등
RE_DIRECTIVE = re.compile(
    r"\[[A-Za-z]{2,6}\s*:.*?\]"
    r"|\[[^\[\]]*(?:랜덤|고정|선택|주관식|입력|기입|자동표시|중단|드롭다운|드롭박스|제시|파이핑)[^\[\]]*\]"
)

# 문항 끝에 붙는 필수응답 표시(`* *`)와 ※ 안내문 꼬리
RE_TAIL_NOISE = re.compile(r"[*\s]+$|\s*※.*$")

# 격자표 첫 열 머리글이 이런 값이면 구성개념명이 아니라 단순 안내이므로 라벨에 쓰지 않는다
GENERIC_TITLES = {"문항 내용", "문항내용", "문항", "속성", "구분", "내용", "항목", ""}

SEQ_NOTE = "확인필요: 설문지에 코드 미표기 -> 1부터 순차 부여"

MULTI_HINTS = ("모두 골라", "모두 선택", "중복", "복수")
RE_MULTI_MAX = re.compile(r"최대\s*\d+\s*개")
TEXT_HINTS = ("주관식", "최소 4byte", "직접 입력", "입력해야함", "입력하게",
              "직접 기입", "자유롭게 적어", "자유롭게 작성")
NUMERIC_HINTS = ("금액 기입", "숫자 기입", "숫자기입", "kWh 기입")
SLIDER_HINTS = ("슬라이더", "(0-100)", "0~100", "0-100")


def unwrap_brackets(text: str) -> str:
    """`[DP: 최초 오답여부: IN1_FAIL 변수 만들어주세요.]` -> 내용만 남긴다.

    clean() 은 지시문 대괄호를 통째로 지운다. 지시문 자체가 라벨이 되어야 하는
    경우(DP 작업지시로 생성되는 변수)에는 내용을 살려야 한다.
    """
    t = re.sub(r"[\[\]]", " ", text or "")
    t = t.replace("\u00a0", " ")
    return re.sub(r"\s+", " ", t).strip(" .*")


def clean(text: str) -> str:
    """지시문·안내문·중복 공백 제거."""
    text = RE_DIRECTIVE.sub(" ", text or "")
    text = text.replace("\u00a0", " ").replace("\n", " ")
    text = RE_TAIL_NOISE.sub("", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip(" .*")


# 변수명 대소문자. "upper" -> Q1 / "lower" -> q1
# SPSS 는 변수명 대소문자를 구분하지 않지만, 저장된 표기 그대로 화면에 보인다.
VAR_NAME_CASE = "upper"

# 순위 변수 접미 (A4_R1 / a4_r1)
RANK_SFX = "R" if VAR_NAME_CASE == "upper" else "r"


def varname(qid: str) -> str:
    """Q13-3 -> q13_3 (SPSS 변수명 규칙: 영문 시작, 하이픈 불가)."""
    name = re.sub(r"\s+", "", qid)
    name = name.upper() if VAR_NAME_CASE == "upper" else name.lower()
    name = re.sub(r"[-‐–]", "_", name)
    name = re.sub(r"[^0-9A-Za-z_]", "_", name)
    if not name[:1].isalpha():
        name = "v" + name
    return name


# ---------------------------------------------------------------- 문서 순회

@dataclass
class Block:
    kind: str          # 'p' | 't'
    text: str = ""
    rows: list[list[str]] = field(default_factory=list)
    is_list: bool = False   # 워드 자동 목록(불릿/번호) 항목인지


def read_blocks(path: str) -> list[Block]:
    """단락과 표를 문서에 나타난 순서대로 읽는다 (순서 보존이 핵심)."""
    doc = Document(path)
    out: list[Block] = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            # 문단 내 소프트 리턴(shift+enter)은 한 문장의 일부이므로 공백으로 합친다.
            para = Paragraph(child, doc)
            t = re.sub(r"[\n\v\r]+", " ", para.text).strip()
            if t and t.strip("\u00a0 "):
                # 워드 자동 목록(불릿/번호)은 본문 텍스트에 기호·번호가 남지 않는다.
                # 보기를 불릿으로만 적은 설문지를 놓치지 않으려면 XML 속성을 봐야 한다.
                is_list = bool(child.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"))
                out.append(Block("p", text=t, is_list=is_list))
        elif tag == "tbl":
            rows = [[c.text.strip() for c in r.cells] for r in Table(child, doc).rows]
            if rows:
                out.append(Block("t", rows=rows))
    return out


# ---------------------------------------------------------------- 표 해석

# 표 셀이 응답 입력란인지. `(        )`, `약 (   )만원`, `[Range: 1~100]`
RE_ENTRY_CELL = re.compile(r"\(\s*\)|Range\s*:")
RE_BLANK = re.compile(r"\(\s*\)")   # 괄호 입력란 하나
RE_CHECKBOX = re.compile(r"[□☐▢]")


def is_entry_cell(cell: str) -> bool:
    return bool(RE_ENTRY_CELL.search(cell or "")) or bool(RE_CHECKBOX.search(cell or ""))


RE_BARE_CODE = re.compile(r"^\s*(\d{1,3})\s*[.)]?\s*$")


def bare_code(cell: str) -> int | None:
    """`3`, `3)`, `3.` 처럼 코드만 든 칸 -> 3. 아니면 None."""
    m = RE_BARE_CODE.match(cell or "")
    return int(m.group(1)) if m else None


def is_numeric_row(cells: list[str]) -> bool:
    vals = [c.strip() for c in cells if c.strip()]
    return bool(vals) and all(bare_code(v) is not None for v in vals)


def parse_anchor(cell: str) -> tuple[int, str] | None:
    """머리행 셀 -> (코드, 앵커라벨). 다양한 표기를 모두 받는다.

    '(3)' -> (3, '')            '(1) 전혀 동의 안함' -> (1, '전혀 동의 안함')
    '7 매우 신뢰한다' -> (7, ...)  '매우 진보\n(0)' -> (0, '매우 진보')
    """
    cell = (cell or "").strip()
    if not cell:
        return None
    # 동그라미 숫자 표기 (`① 매우 불만족`, `③`)
    if m := re.match(rf"^([{CIRCLED_CHARS}])\s*[.)]?\s*(.*)$", cell, re.S):
        return CIRCLED[m.group(1)], clean(m.group(2))
    if m := re.fullmatch(r"\(?\s*(\d{1,3})\s*\)?", cell):
        return int(m.group(1)), ""
    if m := re.match(r"^\(?\s*(\d{1,3})\s*\)?[\s:.]+(.+)$", cell, re.S):
        return int(m.group(1)), clean(m.group(2))
    if m := re.match(r"^(.*?)\(?\s*(\d{1,3})\s*\)?$", cell, re.S):
        return int(m.group(2)), clean(m.group(1))
    return None


def anchors_from_header(header: list[str]) -> dict[int, str]:
    """머리행에서 {코드: 앵커라벨}. 라벨 없는 중간 지점은 빈 문자열."""
    scale: dict[int, str] = {}
    for cell in header:
        if hit := parse_anchor(cell):
            code, label = hit
            scale.setdefault(code, label)
    return scale


def parse_table(rows: list[list[str]]) -> dict[str, Any]:
    """표를 {'kind', 'scale', 'items'} 로 해석."""
    header, body = rows[0], rows[1:]
    first_col = [r[0].strip() for r in body if r and r[0].strip()]

    # 순위형: 첫 열에 '1순위' 류
    if any(re.match(r"^\d\s*순위", c) for c in first_col):
        ranks = [clean(c) for c in first_col if re.match(r"^\d\s*순위", c)]
        return {"kind": "rank", "items": ranks, "scale": {}}

    # 순위 기입 표: 머리행이 `1순위 | 2순위`, 본문은 빈 칸 (응답자가 직접 적음)
    head_ranks = [c for c in header if re.fullmatch(r"\s*(\d)\s*순위\s*", c or "")]
    if len(head_ranks) >= 2 and len(head_ranks) == len([c for c in header if c.strip()]):
        return {"kind": "rank_entry", "items": [clean(c) for c in head_ranks], "scale": {}}

    # 보기 표: 모든 칸이 `N) 라벨` 형태 (17개 시·도처럼 여러 줄에 걸쳐 있어도 모두 수집)
    cells_all = [c.strip() for r in rows for c in r if c.strip()]
    opt_hits = [RE_OPTION.match(c) for c in cells_all]
    if len(cells_all) >= 3 and all(opt_hits):
        opts = {option_code(m): clean(option_text(m)) for m in opt_hits}
        if len(opts) == len(cells_all):
            return {"kind": "options", "items": [], "scale": {}, "options": opts}

    # Code 표: 머리행이 `Code1 … CodeN`, 본문 한 줄에 보기 라벨이 들어간다
    codes = [parse_anchor(c) for c in header if re.match(r"\s*Code\s*\d+", c or "", re.I)]
    if len(codes) >= 2 and len(body) == 1:
        labels = [clean(c) for c in body[0][-len(codes):]]
        opts = {code: lab for (code, _), lab in zip(codes, labels) if lab}
        if opts:
            return {"kind": "options", "items": [], "scale": {}, "options": opts}

    # 1x1 표. 입력 안내문일 때만 주관식으로 보고, 개념 정의·안내 박스는 무시한다.
    if len(rows) == 1 and len(rows[0]) == 1:
        cell = rows[0][0]
        if any(h in cell for h in TEXT_HINTS) or re.search(r"\d+\s*byte", cell):
            return {"kind": "textbox", "items": [], "scale": {}}
        return {"kind": "note", "items": [], "scale": {}}

    # 첫 열이 척도 앵커가 아니면 속성(문항) 열로 보고 척도 해석에서 제외
    first_is_anchor = parse_anchor(header[0]) is not None
    # 두 번째 열 이후가 전부 비어 있고 첫 열에만 항목명이 있는 표 = 기입 양식
    # (기관명 / 담당자 연락처 / 이메일 처럼 응답자가 직접 적는 칸)
    if (len(rows) >= 2 and len(rows[0]) >= 2
            and all(not c.strip() for r in rows for c in r[1:])
            and all(r[0].strip() for r in rows)):
        return {"kind": "entry", "items": [clean(r[0]) for r in rows],
                "cols": [], "cells": [["" for _ in r[1:]] for r in rows], "scale": {}}

    # 입력란 표(숫자·문자 기입)는 척도표가 아니다. 열 제목을 코드로 오인하면
    # 값라벨이 조용히 잘못 붙으므로 먼저 걸러낸다.
    if any(is_entry_cell(c) for r in body for c in r[1:]) or any(
            is_entry_cell(c) for r in rows for c in r[1:]):
        # 1행이 열 제목인지 판정: 입력란이 없고 내용이 있으면 머리행으로 본다.
        first = rows[0]
        has_header = (len(first) > 1
                      and any(c.strip() for c in first[1:])
                      and not any(is_entry_cell(c) for c in first[1:]))
        cols = [clean(c) for c in first[1:]] if has_header else []
        data = rows[1:] if has_header else rows
        return {"kind": "entry", "items": [clean(r[0]) for r in data],
                "cols": cols, "cells": [list(r[1:]) for r in data], "scale": {}}

    # 본문에 항목 텍스트가 하나도 없으면(응답칸만 빈 표) 머리행 전체가 척도점이다.
    # 이때 첫 열을 버리면 첫 척도점("매우 불만족")이 통째로 사라진다.
    has_items = any(r[0].strip() and bare_code(r[0]) is None for r in body)
    use_full_header = first_is_anchor or not has_items
    scale = anchors_from_header(header if use_full_header else header[1:])

    # 코드가 명시된 행이 있으면 그 행을 코드로 채택 (라벨은 머리행 오른쪽부터 매칭)
    for r in body:
        # 행 전체가 숫자면 첫 칸도 코드다. r[1:] 만 보면 코드 1 이 사라진다.
        tail = r if is_numeric_row(r) else (r[1:] if len(r) > 1 else r)
        if is_numeric_row(tail):
            codes = [bare_code(c) for c in tail if bare_code(c) is not None]
            labels = [clean(h) for h in header][-len(codes):]
            if codes and len(codes) == len(labels):
                merged = dict(zip(codes, labels))
                # 앵커 텍스트가 더 정확하면 유지
                for k, v in merged.items():
                    hit = parse_anchor(v) if v else None
                    scale[k] = hit[1] if hit else v
            break

    # 코드가 전혀 표기되지 않은 척도표 -> 라벨 순서대로 1..n 부여 (검수 필요)
    seq_coded = False
    if not scale:
        cand = [clean(h) for h in (header if use_full_header else header[1:]) if h.strip()]
        if len(cand) >= 3:
            scale = {i: lab for i, lab in enumerate(cand, start=1)}
            seq_coded = True

    # 격자형: 첫 열에 속성 텍스트가 있는 표
    items = [c for c in first_col if bare_code(c) is None and len(c) > 1]
    if items and len(header) > 2 and not use_full_header:
        title = clean(header[0])
        if title in GENERIC_TITLES:
            title = ""
        return {"kind": "grid", "items": [clean(i) for i in items], "scale": scale,
                "seq_coded": seq_coded, "title": title}

    if scale:
        return {"kind": "scale", "items": [], "scale": scale, "seq_coded": seq_coded}
    return {"kind": "unknown", "items": [], "scale": {}}


# ---------------------------------------------------------------- 문항 수집

RE_DASH_OPTION = re.compile(r"^\s*[-–—•▪·]\s*(\S.*)$")


@dataclass
class Question:
    qid: str
    label: str
    raw: str = ""          # 지시문을 제거하지 않은 머리글 원문 (유형 판정용)
    dash_options: list[str] = field(default_factory=list)   # 번호 없는 불릿 보기
    lines: list[str] = field(default_factory=list)      # 머리글 이후 일반 단락
    options: dict[int, str] = field(default_factory=dict)
    tables: list[dict[str, Any]] = field(default_factory=list)
    orphan_items: list[str] = field(default_factory=list)  # 번호 없는 속성문 + 척도표 쌍


def collect_questions(blocks: list[Block]) -> list[Question]:
    qs: list[Question] = []
    cur: Question | None = None
    pending_line: str | None = None   # 직전 일반 단락 (번호 없는 속성문 후보)

    for blk in blocks:
        if blk.kind == "p":
            raw = blk.text
            m = RE_QHEAD.match(raw) or RE_QHEAD_NODOT.match(raw)
            if m and not RE_OPTION.match(raw):
                cur = Question(qid=re.sub(r"\s+", "", m.group(1)),
                               label=clean(m.group(2)), raw=raw)
                qs.append(cur)
                pending_line = None
                continue
            if cur is None:
                continue
            # 보기 (한 줄에 여러 개일 수 있음).
            # 줄 끝의 지시문(`[PROG : 10) 기타 제외 …]`)에 보기 표시가 들어 있으면
            # 분리 결과가 어긋나므로 먼저 걷어낸다.
            opt_src = RE_DIRECTIVE.sub(" ", raw).strip()
            parts = RE_OPT_SPLIT.split(opt_src)
            hits = [RE_OPTION.match(p.strip()) for p in parts]
            embedded = any(h and re.search(r"\s\d{1,2}\s*[.)]\s*\S", h.group(0) or "")
                           for h in hits if h)
            if embedded or not (hits and all(hits) and len(parts) > 1):
                # 구분자가 공백 하나뿐인 경우(`1) 개선안A 2) 개선안B`).
                # 코드가 연속된 번호일 때만 인정해 오탐을 막는다.
                loose = RE_OPT_SPLIT_LOOSE.split(opt_src)
                lhits = [RE_OPTION.match(p.strip()) for p in loose]
                if len(loose) > 1 and all(lhits):
                    codes = [option_code(h) for h in lhits]
                    if codes == list(range(codes[0], codes[0] + len(codes))):
                        parts, hits = loose, lhits
            if all(hits) and hits:
                for h in hits:
                    cur.options[option_code(h)] = clean(option_text(h))
                pending_line = None
                continue
            if blk.is_list and not cur.options and len(clean(raw)) <= 80:
                cur.dash_options.append(clean(raw))
                pending_line = None
                continue
            if (m := RE_DASH_OPTION.match(raw)) and not cur.options:
                # 번호 없이 하이픈으로만 적힌 보기. 코드는 나중에 순서대로 부여한다.
                cur.dash_options.append(clean(m.group(1)))
                pending_line = None
                continue
            cur.lines.append(raw)
            pending_line = clean(raw)
        else:
            info = parse_table(blk.rows)
            if cur is None:
                continue
            # 번호 없는 속성문 + 척도표 조합 (A6, C6, D6 ... 형태)
            if info["kind"] == "note":
                continue
            if info["kind"] == "options":
                cur.options.update(info["options"])
                pending_line = None
                continue
            if (info["kind"] == "scale" and pending_line and len(pending_line) > 8
                    and not pending_line.rstrip().endswith(("?", "？"))):
                cur.orphan_items.append(pending_line)
                cur.tables.append({**info, "kind": "scale_item"})
            else:
                cur.tables.append(info)
            pending_line = None
    return qs


# ---------------------------------------------------------------- 유형 판정

DASH_NOTE = "확인필요: 보기에 번호가 없어 1부터 순차 부여 (설문지에 코드 명시 권장)"


def detect_type(q: Question) -> str:
    # 지시문([직접 기입], [금액 기입] 등)이 유형 판정의 핵심 근거이므로 원문을 함께 본다.
    text = " ".join([q.raw, q.label] + q.lines)
    tkinds = [t["kind"] for t in q.tables]

    if "entry" in tkinds:
        return "entry"
    if "rank_entry" in tkinds:
        return "rank_entry"
    if "rank" in tkinds:
        return "rank"
    # 표 없이 문구만으로 순위를 묻는 경우: "순서대로 선택해 주세요 [필수 4개 선택]"
    if re.search(r"순서대로|순위대로", text) and re.search(r"\d+\s*(?:개|가지)\s*(?:선택|기입)", text):
        return "rank"
    if re.search(r"최대\s*\d+\s*순위", text) and q.options:
        return "rank"
    # `[RANGE : 1~23]` 이 있으면 '직접 입력'이라도 숫자 문항이다.
    if re.search(r"\bRANGE\s*[:：]", text, re.I) and not q.options:
        return "numeric"
    if any(h in text for h in TEXT_HINTS) or "textbox" in tkinds:
        return "text"
    if any(h in text for h in SLIDER_HINTS):
        return "slider"
    if "grid" in tkinds:
        # 보기 목록이 따로 있으면 그 표는 격자가 아니라 '설명표'다.
        # (예: 개선안A/B/C 의 정의·특징을 비교한 표 + 아래에 1)2)3) 보기)
        # 행별 응답을 지시한 문항만 진짜 격자로 본다.
        if q.options and not re.search(r"행\s*별|행별|각\s*1개|항목별|항목 별", text):
            pass
        else:
            return "grid"
    if "scale_item" in tkinds:
        return "scale_item"
    opts = q.options or {i: t for i, t in enumerate(q.dash_options, start=1)}
    if (any(h in text for h in MULTI_HINTS) or RE_MULTI_MAX.search(text)) and opts:
        return "multi"
    if opts:
        return "single"
    if "scale" in tkinds:
        return "scale"
    if (any(h in text for h in NUMERIC_HINTS) or "입력" in text
            or re.search(r"\(\s+\)\s*년", text)
            or re.search(r"_{3,}\s*(?:년|원|개|명|회|세|시간)?", text)):
        return "numeric"
    return "unknown"


# ---------------------------------------------------------------- 코드북 생성

MAX_VARLABEL_BYTES = 256   # SPSS 변수라벨 한도 (바이트)


def shorten(text: str, colon_split: bool = True, limit: int = 120) -> str:
    """값 라벨용. `분노: 어쩌구...` -> `분노`. SPSS 값라벨은 120바이트 제한."""
    if colon_split and ":" in text:
        head = text.split(":", 1)[0].strip()
        if 0 < len(head) <= 24:
            return head
    return byte_trim(text, limit)


def byte_trim(text: str, limit: int) -> str:
    """UTF-8 바이트 기준 절단 (한글 1자 = 3바이트)."""
    b = text.encode("utf-8")
    if len(b) <= limit:
        return text
    return b[:limit].decode("utf-8", errors="ignore").rstrip()


def compose_label(stem: str, item: str, title: str = "") -> str:
    """격자 하위변수 라벨 조합.

    문항 서두(stem)가 긴 설문지에서 `stem - item` 을 그대로 쓰면 256바이트 한도에
    걸려 **항목 텍스트가 통째로 잘려나간다**. 그러면 하위변수 라벨이 전부 같아져
    코드북이 쓸모없어진다. 그래서 다음 순서로 고른다.

      1) `구성개념명 - 항목`   (표 첫 열 머리글이 있을 때. 짧고 정보량이 높다)
      2) `문항서두 - 항목`     (한도 안에 들어갈 때만)
      3) `항목`               (둘 다 넘치면 항목만. 항목은 절대 버리지 않는다)
    """
    for cand in ([f"{title} - {item}"] if title else []) + [f"{stem} - {item}", item]:
        if len(cand.encode("utf-8")) <= MAX_VARLABEL_BYTES:
            return cand
    return byte_trim(item, MAX_VARLABEL_BYTES)


@dataclass
class Var:
    name: str
    label: str
    vtype: str                       # numeric | string
    measure: str                     # nominal | ordinal | scale
    values: dict[int, str] = field(default_factory=dict)
    qid: str = ""
    kind: str = ""
    note: str = ""
    missing: str = ""      # 사용자 결측 (예: "99" 또는 "90-99")


MULTI_STYLES = ("category", "position", "dichotomy")


def build_vars(q: Question, colon_split: bool = True,
               multi_style: str = "category") -> list[Var]:
    kind = detect_type(q)
    base = varname(q.qid)
    label_src = q.label
    if not label_src:
        for ln in q.lines:
            c = clean(ln)
            if len(c) > 5:
                label_src = c
                break
    if not (label_src or "").strip():
        label_src = unwrap_brackets(q.raw) or q.qid
    if kind == "slider":
        asks = [clean(ln) for ln in q.lines
                if clean(ln).rstrip(") ").endswith(("?", "？")) or "온도는" in ln]
        if asks:
            label_src = asks[-1]
    qlabel = byte_trim(label_src or q.qid, 256)
    out: list[Var] = []

    def vl(d: dict[int, str]) -> dict[int, str]:
        """라벨이 있는 코드만."""
        return {k: shorten(v, colon_split) for k, v in d.items() if v}

    def vl_scale(d: dict[int, str]) -> dict[int, str]:
        """척도용. 라벨이 없는 중간 지점도 빈 라벨로 남긴다.

        `1 전혀 그렇지 않다 … 9 매우 그렇다` 처럼 양 끝에만 라벨이 붙은 척도에서
        중간 코드를 버리면 값라벨이 1, 9 만 나온다. 코드 체계를 그대로 보여주려면
        2~8 도 빈 라벨로 유지해야 한다.
        """
        return {k: shorten(v, colon_split) if v else "" for k, v in d.items()}

    # 번호 없는 하이픈 보기는 순서대로 코드를 부여한다 (검수 필요).
    dash_note = ""
    if not q.options and q.dash_options:
        q.options = {i: t for i, t in enumerate(q.dash_options, start=1)}
        dash_note = DASH_NOTE

    if kind == "single":
        out.append(Var(base, qlabel, "numeric", "nominal", vl(q.options), q.qid, kind,
                       note=dash_note))

    elif kind == "multi":
        # 복수응답 데이터 저장 방식은 조사기관마다 다르다.
        #   category   : 열마다 선택한 보기 코드가 들어가고 미선택은 공백
        #                -> 전체 보기 값라벨을 붙인다. 열 배치(보기별/응답순서별)와
        #                   무관하게 항상 맞으므로 기본값으로 쓴다.
        #   position   : 보기별 열이며 그 열은 자기 코드만 갖는다 (1열=1, 2열=2 …)
        #                -> 값라벨도 자기 코드 하나만
        #   dichotomy  : 0=비선택 / 1=선택 더미
        # SPSS 에서 category·position 은 MCGROUP, dichotomy 는 MDGROUP 으로 묶인다.
        # 이 구분은 비고에 남겨 코드북을 거쳐도 유지된다 (export 가 이 문자열을 읽는다).
        opts = vl(q.options)
        for code, opt in sorted(q.options.items()):
            label = byte_trim(f"{qlabel} - {shorten(opt, colon_split)}", 256)
            if multi_style == "dichotomy":
                values, note = {0: "비선택", 1: "선택"}, "복수응답 더미(0/1)"
            elif multi_style == "position":
                values, note = {code: opts.get(code, "")}, "복수응답(보기코드)"
            else:
                values, note = opts, "복수응답(보기코드)"
            out.append(Var(f"{base}_{code}", label, "numeric", "nominal",
                           values, q.qid, kind, note=note))

    elif kind == "rank":
        ranks = next((t["items"] for t in q.tables if t["kind"] == "rank"), None)
        if ranks is None:
            joined = " ".join([q.raw, q.label] + q.lines)
            m = (re.search(r"최대\s*(\d+)\s*순위", joined)
                 or re.search(r"(\d+)\s*(?:개|가지)\s*(?:선택|기입)", joined))
            n = int(m.group(1)) if m else 3
            ranks = [f"{i}순위" for i in range(1, min(n, 20) + 1)]
        for i, rk in enumerate(ranks, start=1):
            out.append(Var(f"{base}_{RANK_SFX}{i}", byte_trim(f"{qlabel} - {rk}", 256),
                           "numeric", "nominal", vl(q.options), q.qid, kind,
                           note="순위형"))

    elif kind in ("scale", "scale_item", "grid"):
        si = 0
        for t in q.tables:
            if t["kind"] == "grid":
                # 설문지가 속성마다 번호를 적어둔 경우 그 번호를 변수 접미로 쓴다.
                # 번호가 중간에 빠진 설문지(1,2,3,4,6,7…)에서 위치 기반으로 매기면
                # DP 변수명과 어긋난다. 단, 번호가 일부 항목에만 있으면 위치 기반.
                nums = [re.match(r"^\s*(\d{1,2})\s*[.)]", it) for it in t["items"]]
                explicit = ([int(m.group(1)) for m in nums]
                            if all(nums) and len({m.group(1) for m in nums}) == len(nums)
                            else None)
                for i, item in enumerate(t["items"], start=1):
                    suffix = explicit[i - 1] if explicit else i
                    # 설문지가 항목마다 변수명을 적어둔 경우 그 이름을 그대로 쓴다.
                    # (`1) EQD1. …` -> 변수명 eqd1. 위치 기반 번호는 블록을 넘나들며
                    #  이어지는 경우가 많아 신뢰할 수 없다.)
                    hit = RE_ITEM_VAR.match(item)
                    if hit:
                        vname, item_clean = varname(hit.group(1)), hit.group(2).strip()
                    else:
                        vname = f"{base}_{suffix}"
                        item_clean = re.sub(r"^\d{1,2}\s*[.)]\s*", "", item)
                    out.append(Var(vname, compose_label(qlabel, item_clean, t.get("title", "")),
                                   "numeric", "ordinal", vl_scale(t["scale"]), q.qid, "grid",
                                   note=SEQ_NOTE if t.get("seq_coded") else ""))
            elif t["kind"] == "scale_item":
                si += 1
                item = q.orphan_items[si - 1] if si <= len(q.orphan_items) else f"항목{si}"
                out.append(Var(f"{base}_{si}", compose_label(qlabel, item),
                               "numeric", "ordinal", vl_scale(t["scale"]), q.qid, "scale_item"))
            elif t["kind"] == "scale":
                out.append(Var(base, qlabel, "numeric", "ordinal", vl_scale(t["scale"]), q.qid, "scale",
                               note=SEQ_NOTE if t.get("seq_coded") else ""))

    elif kind == "entry":
        for t in q.tables:
            if t["kind"] != "entry":
                continue
            cols = t.get("cols") or [""]
            for i, item in enumerate(t["items"], start=1):
                cells = t["cells"][i - 1] if i - 1 < len(t["cells"]) else []
                for j, col in enumerate(cols, start=1):
                    cell = cells[j - 1] if j - 1 < len(cells) else ""
                    name = f"{base}_{i}" if len(cols) == 1 else f"{base}_{i}_{j}"
                    label = compose_label(qlabel, item if not col else f"{item} - {col}")
                    if RE_CHECKBOX.search(cell):
                        out.append(Var(name, label, "numeric", "nominal",
                                       {0: "비선택", 1: "선택"}, q.qid, "entry_check",
                                       note="체크박스 (선택/비선택)"))
                        continue
                    numeric = bool(RE_ENTRY_CELL.search(cell)) or any(
                        RE_ENTRY_CELL.search(c or "") for c in cells)
                    notes = ["숫자 기입" if numeric else "문자 기입(주관식)"]
                    if len(RE_BLANK.findall(cell)) > 1 or CIRCLED_CHARS_RE.search(cell):
                        # `① 지상 (  )층 ② 지하 (  )층`, `① (  )㎡ ② 모름` 처럼
                        # 한 칸에 입력란이 둘 이상이거나 '모름/해당없음' 보기가 섞인 경우
                        notes.append("확인필요: 한 칸에 입력란·보기가 둘 이상 "
                                     "(변수 분리 및 코딩 방식 확인)")
                    out.append(Var(name, label, "numeric" if numeric else "string",
                                   "scale" if numeric else "nominal",
                                   {}, q.qid, "entry", note=" / ".join(notes)))

    elif kind == "rank_entry":
        ranks = next((t["items"] for t in q.tables if t["kind"] == "rank_entry"), ["1순위", "2순위"])
        for i, rk in enumerate(ranks, start=1):
            if q.options:
                # 순위 칸 표 + 보기 목록 -> 보기 코드를 담는 숫자 순위 변수
                out.append(Var(f"{base}_{RANK_SFX}{i}", compose_label(qlabel, rk), "numeric", "nominal",
                               vl(q.options), q.qid, "rank",
                               note=" / ".join(n for n in ("순위형", dash_note) if n)))
            else:
                # 응답자가 직접 적는 순위 (사후 코딩 대상)
                out.append(Var(f"{base}_{RANK_SFX}{i}", compose_label(qlabel, rk), "string", "nominal",
                               {}, q.qid, kind, note="주관식 순위 기입"))

    elif kind == "slider":
        out.append(Var(base, qlabel, "numeric", "scale", {}, q.qid, kind, note="0-100 슬라이더"))

    elif kind == "numeric":
        out.append(Var(base, qlabel, "numeric", "scale", {}, q.qid, kind, note="숫자 직접입력"))

    elif kind == "text":
        out.append(Var(base, qlabel, "string", "nominal", {}, q.qid, kind, note="주관식"))

    elif not q.options and not q.tables:
        # 지도·드롭다운·리스트박스로 제시되어 보기가 문서에 없는 문항
        out.append(Var(base, qlabel, "numeric", "nominal", {}, q.qid, "single",
                       note="확인필요: 보기 목록이 설문지에 없음 (코드북에 직접 입력)"))

    else:
        why = ("파일 업로드·첨부 문항 (변수 필요 여부 확인)"
               if re.search(r"첨부|업로드", " ".join([q.raw] + q.lines))
               else "유형 판정 실패")
        out.append(Var(base, qlabel, "numeric", "nominal", vl(q.options), q.qid, "unknown",
                       note=f"확인필요: {why}"))

    # '모름/무응답' 류 보기는 사용자 결측으로 자동 제안 (코드북에서 수정 가능)
    for v in out:
        if v.kind == "multi":
            continue
        flags = [str(c) for c, lab in sorted(v.values.items())
                 if re.search(r"모름|무응답|응답\s*거부", lab)]
        if flags:
            v.missing = ",".join(flags[:3])
            v.note = (v.note + " / " if v.note else "") + f"결측 자동제안({v.missing})"

    # 중복 방지
    seen: set[str] = set()
    for v in out:
        if v.name in seen:
            n = 2
            while f"{v.name}_{n}" in seen:
                n += 1
            v.name = f"{v.name}_{n}"
        seen.add(v.name)
    return out


RE_DP_VAR = re.compile(r"([A-Za-z][A-Za-z0-9_]{1,30})\s*(?:변수|변수를)\s*(?:만들|생성|추가)")


def dp_instruction_vars(blocks: list[Block], existing: set[str]) -> list[Var]:
    """`[DP: … IN1_FAIL 변수 만들어주세요]` 같은 지시로 생성될 변수를 수집한다.

    설문 문항이 아니라 DP 작업 지시이므로 값라벨을 알 수 없다. 누락되는 것보다
    코드북에 '확인필요'로 올려두는 편이 낫다.
    """
    out: list[Var] = []
    for blk in blocks:
        if blk.kind != "p" or "변수" not in blk.text:
            continue
        for hit in RE_DP_VAR.finditer(blk.text):
            name = varname(hit.group(1))
            if name in existing or name in {v.name for v in out}:
                continue
            out.append(Var(name, byte_trim(unwrap_brackets(blk.text), 256), "numeric", "nominal",
                           {}, "", "dp_instruction",
                           note="확인필요: DP 지시로 생성되는 변수 (값라벨 직접 입력)"))
    return out


def parse_docx(path: str, colon_split: bool = True,
               multi_style: str = "category") -> list[Var]:
    """워드 설문지 -> 변수 목록.

    multi_style: 복수응답 저장 방식. category / position / dichotomy 중 하나.
    """
    blocks = read_blocks(path)
    questions = collect_questions(blocks)
    all_ids = {q.qid for q in questions}

    # 의문형 어미만 본다. `주십시오`·`선택` 등은 블록 안내문에도 흔해 제외.
    RE_QUESTIONISH = re.compile(r"\?|습니까|십니까|입니까|무엇|얼마나|어디|언제|어떻게")

    def is_section_header(q: Question) -> bool:
        """하위문항만 갖는 블록 머리글인지.

        `SQ2`(출생년도), `SQ4`(구매 시작 시기)처럼 보기가 표로 없거나 불릿이라
        비어 보이는 실제 문항을 머리글로 오인해 버리면 문항이 통째로 사라진다.
        질문문처럼 보이면 문항으로 유지한다.
        """
        empty = not q.options and not q.tables and not q.dash_options
        has_child = any(i != q.qid and i.startswith(q.qid + "-") for i in all_ids)
        looks_like_question = bool(RE_QUESTIONISH.search(f"{q.label} {' '.join(q.lines)}"))
        return empty and has_child and not looks_like_question

    variables: list[Var] = []
    used: set[str] = set()
    for q in questions:
        if is_section_header(q):
            continue
        for v in build_vars(q, colon_split, multi_style):
            if v.name in used:
                # 설문지 번호 체계상 겹치는 경우(예: SQ2 격자 항목 vs SQ2-1 문항).
                # 이름을 몰래 바꾸면 데이터와 어긋나므로 반드시 눈에 띄게 남긴다.
                original = v.name
                n = 2
                while (f"{original}_DUP{n}" if VAR_NAME_CASE == "upper" else f"{original}_dup{n}") in used:
                    n += 1
                v.name = f"{original}_DUP{n}" if VAR_NAME_CASE == "upper" else f"{original}_dup{n}"
                v.note = (v.note + " / " if v.note else "") + (
                    f"확인필요: 변수명 충돌 (원래 {original}) - 설문지 문항번호 체계 확인")
            used.add(v.name)
            variables.append(v)
    variables += dp_instruction_vars(blocks, used)
    return variables


def rebase_values(variables: list[Var], prefixes: list[str], start: int = 0) -> list[str]:
    """지정한 변수(접두사)의 값라벨 코드를 start부터 다시 매긴다.

    설문지에 코드가 표기되지 않아 1부터 순차 부여된 척도를 0부터로 바꾸는 용도.
    실제로 바뀐 변수명 목록을 반환한다.
    """
    changed: list[str] = []
    for v in variables:
        if not v.values:
            continue
        if not any(v.name.lower().startswith(pfx.lower()) for pfx in prefixes):
            continue
        codes = sorted(v.values)
        if codes[0] == start:
            continue
        offset = codes[0] - start
        v.values = {c - offset: lab for c, lab in v.values.items()}
        v.note = re.sub(r"확인필요: 설문지에 코드 미표기[^/]*", f"코드 {start}부터 적용(사용자 확인)", v.note).strip(" /")
        if v.missing:
            v.missing = ",".join(str(int(m) - offset) for m in v.missing.split(",") if m.strip().lstrip("-").isdigit())
        changed.append(v.name)
    return changed



# ==============================================================================
# 2. 코드북 엑셀 입출력
# ==============================================================================

HEADERS = ["순서", "변수명", "문항번호", "문항유형", "변수라벨", "유형", "측도", "값라벨", "결측값", "비고"]
WIDTHS = [6, 14, 12, 11, 60, 9, 9, 68, 10, 34]
FONT = "Arial"


def values_to_text(values: dict[int, str]) -> str:
    """{1:'남성',2:'여성'} -> '1=남성 | 2=여성'"""
    return " | ".join(f"{k}={v}" for k, v in sorted(values.items()))


def text_to_values(text: str) -> dict[int, str]:
    """'1=남성 | 2=여성' -> {1:'남성',2:'여성'} (형식 오류는 조용히 건너뜀)"""
    out: dict[int, str] = {}
    for chunk in str(text or "").split("|"):
        if "=" not in chunk:
            continue
        code, label = chunk.split("=", 1)
        code, label = code.strip(), label.strip()
        try:
            out[int(code)] = label
        except ValueError:
            continue
    return out


def write_codebook(variables: list[Var], path: str | Path) -> Path:
    path = Path(path)
    wb = Workbook()

    ws = wb.active
    ws.title = "codebook"
    ws.append(HEADERS)
    head_fill = PatternFill("solid", fgColor="D9E1F2")
    for i, _ in enumerate(HEADERS, start=1):
        c = ws.cell(row=1, column=i)
        c.font = Font(name=FONT, bold=True)
        c.fill = head_fill
        c.alignment = Alignment(vertical="center")
        ws.column_dimensions[get_column_letter(i)].width = WIDTHS[i - 1]

    todo_fill = PatternFill("solid", fgColor="FFF2CC")
    for i, v in enumerate(variables, start=1):
        ws.append([i, v.name, v.qid, v.kind, v.label, v.vtype, v.measure,
                   values_to_text(v.values), v.missing, v.note])
        for col in range(1, len(HEADERS) + 1):
            cell = ws.cell(row=i + 1, column=col)
            cell.font = Font(name=FONT)
            cell.alignment = Alignment(vertical="top", wrap_text=col in (5, 8, 10))
            if "확인필요" in v.note:
                cell.fill = todo_fill
    ws.freeze_panes = "A2"

    guide = wb.create_sheet("사용법")
    guide_rows = [
        ["이 파일은 워드 설문지에서 자동 추출한 코드북입니다. 검수 후 SPSS 산출물을 생성하세요."],
        [""],
        ["직접 수정할 열", "설명"],
        ["문항유형", "파서가 판정한 유형(single/multi/grid/rank/scale/slider/text). 참고용입니다."],
        ["변수명", "SPSS 변수명. 영문으로 시작, 공백/하이픈 불가. 데이터 파일의 열 이름과 일치시켜야 라벨이 붙습니다."],
        ["변수라벨", "SPSS VARIABLE LABELS 값. 256바이트 초과분은 자동 절단됩니다."],
        ["유형", "numeric 또는 string"],
        ["측도", "nominal / ordinal / scale"],
        ["값라벨", "'코드=라벨' 을 ' | ' 로 구분. 라벨이 필요 없으면 비워 둡니다."],
        ["결측값", "사용자 결측. 예) 99  또는  99,98  또는  범위형 90-99"],
        ["비고", "메모 열. 산출물에는 반영되지 않습니다."],
        [""],
        ["예시 행", ""],
        ["변수명", "q6"],
        ["변수라벨", "귀하께서는 자신이 다음 중 어느 계층에 속한다고 생각하십니까?"],
        ["유형", "numeric"],
        ["측도", "ordinal"],
        ["값라벨", "1=상위 | 2=중상위 | 3=중간 정도 | 4=중하위 | 5=하위 | 99=모름/무응답"],
        ["결측값", "99"],
        [""],
        ["노란색으로 칠해진 행", "자동 판정이 불확실한 문항입니다. 반드시 눈으로 확인하세요."],
    ]
    for r in guide_rows:
        guide.append(r)
    for row in guide.iter_rows():
        for cell in row:
            cell.font = Font(name=FONT, bold=cell.row in (1, 3, 12))
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    guide.column_dimensions["A"].width = 20
    guide.column_dimensions["B"].width = 88

    wb.save(path)
    return path


def read_codebook(path: str | Path) -> list[Var]:
    wb = load_workbook(path, data_only=True)
    ws = wb["codebook"] if "codebook" in wb.sheetnames else wb.active
    rows = ws.iter_rows(values_only=True)
    header = [str(h or "").strip() for h in next(rows)]
    idx = {name: header.index(name) for name in HEADERS if name in header}

    out: list[Var] = []
    for row in rows:
        def cell(key: str) -> str:
            i = idx.get(key)
            return "" if i is None or row[i] is None else str(row[i]).strip()

        name = cell("변수명")
        if not name:
            continue
        v = Var(
            name=name,
            label=cell("변수라벨"),
            vtype=(cell("유형") or "numeric").lower(),
            measure=(cell("측도") or "nominal").lower(),
            values=text_to_values(cell("값라벨")),
            qid=cell("문항번호"),
            kind=cell("문항유형"),
            note=cell("비고"),
            missing=cell("결측값"),
        )
        out.append(v)
    return out



# ==============================================================================
# 3. SPSS 산출물 (.sps / .sav)
# ==============================================================================

MAX_VALLABEL_BYTES = 120   # SPSS 값라벨 한도 (한글 1자 = 3바이트)


VALUE_STYLES = ("plain", "numbered")


def sq(text: str) -> str:
    """작은따옴표 리터럴. 내부 따옴표는 두 번 써서 이스케이프."""
    return "'" + str(text).replace("'", "''") + "'"


def numbered_label(code: int, label: str) -> str:
    """`  1) 남성` 형태. 조사기관 Label 시트 표기와 맞추기 위한 방식.

    이미 `1) ` 처럼 번호가 붙어 있으면 중복해서 붙이지 않는다.
    번호까지 포함해 120바이트 한도에 맞춰 자른다.
    """
    text = str(label).strip()
    if not text:
        # 라벨 없는 중간 척도점: `  2) ` (번호만)
        return f"  {code}) "
    if re.match(rf"^\s*{code}\s*\)", text):
        body = text
    else:
        body = f"{code}) {text}"
    return byte_trim(f"  {body}", MAX_VALLABEL_BYTES)


def styled_values(var, style: str) -> dict[int, str]:
    """표기 방식을 적용한 {코드: 라벨}."""
    if style != "numbered":
        # 기존 표기에서는 빈 라벨을 내보내지 않는다
        return {c: byte_trim(lab, MAX_VALLABEL_BYTES) for c, lab in var.values.items() if lab}
    return {c: numbered_label(c, lab) for c, lab in var.values.items()}


def q(text: str) -> str:
    """SPSS 문자열 리터럴. 내부 따옴표는 두 번 써서 이스케이프."""
    return '"' + str(text).replace('"', '""') + '"'


def parse_missing(spec: str) -> list[float] | tuple[float, float] | None:
    """'99' -> [99] / '98,99' -> [98,99] / '90-99' -> (90,99)"""
    spec = (spec or "").strip()
    if not spec:
        return None
    if m := re.fullmatch(r"(-?\d+(?:\.\d+)?)\s*[-~]\s*(-?\d+(?:\.\d+)?)", spec):
        return (float(m.group(1)), float(m.group(2)))
    vals = [v.strip() for v in re.split(r"[,\s]+", spec) if v.strip()]
    try:
        return [float(v) for v in vals][:3]   # SPSS 개별 결측은 최대 3개
    except ValueError:
        return None


# ------------------------------------------------------------------ .sps

def build_syntax(variables: list[Var], mrsets: bool = True, source: str = "",
                 value_style: str = "numbered") -> str:
    L: list[str] = [
        "* ==========================================================",
        "* SPSS 라벨 구문 (자동 생성)",
        f"* 원본 설문지: {source or '-'}",
        f"* 생성일: {date.today():%Y-%m-%d}",
        "* 사용법: 데이터셋을 연 상태에서 이 구문 전체를 실행하세요.",
        "* 인코딩: UTF-8(BOM). SPSS가 유니코드 모드여야 한글이 깨지지 않습니다.",
        "* ==========================================================",
        "",
    ]

    labeled = [v for v in variables if v.label]
    if labeled:
        L.append("VARIABLE LABELS")
        for i, v in enumerate(labeled):
            sep = "  " if i == 0 else "  /"
            L.append(f"{sep}{v.name} {q(byte_trim(v.label, MAX_VARLABEL_BYTES))}")
        L.append(".")
        L.append("")

    valued = [v for v in variables if v.values]
    if valued:
        L.append("VALUE LABELS")
        for i, v in enumerate(valued):
            sep = "  " if i == 0 else "  /"
            L.append(f"{sep}{v.name}")
            styled = styled_values(v, value_style)
            for code, label in sorted(styled.items()):
                if value_style == "numbered":
                    # `  1'  1) 남성'` — 코드와 라벨을 붙여 쓰는 표기
                    L.append(f"  {code}{sq(label)}")
                else:
                    L.append(f"    {code} {q(label)}")
        L.append(".")
        L.append("")

    by_level: dict[str, list[str]] = {}
    for v in variables:
        by_level.setdefault(v.measure.upper(), []).append(v.name)
    if by_level:
        # SPSS 구문은 긴 줄에서 문제가 생길 수 있어 변수 8개마다 줄바꿈
        L.append("VARIABLE LEVEL")
        for i, (lvl, names) in enumerate(by_level.items()):
            L.append(("  " if i == 0 else "  /") + f"({lvl})".rjust(0))
            L[-1] = ("  " if i == 0 else "  /") + " ".join(names[:8])
            for j in range(8, len(names), 8):
                L.append("    " + " ".join(names[j:j + 8]))
            L.append(f"    ({lvl})")
        L.append(".")
        L.append("")

    for v in variables:
        spec = parse_missing(getattr(v, "missing", ""))
        if not spec:
            continue
        if isinstance(spec, tuple):
            L.append(f"MISSING VALUES {v.name} ({spec[0]:g} THRU {spec[1]:g}).")
        else:
            L.append(f"MISSING VALUES {v.name} ({', '.join(f'{s:g}' for s in spec)}).")
    if any(parse_missing(getattr(v, "missing", "")) for v in variables):
        L.append("")

    if mrsets:
        groups: dict[str, list[Var]] = {}
        for v in variables:
            if "복수응답" in v.note and (m := re.match(r"^(.*)_\d+$", v.name)):
                groups.setdefault(m.group(1), []).append(v)
        for base, members in groups.items():
            if len(members) < 2:
                continue
            stem = members[0].label.split(" - ")[0]
            names = " ".join(v.name for v in members)
            label = q(byte_trim(stem, MAX_VARLABEL_BYTES))
            # 0/1 더미는 다중이분(MDGROUP), 보기코드 저장은 다중범주(MCGROUP).
            # MCGROUP 은 VALUE·CATEGORYLABELS 를 받지 않는다.
            if any("보기코드" in (v.note or "") for v in members):
                L += [
                    "MRSETS",
                    f"  /MCGROUP NAME=${base}",
                    f"    LABEL={label}",
                    f"    VARIABLES={names}",
                    f"  /DISPLAY NAME=[${base}].",
                    "",
                ]
            else:
                L += [
                    "MRSETS",
                    f"  /MDGROUP NAME=${base}",
                    f"    LABEL={label}",
                    "    CATEGORYLABELS=VARLABELS",
                    f"    VARIABLES={names}",
                    "    VALUE=1",
                    f"  /DISPLAY NAME=[${base}].",
                    "",
                ]

    L.append("* 끝. (라벨 명령은 즉시 적용되므로 EXECUTE가 필요하지 않습니다.)")
    return "\n".join(L) + "\n"


def write_syntax(variables: list[Var], path: str | Path, source: str = "",
                 value_style: str = "numbered") -> Path:
    path = Path(path)
    # SPSS 한글 호환을 위해 BOM 포함 UTF-8로 저장
    path.write_text(build_syntax(variables, source=source, value_style=value_style),
                    encoding="utf-8-sig")
    return path


# ------------------------------------------------------------------ .sav

def load_data(path: str | Path) -> pd.DataFrame:
    """원자료 읽기 (csv / xlsx / sav). 한글 csv는 utf-8 -> cp949 순으로 시도."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext in (".csv", ".txt", ".tsv"):
        sep = "\t" if ext == ".tsv" else ","
        for enc in ("utf-8-sig", "cp949", "euc-kr", "utf-8"):
            try:
                return pd.read_csv(path, sep=sep, encoding=enc)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("csv", b"", 0, 1, "인코딩 판별 실패 (utf-8/cp949 아님)")
    if ext in (".xlsx", ".xlsm", ".xls"):
        return pd.read_excel(path)
    if ext == ".sav":
        df, _ = pyreadstat.read_sav(path, apply_value_formats=False)
        return df
    raise ValueError(f"지원하지 않는 원자료 형식: {ext}")


def write_sav(variables: list[Var], path: str | Path,
              data: pd.DataFrame | None = None,
              value_style: str = "numbered") -> tuple[Path, dict[str, list[str]]]:
    """라벨이 적용된 .sav 저장. data가 없으면 0케이스 템플릿을 만든다.

    반환: (경로, {'labeled': [...], 'missing_in_data': [...], 'unlabeled_in_data': [...]})
    """
    path = Path(path)
    names = [v.name for v in variables]
    report: dict[str, list[str]] = {"labeled": [], "missing_in_data": [], "unlabeled_in_data": []}

    if data is None:
        cols = {}
        for v in variables:
            cols[v.name] = pd.Series(dtype="object" if v.vtype == "string" else "float64")
        df = pd.DataFrame(cols)
        report["labeled"] = names
        use = variables
    else:
        df = data.copy()
        lower = {c.lower(): c for c in df.columns}
        use = []
        for v in variables:
            col = lower.get(v.name.lower())
            if col is None:
                report["missing_in_data"].append(v.name)
                continue
            if col != v.name:
                df = df.rename(columns={col: v.name})
            if v.vtype == "string":
                df[v.name] = df[v.name].astype("object")
            else:
                df[v.name] = pd.to_numeric(df[v.name], errors="coerce")
            report["labeled"].append(v.name)
            use.append(v)
        matched = {v.name for v in use}
        report["unlabeled_in_data"] = [str(c) for c in df.columns if c not in matched]
        # 코드북에 없는 열(ID, 가중치 등)도 라벨 없이 그대로 보존
        df = df[[v.name for v in use] + report["unlabeled_in_data"]]

    labels_by_col = {v.name: byte_trim(v.label or v.name, MAX_VARLABEL_BYTES) for v in use}
    column_labels = [labels_by_col.get(str(c), "") for c in df.columns]
    # .sps 와 .sav 의 값라벨 문구를 항상 같게 맞춘다
    value_labels = {v.name: styled_values(v, value_style) for v in use if v.values}
    measures = {v.name: (v.measure if v.measure in ("nominal", "ordinal", "scale") else "nominal")
                for v in use}
    missing_ranges: dict[str, list[dict[str, float]]] = {}
    for v in use:
        spec = parse_missing(getattr(v, "missing", ""))
        if not spec or v.vtype == "string":
            continue
        if isinstance(spec, tuple):
            missing_ranges[v.name] = [{"lo": spec[0], "hi": spec[1]}]
        else:
            missing_ranges[v.name] = [{"lo": s, "hi": s} for s in spec]

    # 주관식 변수는 폭을 넉넉히 지정 (기본값은 A1이 되어 응답이 잘림)
    fmts = {v.name: ("A2000" if v.vtype == "string" else "F8.2") for v in use}
    fmts = {k: val for k, val in fmts.items() if k in df.columns}

    pyreadstat.write_sav(
        df, str(path),
        column_labels=column_labels,
        variable_format=fmts,
        variable_value_labels=value_labels,
        variable_measure=measures,
        missing_ranges=missing_ranges or None,
    )
    return path, report



# ==============================================================================
# 4. 라벨 검수 (코드북 정합성 / 데이터 대조)
# ==============================================================================

# 값라벨이 없어도 정상인 유형 (숫자·문자 입력)
NO_LABEL_KINDS = {"entry", "text", "numeric", "slider", "rank_entry", "dp_instruction"}

SEVERITY_ORDER = {"오류": 0, "주의": 1, "정보": 2}


@dataclass
class Issue:
    severity: str      # 오류 | 주의 | 정보
    variable: str
    kind: str          # 검사 항목
    detail: str


@dataclass
class Report:
    issues: list[Issue] = field(default_factory=list)
    checked: int = 0
    matched: int = 0            # 데이터와 매칭된 변수 수
    data_rows: int = 0

    def add(self, severity: str, variable: str, kind: str, detail: str) -> None:
        self.issues.append(Issue(severity, variable, kind, detail))

    def counts(self) -> dict[str, int]:
        out = {"오류": 0, "주의": 0, "정보": 0}
        for i in self.issues:
            out[i.severity] = out.get(i.severity, 0) + 1
        return out

    def to_frame(self) -> pd.DataFrame:
        rows = sorted(self.issues, key=lambda i: (SEVERITY_ORDER.get(i.severity, 9), i.variable))
        return pd.DataFrame(
            [{"심각도": i.severity, "변수": i.variable, "검사항목": i.kind, "내용": i.detail}
             for i in rows],
            columns=["심각도", "변수", "검사항목", "내용"],
        )


# ------------------------------------------------------------ 코드북 자체 검사

def multi_groups(variables: list[Var]) -> dict[str, list[Var]]:
    """복수응답 세트를 접두사로 묶는다."""
    groups: dict[str, list[Var]] = {}
    for v in variables:
        if v.kind == "multi" and (m := re.match(r"^(.*)_\d+$", v.name)):
            groups.setdefault(m.group(1), []).append(v)
    return groups


def check_codebook(variables: list[Var], report: Report | None = None,
                   check_gaps: bool = True) -> Report:
    """데이터 없이 코드북만 보고 잡을 수 있는 문제.

    check_gaps: 코드 불연속 검사. 데이터 대조를 함께 할 때는 실제 사용값으로
    판단하는 편이 정확하므로 끌 수 있다.
    """
    rep = report or Report()
    rep.checked = len(variables)

    seen: dict[str, int] = {}
    for v in variables:
        seen[v.name] = seen.get(v.name, 0) + 1
    for name, n in seen.items():
        if n > 1:
            rep.add("오류", name, "변수명 중복", f"같은 변수명이 {n}번 나옵니다")

    for v in variables:
        if not v.label.strip():
            rep.add("오류", v.name, "라벨 없음", "변수라벨이 비어 있습니다")
        if len(v.label.encode("utf-8")) > 256:
            rep.add("오류", v.name, "라벨 길이", "변수라벨이 256바이트를 넘습니다")

        for code, lab in v.values.items():
            if len(lab.encode("utf-8")) > 120:
                rep.add("오류", v.name, "값라벨 길이",
                        f"코드 {code} 라벨이 120바이트를 넘습니다 (SPSS 한도)")

        if v.vtype == "string" and v.values:
            rep.add("오류", v.name, "자료형 불일치",
                    "문자형 변수에 값라벨이 있습니다 (숫자형이어야 함)")

        if (not any(str(x).strip() for x in v.values.values())
                and v.kind not in NO_LABEL_KINDS and v.measure in ("nominal", "ordinal")):
            rep.add("주의", v.name, "값라벨 없음",
                    f"{v.measure} 변수인데 값라벨이 없습니다")

        # 같은 라벨이 두 코드에 붙은 경우 (보기 복사 실수)
        labels_only = [lab for lab in v.values.values() if str(lab).strip()]
        dupes = [lab for lab, n in pd.Series(labels_only).value_counts().items()
                 if n > 1] if labels_only else []
        for lab in dupes[:2]:
            rep.add("주의", v.name, "값라벨 중복", f"'{lab[:30]}' 이 두 코드에 붙어 있습니다")

        # 코드가 끊긴 경우 (보기 누락 의심). 결측 코드(90 이상)는 제외.
        #
        # 척도 문항은 양 끝과 중간에만 라벨을 붙이는 것이 정상이라(예: 1·4·7),
        # 빈 코드가 많으면 오히려 정상이다. 한두 개만 빠진 경우가 진짜 누락 신호다.
        codes = sorted(c for c in v.values if c < 90)
        if check_gaps and len(codes) >= 4:
            gaps = [c for c in range(codes[0], codes[-1]) if c not in codes]
            if 0 < len(gaps) <= 2:
                rep.add("주의", v.name, "코드 불연속",
                        f"코드 {gaps} 이(가) 빠져 있습니다 (보기 누락인지 확인)")

        if "확인필요" in (v.note or ""):
            rep.add("주의", v.name, "검수 미완료", v.note[:80])

        for spec in re.split(r"[,\s]+", (v.missing or "").strip()):
            if spec and re.fullmatch(r"-?\d+", spec) and int(spec) not in v.values:
                rep.add("정보", v.name, "결측 코드",
                        f"결측 {spec} 에 대응하는 값라벨이 없습니다")

    # 복수응답 세트 안에서 방식이 섞이면 안 된다
    for base, members in multi_groups(variables).items():
        styles = {"더미(0/1)" if "0/1" in (m.note or "") else "보기코드" for m in members}
        if len(styles) > 1:
            rep.add("오류", base + "_*", "복수응답 방식 혼재",
                    f"한 세트 안에 {', '.join(sorted(styles))} 가 섞여 있습니다")
    return rep


# ------------------------------------------------------------ 데이터 대조 검사

def check_against_data(variables: list[Var], data: pd.DataFrame,
                       report: Report | None = None,
                       max_examples: int = 5) -> Report:
    """실제 데이터와 대조. 라벨 검수의 핵심."""
    rep = report or Report()
    rep.data_rows = len(data)
    lower = {str(c).lower(): c for c in data.columns}
    used: set[str] = set()
    absent: list[str] = []

    for v in variables:
        col = lower.get(v.name.lower())
        if col is None:
            absent.append(v.name)
            continue
        used.add(str(col))
        rep.matched += 1
        series = data[col]

        if v.vtype == "string":
            continue

        numeric = pd.to_numeric(series, errors="coerce")
        bad = series.notna() & numeric.isna()
        if bad.any():
            samples = series[bad].astype(str).unique()[:3]
            rep.add("오류", v.name, "숫자형에 문자값",
                    f"{int(bad.sum())}건: {', '.join(samples)}")

        present = numeric.dropna()
        if present.empty:
            rep.add("정보", v.name, "전부 결측", "데이터에 유효값이 없습니다")
            continue

        # 결측으로 지정한 코드는 검사 대상에서 제외
        missing_codes = {int(m) for m in re.split(r"[,\s]+", (v.missing or "").strip())
                         if m and re.fullmatch(r"-?\d+", m)}

        if v.values:
            defined = set(v.values) | missing_codes
            actual = {int(x) for x in present.unique() if float(x).is_integer()}
            undefined = sorted(actual - defined)
            if undefined:
                # 가장 중요한 신호. 코드가 밀렸거나 척도 시작값이 틀린 경우 여기서 잡힌다.
                counts = {c: int((present == c).sum()) for c in undefined[:max_examples]}
                rep.add("오류", v.name, "값라벨에 없는 값",
                        f"코드 {list(counts)} 이(가) 데이터에 있으나 라벨이 없습니다 "
                        f"(건수 {list(counts.values())})")
                # 척도 전체가 한 칸 밀린 전형적인 패턴을 따로 알려준다.
                # (0부터 코딩된 척도를 1부터로 라벨했거나 그 반대)
                lo_lab, hi_lab = min(v.values), max(v.values)
                lo_dat, hi_dat = int(present.min()), int(present.max())
                shift = lo_dat - lo_lab
                if shift in (1, -1) and hi_dat - hi_lab == shift:
                    rep.add("오류", v.name, "코드 밀림 의심",
                            f"데이터 {lo_dat}~{hi_dat} vs 라벨 {lo_lab}~{hi_lab} — "
                            f"척도 전체가 {abs(shift)}칸 밀렸습니다 (base0 옵션 확인)")

            unused = sorted(set(v.values) - actual - missing_codes)
            if unused and len(unused) < len(v.values):
                rep.add("정보", v.name, "사용되지 않은 코드",
                        f"코드 {unused[:8]} 이(가) 데이터에 없습니다")

        if v.kind == "multi":
            dummy = "0/1" in (v.note or "")
            vals = {int(x) for x in present.unique() if float(x).is_integer()}
            if dummy and vals - {0, 1}:
                rep.add("오류", v.name, "복수응답 방식 불일치",
                        f"0/1 더미로 라벨했으나 데이터에 {sorted(vals - {0, 1})[:5]} 가 있습니다")
            if not dummy and vals and vals <= {0, 1}:
                rep.add("오류", v.name, "복수응답 방식 불일치",
                        "보기 코드로 라벨했으나 데이터는 0/1 더미로 보입니다")

    # 없는 변수가 20개를 넘으면 개별 항목 대신 한 줄로 묶는다.
    # (일부 문항만 담긴 데이터로 검수하는 경우가 잦다)
    if len(absent) > 20:
        rep.add("주의", f"({len(absent)}개)", "데이터에 없음",
                "코드북에 있으나 데이터에 없는 변수: "
                + ", ".join(absent[:20]) + f" … 외 {len(absent) - 20}개")
    else:
        for name in absent:
            rep.add("주의", name, "데이터에 없음", "코드북에 있으나 데이터에 열이 없습니다")

    extra = [str(c) for c in data.columns if str(c) not in used]
    for col in extra[:30]:
        rep.add("정보", col, "코드북에 없음", "데이터에 있으나 코드북에 없는 열입니다")
    if len(extra) > 30:
        rep.add("정보", "(외)", "코드북에 없음", f"그 외 {len(extra) - 30}개 열")
    return rep


def verify_sav(sav_bytes: bytes, variables: list[Var],
               report: Report | None = None) -> Report:
    """생성된 .sav 를 되읽어 라벨이 실제로 들어갔는지 확인 (왕복 검증)."""
    import tempfile
    from pathlib import Path

    import pyreadstat

    rep = report or Report()
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "check.sav"
        path.write_bytes(sav_bytes)
        _, meta = pyreadstat.read_sav(str(path), user_missing=True)

    labels = meta.column_names_to_labels
    value_labels = meta.variable_value_labels
    for v in variables:
        if v.name not in labels:
            rep.add("오류", v.name, "sav 누락", ".sav 에 변수가 없습니다")
            continue
        got = labels.get(v.name) or ""
        if v.label and got.strip() != v.label.strip():
            rep.add("주의", v.name, "sav 라벨 불일치",
                    f"코드북 '{v.label[:30]}' vs sav '{got[:30]}'")
        if v.values:
            saved = value_labels.get(v.name, {})
            if len(saved) != len(v.values):
                rep.add("오류", v.name, "sav 값라벨 개수",
                        f"코드북 {len(v.values)}개 vs sav {len(saved)}개")
    return rep


def validate(variables: list[Var], data: pd.DataFrame | None = None,
             sav_bytes: bytes | None = None) -> Report:
    """전체 검수. 데이터가 있으면 대조까지, .sav 가 있으면 왕복 검증까지."""
    rep = check_codebook(variables, check_gaps=data is None)
    if data is not None:
        check_against_data(variables, data, rep)
    if sav_bytes is not None:
        verify_sav(sav_bytes, variables, rep)
    return rep



# ==============================================================================
# 5. 화면용 bytes 입출력 계층
# ==============================================================================

FIELDS = ["변수명", "문항번호", "문항유형", "변수라벨", "유형", "측도", "값라벨", "결측값", "비고"]


# ------------------------------------------------------------- .doc 변환

class LegacyDocError(RuntimeError):
    """구형 .doc 파일을 변환할 수 없을 때."""


def looks_like_legacy_doc(data: bytes, filename: str = "") -> bool:
    """OLE2 복합문서(.doc) 서명 확인. 확장자만으로는 판별할 수 없다."""
    return data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" or (
        filename.lower().endswith(".doc") and data[:2] != b"PK")


def convert_doc_to_docx(data: bytes) -> bytes:
    """LibreOffice(soffice)로 .doc -> .docx 변환.

    python-docx 는 Word 97-2003 이진 포맷을 읽지 못한다. 서버에 soffice 가
    설치돼 있으면 자동 변환하고, 없으면 사용자에게 직접 변환을 안내한다.
    """
    import shutil
    import subprocess

    exe = shutil.which("soffice") or shutil.which("libreoffice")
    if not exe:
        raise LegacyDocError(
            "구형 .doc 파일입니다. 서버에 LibreOffice 가 없어 자동 변환할 수 없습니다. "
            "워드에서 '다른 이름으로 저장 > .docx' 로 변환한 뒤 올려 주세요."
        )
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "input.doc"
        src.write_bytes(data)
        try:
            subprocess.run([exe, "--headless", "--convert-to", "docx",
                            "--outdir", tmp, str(src)],
                           check=True, capture_output=True, timeout=180)
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            raise LegacyDocError(f"LibreOffice 변환 실패: {exc}") from exc
        out = Path(tmp) / "input.docx"
        if not out.exists():
            raise LegacyDocError("변환 결과 파일이 생성되지 않았습니다.")
        return out.read_bytes()


# ------------------------------------------------------------- 파싱

def parse_upload(docx_bytes: bytes, base0: list[str] | None = None,
                 full_labels: bool = False, multi_style: str = "category") -> list[Var]:
    """업로드된 워드 파일 bytes -> 변수 목록.

    multi_style: 복수응답 저장 방식 (category / position / dichotomy).
    """
    if looks_like_legacy_doc(docx_bytes):
        docx_bytes = convert_doc_to_docx(docx_bytes)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        variables = parse_docx(tmp.name, colon_split=not full_labels,
                               multi_style=multi_style)
    if base0:
        rebase_values(variables, base0, start=0)
    return variables


# ------------------------------------------------------------- 표 <-> 변수

def vars_to_frame(variables: list[Var]) -> pd.DataFrame:
    """편집 가능한 표로 변환 (st.data_editor에 그대로 넘길 수 있음)."""
    rows = [{
        "변수명": v.name, "문항번호": v.qid, "문항유형": v.kind, "변수라벨": v.label,
        "유형": v.vtype, "측도": v.measure, "값라벨": values_to_text(v.values),
        "결측값": v.missing, "비고": v.note,
    } for v in variables]
    return pd.DataFrame(rows, columns=FIELDS)


def frame_to_vars(df: pd.DataFrame) -> list[Var]:
    """편집된 표 -> 변수 목록. 변수명이 빈 행은 버린다."""
    out: list[Var] = []
    for rec in df.to_dict("records"):
        def get(key: str) -> str:
            val = rec.get(key)
            return "" if val is None or pd.isna(val) else str(val).strip()

        name = get("변수명")
        if not name:
            continue
        out.append(Var(
            name=name,
            label=get("변수라벨"),
            vtype=(get("유형") or "numeric").lower(),
            measure=(get("측도") or "nominal").lower(),
            values=text_to_values(get("값라벨")),
            qid=get("문항번호"),
            kind=get("문항유형"),
            note=get("비고"),
            missing=get("결측값"),
        ))
    return out


# ------------------------------------------------------------- 산출물 bytes

def codebook_bytes(variables: list[Var]) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "codebook.xlsx"
        write_codebook(variables, path)
        return path.read_bytes()


def codebook_upload_to_vars(xlsx_bytes: bytes) -> list[Var]:
    """검수 완료된 코드북을 다시 올려 이어서 작업할 때."""
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "codebook.xlsx"
        path.write_bytes(xlsx_bytes)
        return read_codebook(path)


def syntax_bytes(variables: list[Var], source: str = "",
                 value_style: str = "numbered") -> bytes:
    # SPSS 한글 호환을 위해 BOM 포함
    return build_syntax(variables, source=source,
                        value_style=value_style).encode("utf-8-sig")


def sav_bytes(variables: list[Var], data: pd.DataFrame | None = None,
              value_style: str = "numbered") -> tuple[bytes, dict[str, list[str]]]:
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "out.sav"
        _, report = write_sav(variables, path, data=data, value_style=value_style)
        return path.read_bytes(), report


def read_data_upload(file_bytes: bytes, filename: str) -> pd.DataFrame:
    """원자료 업로드(csv/xlsx/sav) -> DataFrame."""
    suffix = Path(filename).suffix.lower() or ".csv"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=True) as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        return load_data(tmp.name)


# ------------------------------------------------------------- 요약

def summarize(variables: list[Var]) -> dict[str, Any]:
    kinds: dict[str, int] = {}
    for v in variables:
        kinds[v.kind or "-"] = kinds.get(v.kind or "-", 0) + 1
    todo = [v.name for v in variables if "확인필요" in (v.note or "")]
    return {
        "total": len(variables),
        "kinds": dict(sorted(kinds.items())),
        "with_labels": sum(1 for v in variables if v.values),
        "todo": todo,
    }


def zip_bytes(files: dict[str, bytes]) -> bytes:
    """여러 산출물을 한 번에 내려주고 싶을 때."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, blob in files.items():
            zf.writestr(name, blob)
    return buf.getvalue()
