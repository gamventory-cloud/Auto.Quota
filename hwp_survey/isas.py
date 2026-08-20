# -*- coding: utf-8 -*-
"""ISAS 표준 설문지 서식.

납품본(.docx)에서 실측한 값을 그대로 쓴다.
  * 나눔고딕 9pt, 여백 좌우 1.91cm / 상하 2.54cm (본문 폭 17.17cm)
  * 행별 표: 문항열 8.82cm + 보기열 1.67cm × 5, 첫 칸은 비움
  * 보기 머리칸: 라벨을 어절마다 줄바꿈한 뒤 마지막 줄에 보기 코드
  * 표 안쪽은 보기 코드(1 2 3 …)로 모두 채움
  * [PROG: …] 파란색, [DATA: …] 초록색, 미분류 내용 빨간색

문항 번호는 ISAS 관행을 따른다.
  선별 문항 -> SQ1, SQ2 …      동의 문항 -> AQ
  첫 PART   -> Q1-1, Q1-2 …    이후      -> Q2, Q3 …
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .dp import RE_LABEL, RE_TAG, items_to_dp_dsl, parse_dp, summarize_dp

#: DP 태그 -> ISAS 표기 (납품본은 숫자와 '개' 사이를 띄운다)
TAGS = {
    "1개선택": "1개 선택",
    "모두선택": "모두 선택",
    "행별 1개선택": "행별 1개 선택",
    "행별1개선택": "행별 1개 선택",
    "출생년도 입력": "출생연도로 응답",
    "지도에서 선택": "1개 선택",
    "표 응답": "표 응답",
}
RE_PART = re.compile(r"^##\s*(?:PART|파트)\s*(\d+)", re.I)
RE_SCREEN = re.compile(r"선별|선정|스크리닝|screen", re.I)
RE_DEMO = re.compile(r"배경|인적\s*사항|demo|배문", re.I)


# =====================================================================
# 1) 중간 텍스트 (DP 문법을 그대로 쓰되 번호와 태그를 ISAS 관행으로)
# =====================================================================
def items_to_isas_dsl(items, add_matrix_hint=False, **kwargs) -> str:
    # ISAS 규칙 1)은 원문 텍스트를 그대로 쓰라고 한다 -> 안내 문장을 덧붙이지 않는다
    dsl = items_to_dp_dsl(items, add_matrix_hint=add_matrix_hint, **kwargs)
    return renumber_isas(retag_isas(dsl))


def retag_isas(dsl: str) -> str:
    out = []
    for line in dsl.splitlines():
        tag = RE_TAG.search(line.rstrip())
        if tag and (RE_LABEL.match(line) or line.startswith("?.")):
            name = tag.group(1).strip()
            out.append(line[: tag.start()].rstrip() + f" [{TAGS.get(name, name)}]")
        else:
            out.append(line)
    return "\n".join(out)


def renumber_isas(dsl: str) -> str:
    """ISAS 번호 체계로 다시 매긴다.

      * 첫 구역(선별/스크리닝) 이전·안쪽 문항 -> SQ1, SQ2 …
      * 참여 동의 문항                        -> AQ
      * 본문 구역 k번째                       -> 문항이 여럿이면 Qk-1, Qk-2 …
                                               하나면 Qk
      * 배경/인적사항 구역                    -> DQ1, DQ2 …
    """
    lines = dsl.splitlines()
    if not any(l.startswith("##") for l in lines):
        return renumber_without_sections(lines)

    # 1) 구역 경계와 구역별 문항 수를 먼저 센다
    zones = []            # (시작줄, 종류, 문항 수)
    kind = "screen"
    start, count = 0, 0
    for n, line in enumerate(lines):
        if line.startswith("##"):
            zones.append((start, kind, count))
            name = line.lstrip("# ").strip()
            kind = ("demo" if RE_DEMO.search(name) else
                    "screen" if RE_SCREEN.search(name) else "body")
            start, count = n, 0
        elif is_question_line(line):
            count += 1
    zones.append((start, kind, count))

    # 2) 줄마다 번호를 매긴다
    zone_of = {}
    for i, (begin, zkind, cnt) in enumerate(zones):
        zone_of[begin] = (i, zkind, cnt)

    sq = dq = body_no = 0
    sub = 0
    cur_kind, cur_cnt = "screen", 0
    consent_done = False

    for n, line in enumerate(lines):
        if n in zone_of:
            _, cur_kind, cur_cnt = zone_of[n]
            if cur_kind == "body" and cur_cnt:
                body_no += 1
                sub = 0
        if not is_question_line(line):
            continue

        m = RE_LABEL.match(line)
        body = line[2:].strip() if line.startswith("?.") else m.group(2).strip()

        if not consent_done and "응답해 주시기" in body and "동의" not in body:
            lines[n] = f"AQ. {body}"              # 참여 동의 문항
            consent_done = True
            continue
        if cur_kind == "demo":
            dq += 1
            lines[n] = f"DQ{dq}. {body}"
        elif cur_kind == "body":
            no = body_no or 1
            if cur_cnt > 1:
                sub += 1
                lines[n] = f"Q{no}-{sub}. {body}"
            else:
                lines[n] = f"Q{no}. {body}"
        else:
            sq += 1
            lines[n] = f"SQ{sq}. {body}"
    return "\n".join(lines)


def renumber_without_sections(lines) -> str:
    """구역 표시가 없는 설문지: 첫 행별 문항 앞은 SQ, 그 뒤는 Q."""
    first_matrix = next((n for n, l in enumerate(lines)
                         if is_question_line(l) and "행별" in l), None)
    sq = q = 0
    for n, line in enumerate(lines):
        if not is_question_line(line):
            continue
        m = RE_LABEL.match(line)
        body = line[2:].strip() if line.startswith("?.") else m.group(2).strip()
        if first_matrix is not None and n >= first_matrix:
            q += 1
            lines[n] = f"Q{q}. {body}"
        else:
            sq += 1
            lines[n] = f"SQ{sq}. {body}"
    return "\n".join(lines)


def is_question_line(line: str) -> bool:
    return bool(RE_TAG.search(line.rstrip())
                and (RE_LABEL.match(line) or line.startswith("?.")))


# =====================================================================
# 2) 문서 생성
# =====================================================================
class ISASWriter:
    BLUE, GREEN, RED, GRAY = "0000FF", "006600", "FF0000", "7F7F7F"
    FILL = "F2F2F2"

    def __init__(self, font="나눔고딕", base_pt=9.0, title_pt=12.0,
                 content_cm=17.17, row_label_cm=8.82, spec_label_cm=2.24,
                 doc_header=""):
        self.font_name = font
        self.base = Pt(base_pt)
        self.title_pt = title_pt
        self.content_cm = content_cm
        self.row_label_cm = row_label_cm
        self.spec_label_cm = spec_label_cm
        self.doc_header = doc_header
        self.doc = self._new_document()

    # ------------------------------------------------------------ 기본
    def _new_document(self):
        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = self.font_name
        normal.font.size = self.base
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), self.font_name)
        pf = normal.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.15

        sec = doc.sections[0]
        sec.left_margin = sec.right_margin = Cm(1.91)
        sec.top_margin = sec.bottom_margin = Cm(2.54)
        return doc

    def run(self, paragraph, text, bold=False, color=None, size=None):
        r = paragraph.add_run(text)
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            rFonts.set(qn(attr), self.font_name)
        r.font.size = size or self.base
        r.font.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        return r

    def para(self, text="", bold=False, color=None, align=None, size=None,
             after=0):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        if align is not None:
            p.alignment = align
        if text:
            self.run(p, text, bold=bold, color=color, size=size)
        return p

    @staticmethod
    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {})
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    @staticmethod
    def fix_layout(table, header_repeat=True):
        tblPr = table._tbl.tblPr
        layout = tblPr.makeelement(qn("w:tblLayout"), {})
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        for i, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
            if header_repeat and i == 0:
                trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))

    def cell_text(self, cell, text, bold=False, color=None, align=None,
                  size=None, valign=None):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if align is not None:
            p.alignment = align
        for n, line in enumerate(str(text).split("\n")):
            target = p if n == 0 else cell.add_paragraph()
            if n:
                target.paragraph_format.space_after = Pt(0)
                target.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
            self.run(target, line, bold=bold, color=color, size=size)
        if valign is not None:
            cell.vertical_alignment = valign
        return cell

    def widths(self, table, widths):
        for i, col in enumerate(table.columns):
            col.width = widths[i]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths):
                    cell.width = widths[i]

    def page_furniture(self, title):
        """머리말(문서 제목)과 꼬리말(페이지 번호)."""
        sec = self.doc.sections[0]
        if title:
            p = sec.header.paragraphs[0]
            self.run(p, title, bold=True, color=self.GRAY, size=Pt(9))
        foot = sec.footer.paragraphs[0]
        foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.run(foot, "페이지 ", size=Pt(9))
        self._page_field(foot, "PAGE")
        self.run(foot, " / ", size=Pt(9))
        self._page_field(foot, "NUMPAGES")

    def _page_field(self, paragraph, code):
        r = self.run(paragraph, "", size=Pt(9))._element
        for kind, text in (("begin", None), (None, code), ("end", None)):
            if kind:
                el = r.makeelement(qn("w:fldChar"), {})
                el.set(qn("w:fldCharType"), kind)
            else:
                el = r.makeelement(qn("w:instrText"), {})
                el.set(qn("xml:space"), "preserve")
                el.text = f" {text} "
            r.append(el)

    # ------------------------------------------------------------ 구성 요소
    def title_box(self, title):
        table = self.doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        self.cell_text(cell, title, bold=True, size=Pt(self.title_pt),
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       valign=WD_ALIGN_VERTICAL.CENTER)
        self.widths(table, [Cm(self.content_cm)])
        self.fix_layout(table, header_repeat=False)
        self.para(after=6)

    def spec_table(self, doc):
        rows = [("조사대상", doc["대상자"]), ("샘플 수(명)", doc["샘플수"])]
        if doc["쿼터"] or doc["쿼터표"]:
            rows.append(("할당", doc["쿼터"]))
        if not any(v for _, v in rows) and not doc["쿼터표"]:
            return
        table = self.doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        for r, (label, value) in enumerate(rows):
            head, body = table.rows[r].cells
            self.cell_text(head, label, valign=WD_ALIGN_VERTICAL.CENTER)
            self.shade(head, self.FILL)
            self.cell_text(body, value)
            if label == "할당" and doc["쿼터표"]:
                self.quota_table(body, doc["쿼터표"])
        self.widths(table, [Cm(self.spec_label_cm),
                            Cm(self.content_cm - self.spec_label_cm)])
        self.fix_layout(table, header_repeat=False)
        self.para(after=6)

    def quota_table(self, cell, rows):
        cols = max(len(r) for r in rows)
        inner = cell.add_table(rows=0, cols=cols)
        inner.style = "Table Grid"
        width = Cm((self.content_cm - self.spec_label_cm - 0.3) / cols)
        for n, row in enumerate(rows):
            cells = inner.add_row().cells
            for c in range(cols):
                self.cell_text(cells[c], row[c] if c < len(row) else "",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
                cells[c].width = width
                if n == 0:
                    self.shade(cells[c], self.FILL)
        for col in inner.columns:
            col.width = width
        self.fix_layout(inner, header_repeat=False)

    def box(self, texts):
        table = self.doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        for n, text in enumerate(texts):
            p = cell.paragraphs[0] if n == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            self.run(p, text)
        self.widths(table, [Cm(self.content_cm)])
        self.fix_layout(table, header_repeat=False)
        self.para(after=6)

    def matrix(self, block):
        scale = block["scale"] or ["1", "2", "3", "4", "5"]
        rows = block["options"] or [{"type": "row", "text": "항목"}]
        first = Cm(self.row_label_cm)
        rest = Cm((self.content_cm - self.row_label_cm) / len(scale))

        table = self.doc.add_table(rows=1, cols=len(scale) + 1)
        table.style = "Table Grid"
        self.cell_text(table.rows[0].cells[0], "")          # 첫 칸은 비운다
        for i, label in enumerate(scale, 1):
            wrapped = "\n".join(str(label).split())          # 어절마다 줄바꿈
            self.cell_text(table.rows[0].cells[i], f"{wrapped}\n{i}",
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           valign=WD_ALIGN_VERTICAL.BOTTOM)

        for item in rows:
            cells = table.add_row().cells
            if item["type"] == "group":          # 규칙 2-5: 대분류는 병합 + 빨간색
                merged = cells[0].merge(cells[-1])
                self.cell_text(merged, item["text"], bold=True, color=self.RED)
                continue
            self.cell_text(cells[0], item["text"],
                           valign=WD_ALIGN_VERTICAL.CENTER)
            for i in range(1, len(scale) + 1):
                self.cell_text(cells[i], str(i),             # 보기 코드로 채움
                               align=WD_ALIGN_PARAGRAPH.CENTER,
                               valign=WD_ALIGN_VERTICAL.CENTER)

        self.widths(table, [first] + [rest] * len(scale))
        self.fix_layout(table)
        self.para(after=6)

    def grid(self, rows):
        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=0, cols=cols)
        table.style = "Table Grid"
        width = Cm(self.content_cm / cols)
        for n, row in enumerate(rows):
            cells = table.add_row().cells
            for c in range(cols):
                self.cell_text(cells[c], row[c] if c < len(row) else "",
                               align=WD_ALIGN_PARAGRAPH.CENTER)
                cells[c].width = width
                if n == 0:
                    self.shade(cells[c], self.FILL)
        for col in table.columns:
            col.width = width
        self.fix_layout(table, header_repeat=False)
        self.para(after=6)

    def options(self, opts):
        """보기 두 개면 한 줄에, 그보다 많으면 한 줄에 하나씩."""
        rows = [o for o in opts if o["type"] == "row"]
        labelled = [o["text"] if re.match(r"^\d{1,4}[.)]", o["text"])
                    else f"{n}) {o['text']}" for n, o in enumerate(rows, 1)]
        if len(labelled) == 2 and all(len(t) <= 12 for t in labelled):
            self.para("\t\t\t".join(labelled))
        else:
            for text in labelled:
                self.para(text)

    # ------------------------------------------------------------ 본문
    def write(self, doc):
        title = doc["제목"] or "OO조사 설문지"
        self.page_furniture(self.doc_header or title)
        self.title_box(title)
        self.spec_table(doc)

        box_buffer: list[str] = []

        def flush_box():
            if box_buffer:
                self.box(list(box_buffer))
                box_buffer.clear()

        for b in doc["blocks"]:
            if b["kind"] == "box":
                box_buffer.append(b["text"])
                continue
            if b["kind"] == "note" and box_buffer and len(b["text"]) <= 30:
                box_buffer.append(f"({b['text']})")
                continue
            flush_box()

            if b["kind"] == "question":
                self.para()
                label = f"{b['label']}. " if b["label"] else ""
                stem = self.para(f"{label}{b['text']} [{b['tag']}]")
                if b["tag"].startswith("행별") or b["tag"].startswith("열별"):
                    stem.paragraph_format.keep_with_next = True
                    self.matrix(b)
                elif "연도" in b["tag"] or "수치" in b["tag"]:
                    self.para("(          )")
                else:
                    self.options(b["options"])
            elif b["kind"] == "section":
                self.para()
                self.para(b["text"], bold=True)
            elif b["kind"] == "grid":
                self.grid(b["rows"])
            elif b["kind"] == "prog":
                self.para(f"[PROG: {b['text']}]", color=self.BLUE)
            elif b["kind"] == "verify":
                self.para(f"[DATA: {b['text']}]", color=self.GREEN)
            else:
                self.para(b["text"], color=self.RED)   # 규칙 2-8: 미분류는 빨간색

        flush_box()
        return self

    def save(self, path):
        self.doc.save(path)
        return path

    def to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


def build_isas_docx(doc, **opts) -> bytes:
    return ISASWriter(**opts).write(doc).to_bytes()


def parse_isas(text: str):
    return parse_dp(text)


def summarize_isas(doc) -> dict:
    qs = [b for b in doc["blocks"] if b["kind"] == "question"]
    return {
        "문항": len(qs),
        "선별문항(SQ)": sum(1 for q in qs if q["label"].startswith("SQ")),
        "행별 표": sum(1 for q in qs if q["tag"].startswith(("행별", "열별"))),
        "PROG 지시문": sum(1 for b in doc["blocks"] if b["kind"] == "prog"),
        "일반 표": sum(1 for b in doc["blocks"] if b["kind"] == "grid"),
    }
