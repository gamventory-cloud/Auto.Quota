# -*- coding: utf-8 -*-
"""리서치사 납품용 DP 스크립트 형식.

인쇄용 설문지(writer.py)와 목적이 다르다. DP 스크립트는 조사 개요표,
SQ/Q 문항 번호, `[1개선택]` 같은 응답 방식 태그, 파란색 `[PROG: ...]`
프로그래밍 지시문, 빨간색 데이터 검증 지시문으로 이루어진다.

중간 텍스트(DP DSL)
    @제목: ...
    @대상자: ...
    @샘플수: 250명
    @쿼터: 성별*연령대별 균등할당
    @쿼터표: ,남자,여자
    @쿼터표: 만 20-29세,62,63
    @제외: 2026060452 참여자 제외
    ~ 상자글(용어 정의)
    SQ1. 문항 [1개선택]
    - 보기
    %PROG: 만20~39세만 진행
    Q1. 문항 [행별 1개선택]
    @행별: 전혀 그렇지 않다,그렇지 않다,보통이다,그렇다,매우 그렇다
    - 진술문
    %검증: Q1~Q2 일자찍기는 최종 납품 데이터에서 제외
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .parser import CIRCLED

# ---------------------------------------------------------------- 상수
SCALE_5 = ["전혀 그렇지 않다", "그렇지 않다", "보통이다", "그렇다", "매우 그렇다"]
MATRIX_HINT = "귀하의 의견과 가장 일치하는 정도에 체크해 주세요."

RE_LABEL = re.compile(r"^\s*((?:SQ|Q|DQ|A)\s*\d+(?:-\d+)?)\s*[.)]\s*(.*)$", re.I)
RE_FIELD = re.compile(r"^@(제목|대상자|샘플수|쿼터|쿼터표|제외|행별)\s*:\s*(.*)$")
RE_PROG_SRC = re.compile(r"^\s*\[?\s*PROG\s*[:：]\s*(.+?)\s*\]?\s*$", re.I)
RE_TAG = re.compile(r"\[([^\[\]]+)\]\s*$")
RE_OPT_LINE = re.compile(rf"^\s*(?:[{CIRCLED}]|\(\s*\d+\s*\)|\d{{1,2}}\s*[).])\s*(.+)$")
RE_PAREN_ONLY = re.compile(r"^\s*[(（]\s*(.+?)\s*[)）]\s*$")

ALONE_HINTS = ("이 중 없음", "없음", "이용하지 않음", "해당 없음")
#: 이 보기가 있으면 사실상 복수응답 문항이다(단독선택 보기를 따로 둔 것)
MULTI_SIGNS = ("이 중 없음", "이용하지 않음", "해당 없음")
SINGLE_TAG, MULTI_TAG, MATRIX_TAG = "1개선택", "모두선택", "행별 1개선택"


# =====================================================================
# 1) 추출 결과 -> DP 중간 텍스트
# =====================================================================
def items_to_dp_dsl(items, add_matrix_hint=True, add_alone_prog=True) -> str:
    lines: list[str] = []
    quota_rows: list[list[str]] = []
    state = {"mode": None}          # '쿼터' 수집 중 여부

    flat = [(k, v) for k, v in items]
    i = 0
    while i < len(flat):
        kind, payload = flat[i]
        i += 1

        if kind == "table":                       # 원본이 표로 쿼터를 짠 경우
            quota_rows.extend([c for c in row] for row in payload)
            continue

        text = payload.strip()
        if not text or set(text) <= {"-", "―", "—", "="}:      # 구분선
            continue

        # 머리말: <제목>, ■ 필드
        m = re.match(r"^[<〈《【]\s*(.+?)\s*[>〉》】]$", text)
        if m and not lines:
            lines.append(f"@제목: {m.group(1)}")
            continue
        m = re.match(r"^[■◆●]\s*(.+?)\s*[:：]\s*(.*)$", text)
        if m:
            name, value = m.group(1).strip(), m.group(2).strip()
            key = ("대상자" if "대상" in name else
                   "샘플수" if "샘플" in name else
                   "쿼터" if "쿼터" in name or "할당" in name else None)
            state["mode"] = "쿼터" if key == "쿼터" else None
            if key:
                lines.append(f"@{key}: {value}")
            else:
                lines.append(f"~ {text}")
            continue

        # 쿼터 표를 문단으로 흘려 쓴 경우: 공백/탭으로 나뉜 숫자 행
        if state["mode"] == "쿼터":
            if looks_like_quota_row(text):
                quota_rows.append([c for c in re.split(r"\s+|\t", text.strip()) if c])
                continue
            state["mode"] = None                    # 표가 끝났다

        # 섹션 안내(SQ. 선정질문 등)는 DP 스크립트에서 쓰지 않는다
        if re.match(r"^(SQ|Q)\s*[.)]\s*\S", text) and not RE_LABEL.match(text):
            continue

        prog = RE_PROG_SRC.match(text)
        if prog:
            lines.append(f"%PROG: {prog.group(1)}")
            continue

        m = RE_LABEL.match(text)
        if m:
            label, stem = m.group(1).upper().replace(" ", ""), m.group(2).strip()
            block, i = read_question(flat, i, label, stem,
                                     add_matrix_hint, add_alone_prog)
            lines.append("")
            lines.extend(block)
            continue

        opt = RE_OPT_LINE.match(text)
        if opt:
            lines.append(f"- {opt.group(1).strip()}")
            continue

        paren = RE_PAREN_ONLY.match(text)
        if paren and len(text) <= 30:
            lines.append(f"! {paren.group(1)}")
            continue

        lines.append(f"~ {text}")

    if quota_rows:
        width = max(len(r) for r in quota_rows)
        quota_rows = [([""] * (width - len(r))) + r for r in quota_rows]  # 좌상단 빈 칸
        head = [f"@쿼터표: {','.join(r)}" for r in quota_rows]
        anchor = next((n for n, l in enumerate(lines)
                       if l.startswith("@쿼터:")), len(lines) - 1)
        lines[anchor + 1:anchor + 1] = head

    return "\n".join(collapse(lines)).strip()


def looks_like_quota_row(text: str) -> bool:
    """쿼터 표를 문단으로 흘려 쓴 줄. 머리 행('20대 30대 합계')과 숫자 행 모두.

    문항·보기·PROG 줄을 삼키지 않도록 조건을 좁게 잡는다.
    """
    if "?" in text or len(text) > 40:
        return False
    if RE_LABEL.match(text) or RE_OPT_LINE.match(text) or text.startswith("["):
        return False
    cells = [c for c in re.split(r"\s+|\t", text.strip()) if c]
    if not 2 <= len(cells) <= 6:
        return False
    return all(len(c) <= 6 for c in cells)


def read_question(flat, i, label, stem, add_matrix_hint, add_alone_prog):
    """문항 한 덩어리(보기·척도 안내·PROG)를 읽어 DP DSL 줄로."""
    options: list[str] = []
    progs: list[str] = []
    notes: list[str] = []
    scale: list[str] | None = None
    tag = None

    while i < len(flat):
        kind, payload = flat[i]
        if kind != "p":
            break
        text = payload.strip()
        if not text or set(text) <= {"-", "―", "—", "="}:
            i += 1
            continue
        if RE_LABEL.match(text):
            break

        prog = RE_PROG_SRC.match(text)
        if prog:
            progs.append(prog.group(1))
            i += 1
            continue

        opt = RE_OPT_LINE.match(text)
        if opt:
            options.append(opt.group(1).strip())
            i += 1
            continue

        paren = RE_PAREN_ONLY.match(text)
        if paren:
            inner = paren.group(1)
            if "점 척도" in inner:
                scale = scale_labels(inner)
                tag = MATRIX_TAG
            elif "출생" in inner or "년도" in inner:
                tag = "출생년도 입력"
            elif "동일" in inner or "지도" in inner:
                tag = "지도에서 선택"
                notes.append(inner)
            else:
                notes.append(inner)
            i += 1
            continue
        break                                   # 다음 문단은 문항 바깥(상자글 등)

    if tag is None:
        if re.search(r"복수\s*응답|모두", stem):
            tag = MULTI_TAG
        elif len(options) >= 3 and any(o.strip().startswith(h)
                                       for o in options for h in MULTI_SIGNS):
            tag = MULTI_TAG          # 단독선택용 배타 보기를 둔 문항
        else:
            tag = SINGLE_TAG
    stem = re.sub(r"[(（]\s*복수\s*응답\s*[)）]", "", stem).strip()

    out = []
    if tag == MATRIX_TAG:
        if add_matrix_hint and MATRIX_HINT not in stem:
            stem = f"{stem} {MATRIX_HINT}".replace("  ", " ")
        out.append(f"{label}. {stem} [{tag}]")
        out.append(f"@행별: {','.join(scale or SCALE_5)}")
    else:
        out.append(f"{label}. {stem} [{tag}]")

    out += [f"! {n}" for n in notes]
    out += [f"- {o}" for o in options]

    if add_alone_prog and tag == MULTI_TAG:
        for n, opt in enumerate(options, 1):
            if any(opt.strip().startswith(h) or opt.strip() == h for h in ALONE_HINTS):
                text = f"{n}번 보기는 단독선택만 가능"
                if not any(text in p for p in progs):
                    progs.insert(0, text)
    if tag == "지도에서 선택" and not progs:
        progs.append("17개시도 지도제시")
    out += [f"%PROG: {p}" for p in progs]
    return out, i


def scale_labels(inner: str) -> list[str]:
    n = int(re.search(r"(\d+)\s*점", inner).group(1))
    if n == 5:
        return SCALE_5
    if n == 7:
        return ["전혀 그렇지 않다", "그렇지 않다", "약간 그렇지 않다", "보통이다",
                "약간 그렇다", "그렇다", "매우 그렇다"]
    return [f"{i}" for i in range(1, n + 1)]


def collapse(lines):
    out = []
    for line in lines:
        if not line.strip() and (not out or not out[-1].strip()):
            continue
        out.append(line)
    return out


# =====================================================================
# 2) DP 중간 텍스트 -> 블록
# =====================================================================
def parse_dp(text: str) -> dict:
    doc = {"제목": "", "대상자": "", "샘플수": "", "쿼터": "", "쿼터표": [],
           "제외": "", "blocks": []}
    cur = None

    for raw in text.splitlines():
        s = raw.strip()
        if not s:
            continue

        field = RE_FIELD.match(s)
        if field:
            name, value = field.group(1), field.group(2).strip()
            if name == "쿼터표":
                doc["쿼터표"].append([c.strip() for c in value.split(",")])
            elif name == "행별":
                if cur:
                    cur["scale"] = [c.strip() for c in value.split(",") if c.strip()]
            else:
                doc[name] = value
            continue

        if s.startswith("%PROG:") or s.startswith("%prog:"):
            doc["blocks"].append({"kind": "prog", "text": s.split(":", 1)[1].strip()})
            continue
        if s.startswith("%검증:"):
            doc["blocks"].append({"kind": "verify", "text": s.split(":", 1)[1].strip()})
            continue
        if s.startswith("~"):
            cur = None
            doc["blocks"].append({"kind": "box", "text": s.lstrip("~ ").strip()})
            continue
        if s.startswith("!"):
            doc["blocks"].append({"kind": "note", "text": s.lstrip("! ").strip()})
            continue
        if s.startswith("-"):
            if cur:
                cur["options"].append(s.lstrip("- ").strip())
            continue

        m = RE_LABEL.match(s)
        label, body = (m.group(1).upper(), m.group(2)) if m else ("", s)
        tag = SINGLE_TAG
        t = RE_TAG.search(body)
        if t:
            tag = t.group(1).strip()
            body = body[: t.start()].strip()
        cur = {"kind": "question", "label": label, "text": body.strip(),
               "tag": tag, "options": [], "scale": None}
        doc["blocks"].append(cur)

    return doc


def summarize_dp(doc) -> dict:
    qs = [b for b in doc["blocks"] if b["kind"] == "question"]
    return {
        "문항": len(qs),
        "선정문항(SQ)": sum(1 for q in qs if q["label"].startswith("SQ")),
        "행별 표": sum(1 for q in qs if q["tag"] == MATRIX_TAG),
        "PROG 지시문": sum(1 for b in doc["blocks"] if b["kind"] == "prog"),
    }


# =====================================================================
# 3) 블록 -> DP 스크립트 워드 문서
# =====================================================================
class DPWriter:
    """리서치사 DP 스크립트 서식. 기본값은 납품 문서에서 실측한 값."""

    NAVY = "002060"
    BLUE = "0000FF"
    RED = "FF0000"

    def __init__(self, font="나눔고딕", base_pt=10.0, content_cm=18.6,
                 row_label_cm=9.85, spec_label_cm=2.5):
        self.font_name = font
        self.base = Pt(base_pt)
        self.content_cm = content_cm
        self.row_label_cm = row_label_cm
        self.spec_label_cm = spec_label_cm
        self.doc = self._new_document()

    # ------------------------------------------------------------ 기본
    def _new_document(self):
        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = self.font_name
        normal.font.size = self.base
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), self.font_name)
        pf = normal.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        sec = doc.sections[0]
        sec.top_margin = sec.bottom_margin = Cm(1.27)
        sec.left_margin = sec.right_margin = Cm(1.27)
        return doc

    def run(self, paragraph, text, bold=False, color=None, size=None,
            highlight=False):
        r = paragraph.add_run(text)
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            rFonts.set(qn(attr), self.font_name)
        r.font.size = size or self.base
        r.font.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        if highlight:
            hl = rPr.makeelement(qn("w:highlight"), {})
            hl.set(qn("w:val"), "yellow")
            rPr.append(hl)
        return r

    def para(self, text="", bold=False, color=None, align=None, after=0,
             size=None, highlight=False):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.0
        if align is not None:
            p.alignment = align
        if text:
            self.run(p, text, bold=bold, color=color, size=size,
                     highlight=highlight)
        return p

    @staticmethod
    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {})
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    @staticmethod
    def fix_layout(table, header_repeat=False):
        tblPr = table._tbl.tblPr
        layout = tblPr.makeelement(qn("w:tblLayout"), {})
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        for i, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
            if header_repeat and i == 0:
                trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))

    def set_widths(self, table, widths):
        for i, col in enumerate(table.columns):
            col.width = widths[i]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths):
                    cell.width = widths[i]

    def cell_text(self, cell, text, bold=False, color=None, align=None,
                  highlight=False, size=None):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if align is not None:
            p.alignment = align
        for n, line in enumerate(str(text).split("\n")):
            target = p if n == 0 else cell.add_paragraph()
            if n:
                target.paragraph_format.space_after = Pt(0)
                target.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
            self.run(target, line, bold=bold, color=color, highlight=highlight,
                     size=size)
        return cell

    # ------------------------------------------------------------ 구성 요소
    def spec_table(self, doc):
        rows = [("대상자", doc["대상자"]), ("샘플수", doc["샘플수"])]
        if doc["쿼터"] or doc["쿼터표"]:
            rows.append(("쿼터", doc["쿼터"]))
        if doc["제외"]:
            rows.append(("", doc["제외"]))
        if not any(v for _, v in rows):
            return

        table = self.doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        widths = [Cm(self.spec_label_cm),
                  Cm(self.content_cm - self.spec_label_cm)]

        for r, (label, value) in enumerate(rows):
            head, body = table.rows[r].cells
            self.cell_text(head, label, bold=True, color="FFFFFF",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            head.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self.shade(head, self.NAVY)
            self.cell_text(body, value, highlight=(label == "" and bool(value)))
            if label == "쿼터" and doc["쿼터표"]:
                self.quota_table(body, doc["쿼터표"])

        self.set_widths(table, widths)
        self.fix_layout(table)
        self.para(after=6)

    def quota_table(self, cell, rows):
        cols = max(len(r) for r in rows)
        inner = cell.add_table(rows=0, cols=cols)
        inner.style = "Table Grid"
        width = Cm((self.content_cm - self.spec_label_cm - 0.4) / cols)
        for row in rows:
            cells = inner.add_row().cells
            for c in range(cols):
                value = row[c] if c < len(row) else ""
                self.cell_text(cells[c], value)
                cells[c].width = width
        for col in inner.columns:
            col.width = width
        self.fix_layout(inner)

    def box(self, texts):
        table = self.doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        cell.width = Cm(self.content_cm)
        for n, text in enumerate(texts):
            p = cell.paragraphs[0] if n == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            self.run(p, text)
        for col in table.columns:
            col.width = Cm(self.content_cm)
        self.fix_layout(table)
        self.para(after=6)

    def matrix(self, block):
        scale = block["scale"] or SCALE_5
        rows = block["options"] or ["항목"]
        first = Cm(self.row_label_cm)
        rest = Cm((self.content_cm - self.row_label_cm) / len(scale))

        table = self.doc.add_table(rows=1, cols=len(scale) + 1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        self.cell_text(table.rows[0].cells[0], "")
        for i, label in enumerate(scale, 1):
            wrapped = "\n".join(label.split())      # '전혀 그렇지 않다' -> 3줄
            cell = table.rows[0].cells[i]
            self.cell_text(cell, f"{wrapped}\n{i}", align=WD_ALIGN_PARAGRAPH.CENTER)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

        for text in rows:
            cells = table.add_row().cells
            self.cell_text(cells[0], text)
            for i in range(1, len(scale) + 1):
                self.cell_text(cells[i], "", align=WD_ALIGN_PARAGRAPH.CENTER)

        self.set_widths(table, [first] + [rest] * len(scale))
        self.fix_layout(table, header_repeat=True)
        self.para()

    # ------------------------------------------------------------ 본문
    def write(self, doc):
        if doc["제목"]:
            self.para(doc["제목"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                      after=6, size=Pt(12))
        self.spec_table(doc)

        box_buffer: list[str] = []
        pending_input = None

        def flush_box():
            if box_buffer:
                self.box(list(box_buffer))
                box_buffer.clear()

        for b in doc["blocks"]:
            if b["kind"] == "note" and box_buffer:      # 상자 안 (예시) 줄
                box_buffer.append(f"({b['text']})")
                continue
            if b["kind"] != "box":
                flush_box()

            if b["kind"] == "box":
                box_buffer.append(b["text"])
            elif b["kind"] == "question":
                self.para()
                stem = self.para(f"{b['label']}. {b['text']} [{b['tag']}]".lstrip(". "))
                if b["tag"] == MATRIX_TAG:
                    stem.paragraph_format.keep_with_next = True
                    self.matrix(b)
                    pending_input = None
                elif "입력" in b["tag"]:
                    pending_input = self.para("________ 년 " if "년" in b["tag"]
                                              or "출생" in b["tag"] else "________ ")
                else:
                    pending_input = None
                    for n, opt in enumerate(b["options"], 1):
                        self.para(f"{n}) {opt}")
            elif b["kind"] == "prog":
                if pending_input is not None:          # 입력형은 같은 줄에 붙인다
                    self.run(pending_input, f"[PROG : {b['text']}]", color=self.BLUE)
                    pending_input = None
                else:
                    self.para(f"[PROG: {b['text']}]", color=self.BLUE)
            elif b["kind"] == "verify":
                self.para()
                self.para(f"[데이터 검증: {b['text']}]", bold=True, color=self.RED)
            else:
                self.para(b["text"])

        flush_box()
        return self

    def save(self, path):
        self.doc.save(path)
        return path

    def to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


def build_dp_docx(doc, **opts) -> bytes:
    return DPWriter(**opts).write(doc).to_bytes()
