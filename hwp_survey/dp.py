# -*- coding: utf-8 -*-
"""리서치사 납품용 DP 스크립트 형식.

인쇄용 설문지(writer.py)와 목적이 다르다. DP 스크립트는 조사 개요표,
SQ/Q 문항 번호, `[1개선택]` 같은 응답 방식 태그, 파란색 `[PROG: ...]`
프로그래밍 지시문, 빨간색 데이터 검증 지시문으로 이루어진다.

중간 텍스트(DP DSL)
    @제목: ...
    @대상자: ...
    @샘플수: 250명
    @쿼터: 성별*연령대별 균등할당
    @쿼터표: ,남자,여자
    @쿼터표: 만 20-29세,62,63
    @제외: 2026060452 참여자 제외
    ~ 상자글(용어 정의)
    SQ1. 문항 [1개선택]
    - 보기
    %PROG: 만20~39세만 진행
    Q1. 문항 [행별 1개선택]
    @행별: 전혀 그렇지 않다,그렇지 않다,보통이다,그렇다,매우 그렇다
    - 진술문
    %검증: Q1~Q2 일자찍기는 최종 납품 데이터에서 제외
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .parser import (CIRCLED, END, RE_BOX_SPLIT, RE_FIELDWORK, RE_FOOTNOTE,
                     RE_LEADIN, RE_OPT_CODE, RE_OPT_SPLIT, RE_RESP_TAG,
                     RE_ROMAN_HEAD, code_options, detect_label_style, grid_lines,
                     header_matrix, is_banner, is_prose_table, matrix_rows,
                     option_cell, option_table, scale_columns, scale_from_note,
                     scale_header, screening_rows)

# ---------------------------------------------------------------- 상수
SCALE_5 = ["전혀 그렇지 않다", "그렇지 않다", "보통이다", "그렇다", "매우 그렇다"]
MATRIX_HINT = "귀하의 의견과 가장 일치하는 정도에 체크해 주세요."

#: 번호 없는 라벨(AQ, DQ)도 쓰이므로 숫자는 있어도 없어도 된다
#: 'SQ1.', '문 3.', 'A0.', 'B2.', 'A3-2-1.' 처럼 쓰이는 문항 번호.
#: 파트 구분에 알파벳 한 글자를 쓰는 설문지가 많아 A~Z를 모두 받는다.
RE_LABEL = re.compile(
    rf"^\s*((?:SQ|DQ|AQ|[A-Z]|문|배문)\s*\d*(?:[-_]\d+){{0,3}})\s*{END}\s*(.*)$",
    re.I)
#: '(□ 예, □ 아니오) 나는 …' 형태의 동의 문항
RE_CONSENT = re.compile(r"^\s*\([^)]*[□☐][^)]*\)\s*(.+)$")
#: 보기 안의 '(설문 종료)' 지시
RE_STOP = re.compile(r"[(\[]\s*설문\s*(?:지\s*)?종료\s*[)\]]")
#: 'PART 1. 일반 특성', '제2부' 처럼 길이와 무관하게 구역을 나누는 줄
RE_PART_LINE = re.compile(r"^\s*(?:PART|파트|SECTION)\s*\d+|^\s*제\s*\d+\s*부", re.I)
#: 구역 제목으로 볼 낱말. 짧은 줄을 무조건 구역으로 보면 연구자 소속·연락처까지
#: 구역이 되어 버리므로 낱말로 좁힌다.
RE_SECTION_WORD = re.compile(
    r"선별|선정|스크리닝|screen|본문|demo|배경|인적\s*사항|일반\s*사항|이용$", re.I)
RE_FIELD = re.compile(r"^@(제목|대상자|샘플수|쿼터|쿼터표|제외|행별)\s*:\s*(.*)$")
RE_PROG_SRC = re.compile(r"^\s*\[?\s*PROG\s*[:：]\s*(.+?)\s*\]?\s*$", re.I)
#: '[DP: 20대, 30대 … 구간별 결과 분석]' 처럼 데이터 처리에 대한 지시문
RE_DP_SRC = re.compile(r"^\s*\[?\s*(?:DP|DATA|데이터)\s*[:：]\s*(.+?)\s*\]?\s*$", re.I)
#: '(  )명', '( )원' 같은 수치 기입란
RE_NUM_INPUT = re.compile(r"^\s*[(（]\s*[)）]\s*\S{0,4}\s*$")
RE_TAG = re.compile(r"\[([^\[\]]+)\]\s*$")
RE_NUM_Q = re.compile(rf"^\s*(?:문\s*)?(\d{{1,2}}(?:-\d{{1,2}})?)\s*{END}\s*(.+)$")
#: '귀하는 … 동의하십니까?' 처럼 번호 없이 문장만 있는 문항
RE_QUESTION_SENTENCE = re.compile(r"[?？]\s*$")
#: '② 아니오 응답 시 ☞ 문 3으로 이동' 같은 분기 지시문
RE_SKIP_NOTE = re.compile(r"응답\s*시|이동|종료|해당자만|건너뛰")
RE_OPT_LINE = re.compile(rf"^\s*(?:[{CIRCLED}]|\(\s*\d+\s*\)|\d{{1,2}}\s*[).])\s*(.+)$")
RE_PAREN_ONLY = re.compile(r"^\s*[(（]\s*(.+?)\s*[)）]\s*$")

ALONE_HINTS = ("이 중 없음", "없음", "이용하지 않음", "해당 없음")
#: 이 보기가 있으면 사실상 복수응답 문항이다(단독선택 보기를 따로 둔 것)
MULTI_SIGNS = ("이 중 없음", "이용하지 않음", "해당 없음")
SINGLE_TAG, MULTI_TAG, MATRIX_TAG = "1개선택", "모두선택", "행별 1개선택"


# =====================================================================
# 1) 추출 결과 -> DP 중간 텍스트
# =====================================================================
def items_to_dp_dsl(items, add_matrix_hint=True, add_alone_prog=True) -> str:
    """추출 결과를 DP 중간 텍스트로.

    원본이 문단만 있는 초안이든, 표로 리커트를 짠 학술 설문지든 모두 받는다.
    문항 번호가 SQ/Q 형태가 아니면 마지막에 자동으로 매긴다.
    """
    lines: list[str] = []
    quota_rows: list[list[str]] = []
    pending_scale: list[str] | None = None
    state = {"mode": None}
    consent_rows: list[str] = []
    style = detect_label_style(items)      # '문1.'/'SQ1.' 표기를 쓰는 설문지인지

    flat = list(items)
    i = 0
    while i < len(flat):
        kind, payload = flat[i]
        i += 1

        # ---------------------------------------------------------- 표
        if kind == "table":
            rows = payload
            joined = " ".join(c for r in rows for c in r)

            # 한 칸짜리 표 안에 문항과 보기가 통째로 들어 있는 경우가 많다.
            # 줄 단위로 풀어서 본문과 똑같이 처리한다.
            if expandable_cell(rows):
                flat[i:i] = [("p", ln) for ln in rows[0][0].splitlines()
                             if ln.strip()]
                continue

            opts = option_cell(rows)                # 보기만 담긴 한 칸짜리 표
            if opts:
                lines += [f"- {strip_circle(o).strip()}"
                          for o in RE_OPT_SPLIT.findall(opts) if strip_circle(o).strip()]
                continue

            if is_banner(rows):                     # '신문 이용' 같은 영역 배너
                pending_scale = flush_scale(lines, pending_scale)
                lines.append("")
                lines.append(f"## {rows[0][0].strip()}")
                continue

            screening = screening_rows(rows)         # 'SQ1. 거주 | 1. 서울 2. 부산 …'
            if screening:
                for raw_label, body in screening:
                    m = RE_LABEL.match(raw_label)
                    label = re.sub(r"\s+", "", m.group(1)).upper() if m else "?"
                    lead = (m.group(2).strip() if m else raw_label.strip())
                    block, _ = read_question(
                        [], 0, label, f"{lead} {body}".strip(),
                        add_matrix_hint, add_alone_prog, style, inline_only=True)
                    lines.append("")
                    lines.extend(block)
                continue

            cols = scale_columns(rows)                  # 척도 안내 표
            if cols:
                pending_scale = fill_scale(cols)
                continue

            body_rows, head_labels = rows, scale_header(rows)
            if head_labels:                         # 첫 행이 보기 라벨 줄이면 분리
                body_rows = rows[1:]
                # 이 표는 자체 라벨을 가지므로, 앞서 읽어둔 척도 안내표는
                # 바로 앞 문항의 보기였다는 뜻이다.
                pending_scale = flush_scale(lines, pending_scale)
            # 보기가 둘뿐인 표(예/아니오)는 행마다 기호도 둘뿐이다
            min_marks = max(2, min(3, len(head_labels))) if head_labels else 3
            matrix = matrix_rows(body_rows, min_marks=min_marks)
            head = head_labels or pending_scale
            if matrix is None:
                fallback = header_matrix(body_rows)
                if fallback:
                    head_raw, matrix = fallback
                    numeric = all(re.fullmatch(r"\d+", h.strip())
                                  for h in head_raw if h.strip())
                    # 표 머리가 숫자뿐이면 ※ 안내문에서 읽은 라벨을 쓴다
                    head = (pending_scale if numeric and pending_scale
                            and len(pending_scale) == len(head_raw)
                            else fill_scale(head_raw))
            if matrix:
                label, stem = pop_stem(lines)
                stem = stem or "다음 각 항목에 대해 응답해 주십시오."
                if add_matrix_hint and not re.search(r"체크|표시|응답", stem):
                    stem = f"{stem} {MATRIX_HINT}"
                lines.append("")
                lines.append(f"{label or '?'}. {stem} [{MATRIX_TAG}]")
                lines.append(f"@행별: {','.join(head or SCALE_5)}")
                lines.extend(strip_row_number(r) for r in matrix)  # 항목 코드 제거
                pending_scale = None
                continue

            if len(rows) == 1 and len(rows[0]) >= 2 and RE_ROMAN_HEAD.match(rows[0][0]):
                lines.append(f"! {' '.join(rows[0][1:]).strip()}")   # 섹션 안내문
                continue

            if "동의" in joined and "□" in joined:      # 참여 동의 표
                options = [c.lstrip("□ ").strip() for r in rows for c in r
                           if c.strip().startswith("□")]
                for r in rows:
                    for c in r:
                        if c.strip() and not c.strip().startswith("□"):
                            lines.append(f"~ {c.strip()}")
                if options:
                    lines.append("")
                    lines.append(f"?. 위 내용에 동의하십니까? [{SINGLE_TAG}]")
                    lines += [f"- {o}" for o in options]
                continue

            if state["mode"] == "쿼터" or is_quota_grid(rows):
                quota_rows.extend([c for c in r] for r in rows)
                continue

            opts_grid = option_table(rows)          # 보기를 여러 칸에 나눈 표
            if opts_grid:
                pending_scale = flush_scale(lines, pending_scale)
                lines += [f"- {o}" for o in opts_grid if o]
                continue

            if len(rows) >= 2 and not is_prose_table(rows):
                lines.extend(grid_lines(rows))       # 빈도·기입 표는 격자 그대로
                continue

            for r in rows:                              # 그 밖의 상자(표지/정의)
                for c in r:
                    if c.strip():
                        lines.extend(classify_free_line(c.strip(), lines))
            continue

        # ---------------------------------------------------------- 문단
        text = payload.strip()
        if not text or set(text) <= {"-", "―", "—", "="}:
            continue

        m = re.match(r"^[<〈《【]\s*(.+?)\s*[>〉》】]$", text)
        if m and not lines:
            lines.append(f"@제목: {m.group(1)}")
            continue
        m = re.match(r"^[■◆●]\s*(.+?)\s*[:：]\s*(.*)$", text)
        if m:
            name, value = m.group(1).strip(), m.group(2).strip()
            key = ("대상자" if "대상" in name else
                   "샘플수" if "샘플" in name else
                   "쿼터" if "쿼터" in name or "할당" in name else None)
            state["mode"] = "쿼터" if key == "쿼터" else None
            lines.append(f"@{key}: {value}" if key else f"~ {text}")
            continue

        if state["mode"] == "쿼터":
            if looks_like_quota_row(text):
                quota_rows.append([c for c in re.split(r"\s+|\t", text) if c])
                continue
            state["mode"] = None

        if re.match(r"^(SQ|Q)\s*[.)]\s*\S", text) and not RE_LABEL.match(text):
            continue                                     # 'SQ. 선정질문' 같은 구역 표시

        prog = RE_PROG_SRC.match(text)
        if prog:
            lines.append(f"%PROG: {prog.group(1)}")
            continue

        note_scale = scale_from_note(text, SCALE_5) if text.startswith("※") else None
        if note_scale:
            pending_scale = note_scale                 # ※ 1=…, 3=…, 5=…
            continue

        consent = RE_CONSENT.match(text)
        if consent:                                    # (□ 예, □ 아니오) 나는 …
            consent_rows.append(consent.group(1).strip())
            continue

        data_note = RE_DP_SRC.match(text)              # [DP: …]
        if data_note:
            lines.append(f"%검증: {data_note.group(1)}")
            continue

        field = RE_FIELDWORK.match(text)               # ▷ 조사원: …
        if field:
            lines.append(f"%PROG: 조사원 - {field.group(1).strip()}")
            continue
        if RE_FOOTNOTE.match(text) or RE_LEADIN.match(text):
            lines.append(f"! {text}")
            continue

        if (is_marked_option(text) and last_is_question(lines)
                and not RE_LABEL.match(text) and not RE_NUM_Q.match(text)):
            lines += option_lines(text)            # 앞 문항에 딸린 보기 줄
            continue

        m = RE_LABEL.match(text)
        if not m and (style == "bare" or has_inline_options(text)
                      or is_sub_numbered(text)):
            m = RE_NUM_Q.match(text)
        if m and pending_scale:
            pending_scale = flush_scale(lines, pending_scale)
        if not m and RE_QUESTION_SENTENCE.search(text) and options_follow(flat, i):
            block, i = read_question(flat, i, "?", text, add_matrix_hint,
                                     add_alone_prog, style)
            lines.append("")
            lines.extend(block)
            continue
        if m:
            label = (re.sub(r"\s+", "", m.group(1)).upper()
                     if RE_LABEL.match(text) else "?")
            block, i = read_question(flat, i, label, m.group(2).strip(),
                                     add_matrix_hint, add_alone_prog, style)
            lines.append("")
            lines.extend(block)
            continue

        lines.extend(classify_free_line(text, lines))

    pending_scale = flush_scale(lines, pending_scale)

    if consent_rows:
        anchor = next((n for n, l in enumerate(lines)
                       if RE_LABEL.match(l) or l.startswith("?.")), len(lines))
        block = ["", f"?. 아래 항목을 읽고, 응답해 주시기 바랍니다. [{MATRIX_TAG}]",
                 "@행별: 예,아니오"] + [f"- {t}" for t in consent_rows] + [""]
        lines[anchor:anchor] = block

    headings_to_sections(lines)
    retag_grid_questions(lines)
    promote_title(lines)

    if quota_rows:
        width = max(len(r) for r in quota_rows)
        quota_rows = [([""] * (width - len(r))) + r for r in quota_rows]
        head = [f"@쿼터표: {','.join(r)}" for r in quota_rows]
        anchor = next((n for n, l in enumerate(lines)
                       if l.startswith("@쿼터:")), len(lines) - 1)
        lines[anchor + 1:anchor + 1] = head

    return "\n".join(collapse(number_questions(lines))).strip()


GRID_TAG = "표 응답"
GENERIC_TITLES = ("설문지", "설 문 지", "질문지", "조사표")


#: '다음 문항은 … 묻는 문항입니다' 처럼 뒤따르는 문항들을 묶는 안내 줄
RE_GROUP_HEAD = re.compile(r"(문항입니다|질문입니다|문항이다|묻는 문항|관한 문항)")


def headings_to_sections(lines):
    """보기도 표도 없이 뒤 문항을 묶기만 하는 줄은 문항이 아니라 구역 제목이다."""
    for n, line in enumerate(lines):
        if not (RE_LABEL.match(line) or line.startswith("?.")):
            continue
        tag = RE_TAG.search(line.rstrip())
        if not tag or MATRIX_TAG in line or not RE_GROUP_HEAD.search(line):
            continue
        following = [l for l in lines[n + 1:] if l.strip()]
        nxt = following[0] if following else ""
        if nxt.startswith(("-", "@행별", "@표")):
            continue                              # 보기가 딸린 진짜 문항
        lines[n] = "## " + line[: tag.start()].strip()
    return lines


def retag_grid_questions(lines):
    """보기가 없고 바로 뒤가 표인 문항은 응답을 표에서 받는다."""
    for n, line in enumerate(lines):
        if not RE_TAG.search(line) or MATRIX_TAG in line:
            continue
        if not (RE_LABEL.match(line) or line.startswith("?.")):
            continue
        nxt = next((l for l in lines[n + 1:]
                    if l.strip() and not l.startswith("!")), "")
        if nxt.startswith("@표:"):
            lines[n] = RE_TAG.sub(f"[{GRID_TAG}]", line.rstrip())
    return lines


def promote_title(lines):
    """@제목이 '설문지'처럼 형식적이면 바로 뒤의 주제 줄을 제목으로 올린다."""
    for n, line in enumerate(lines):
        if not line.startswith("@제목:"):
            continue
        if line.split(":", 1)[1].strip() not in GENERIC_TITLES:
            return
        for m in range(n + 1, min(n + 4, len(lines))):
            body = lines[m].lstrip("~ ").strip()
            if lines[m].startswith("~") and 10 <= len(body) <= 70:
                lines[n] = f"@제목: {body}"
                lines.pop(m)
                return
        return


def classify_free_line(text: str, lines) -> list[str]:
    """문항이 아닌 줄: 제목 / 영역 배너 / 인사말·유의사항 상자 / 괄호 안내."""
    if not lines and len(text) <= 40:
        return [f"@제목: {text}"]
    if RE_PART_LINE.match(text):
        return ["", f"## {text}"]                  # 'PART 1. …' 구역 제목
    if len(text) <= 20 and RE_SECTION_WORD.search(text):
        return ["", f"## {text}"]                  # '선별 문항' 같은 영역 이름
    paren = RE_PAREN_ONLY.match(text)
    if paren and len(text) <= 30:
        return [f"! {paren.group(1)}"]
    return [f"~ {text}"]


def expandable_cell(rows) -> bool:
    """문항·보기가 여러 줄로 들어 있는 한 칸짜리 표인지."""
    if len(rows) != 1 or len(rows[0]) != 1 or "\n" not in rows[0][0]:
        return False
    lines = [l.strip() for l in rows[0][0].splitlines() if l.strip()]
    if len(lines) < 2:
        return False
    marked = sum(1 for l in lines
                 if is_option_line(l) or RE_LABEL.match(l) or RE_NUM_Q.match(l))
    return marked >= 2


def is_option_line(text: str) -> bool:
    return bool(RE_OPT_LINE.match(text) or has_inline_options(text))


def is_marked_option(text: str) -> bool:
    """기호(①, □)로 시작하는 보기 줄만. 숫자로 시작하는 줄은 문항일 수 있다."""
    head = text.strip()[:1]
    return bool(head and (head in CIRCLED or head in "□☐▢")) or has_inline_options(text)


def option_lines(text: str) -> list[str]:
    """'① 예 ② 아니오' 한 줄을 보기 여러 줄로."""
    inline = [strip_circle(o).strip() for o in RE_OPT_SPLIT.findall(text)]
    if len(inline) >= 2:
        return [f"- {o}" for o in inline if o]
    boxes = [strip_box(o).strip() for o in RE_BOX_SPLIT.findall(text)]
    if len(boxes) >= 2:
        return [f"- {o}" for o in boxes if o]
    m = RE_OPT_LINE.match(text)
    return [f"- {m.group(1).strip()}"] if m else []


def last_is_question(lines) -> bool:
    """마지막으로 쓴 줄이 문항이거나 그 보기인지."""
    for line in reversed(lines):
        s = line.strip()
        if not s:
            continue
        if s.startswith("-") or s.startswith("!") or s.startswith("%"):
            return True
        return bool(RE_LABEL.match(s) or s.startswith("?."))
    return False


def is_sub_numbered(text: str) -> bool:
    """'5-1.' '10-2.' 같은 하위 번호는 보기가 아니라 문항이다."""
    m = RE_NUM_Q.match(text)
    return bool(m and "-" in m.group(1))


def options_follow(flat, i) -> bool:
    """다음 항목이 보기 줄이거나 보기만 담긴 표인지."""
    if i >= len(flat):
        return False
    kind, payload = flat[i]
    if kind == "table":
        return option_cell(payload) is not None
    text = payload.strip()
    return bool(RE_OPT_LINE.match(text) or has_inline_options(text))


def has_inline_options(text: str) -> bool:
    """'1. 최종 학력 □ 고등학교 이하 □ 전문대학 …' 처럼 보기를 품은 줄."""
    return (len(RE_BOX_SPLIT.findall(text)) >= 2
            or len(RE_OPT_SPLIT.findall(text)) >= 2)


def flush_scale(lines, pending, used: bool = False):
    """행별 표에 쓰이지 못한 척도 안내표는 앞 문항의 보기로 돌린다.

    '전혀 없음 ① ② … ⑧ 매우 빈번함'처럼 문항 하나를 척도로 받는 형태다.
    """
    if not pending or used:
        return None if used else pending
    if last_is_question(lines):
        lines += [f"- {label}" for label in pending]
    return None


def strip_box(text: str) -> str:
    return re.sub(r"^[□☐▢]\s*", "", text).strip()


def clean_stem(text: str) -> str:
    """문항 문장 앞에 남은 기호·번호 찌꺼기를 떼어낸다."""
    return re.sub(r"^[\s.·]+", "", text or "").strip()


def strip_circle(text: str) -> str:
    """'① 전혀그렇지않다' -> '전혀그렇지않다' (DP 서식은 열 머리에 번호를 따로 찍는다)"""
    return re.sub(rf"^[{CIRCLED}]\s*", "", text).strip()


def fill_scale(cols: list[str]) -> list[str]:
    """양 끝만 라벨이 있는 척도('전혀그렇지않다 … 매우 그렇다')의 가운데를 채운다."""
    labels = [strip_circle(c) for c in cols]
    if all(labels):
        return labels
    if len(labels) == 5 and labels[0] and labels[-1]:
        filled = SCALE_5[:]                      # 5점은 표준 라벨로 채운다
        filled[0], filled[-1] = labels[0], labels[-1]
        return filled
    return [l or str(i) for i, l in enumerate(labels, 1)]


def strip_row_number(line: str) -> str:
    """행 머리의 항목 코드를 떼어낸다. '- 1. 나는 …', '- P2-1. 반려견은 …'"""
    return re.sub(r"^-\s*(?:\d{1,2}|[A-Za-z]{1,4}\s*\d{0,2}(?:[-_]\d{1,2})?)\s*[.)]\s*",
                  "- ", line)


def pop_stem(lines) -> tuple[str | None, str | None]:
    """표 바로 앞의 문항/안내문을 행별 문항의 (번호, 문장)으로 끌어올린다."""
    for i in range(len(lines) - 1, -1, -1):
        s = lines[i].strip()
        if not s:
            continue
        if s.startswith("!") and len(s) > 6:
            body = s.lstrip("! ").strip()
            if RE_SKIP_NOTE.search(body) or body[0] in CIRCLED:
                continue                          # '② … 응답 시 문3으로 이동'
            return None, lines.pop(i).lstrip("! ").strip()
        if s.startswith("~") and len(s) > 12:
            body = lines[i].lstrip("~ ").strip()
            if body[0] in CIRCLED or body.startswith("*"):
                continue                          # 보기 조각·각주는 문항이 아니다
            lines.pop(i)
            return None, clean_stem(body)
        if s.startswith("?."):                       # 번호 자리표시자 문항
            lines.pop(i)
            return None, clean_stem(RE_TAG.sub("", s[2:]))
        m = RE_LABEL.match(s)
        if m:                                        # '문40. …' 번호를 지킨다
            lines.pop(i)
            return re.sub(r"\s+", "", m.group(1)).upper(), \
                RE_TAG.sub("", m.group(2)).strip()
        m = RE_NUM_Q.match(s)
        if m:
            lines.pop(i)
            return None, RE_TAG.sub("", m.group(2)).strip()
        return None, None
    return None, None


def is_quota_grid(rows) -> bool:
    """숫자가 대부분인 작은 격자만 쿼터 표로 본다."""
    cells = [c.strip() for r in rows for c in r if c.strip()]
    if not cells or len(rows) > 6 or max(len(r) for r in rows) > 6:
        return False
    digits = sum(1 for c in cells if c.replace(",", "").isdigit())
    return digits >= len(cells) / 2 and all(len(c) <= 8 for c in cells)


def number_questions(lines):
    """번호가 없는 문항('?.')에 SQ/Q 번호를 매긴다.

    첫 행별 문항 앞은 선정 문항(SQ), 그 뒤는 본 문항(Q)으로 본다.
    """
    first_matrix = next((n for n, l in enumerate(lines)
                         if MATRIX_TAG in l and (l.startswith("?.")
                                                 or RE_LABEL.match(l))), None)
    sq = q = 0
    for n, line in enumerate(lines):
        if not line.startswith("?."):
            continue
        body = line[2:].strip()
        if first_matrix is not None and n < first_matrix:
            sq += 1
            lines[n] = f"SQ{sq}. {body}"
        else:
            q += 1
            lines[n] = f"Q{q}. {body}"
    return lines


def looks_like_quota_row(text: str) -> bool:
    """쿼터 표를 문단으로 흘려 쓴 줄. 머리 행('20대 30대 합계')과 숫자 행 모두.

    문항·보기·PROG 줄을 삼키지 않도록 조건을 좁게 잡는다.
    """
    if "?" in text or len(text) > 40:
        return False
    if RE_LABEL.match(text) or RE_OPT_LINE.match(text) or text.startswith("["):
        return False
    cells = [c for c in re.split(r"\s+|\t", text.strip()) if c]
    if not 2 <= len(cells) <= 6:
        return False
    return all(len(c) <= 6 for c in cells)


def read_question(flat, i, label, stem, add_matrix_hint, add_alone_prog,
                  style="bare", inline_only=False):
    """문항 한 덩어리(보기·척도 안내·PROG)를 읽어 DP DSL 줄로."""
    options: list[str] = []
    progs: list[str] = []
    datas: list[str] = []
    notes: list[str] = []
    scale: list[str] | None = None
    tag = None

    resp = RE_RESP_TAG.search(stem)                 # '[복수 응답]' 같은 꼬리표
    if resp:
        keyword = resp.group(1)
        stem = stem[: resp.start()].strip()
        tag = MULTI_TAG if ("복수" in keyword or "중복" in keyword
                            or "모두" in keyword) else SINGLE_TAG

    inline_stem = [strip_circle(o).strip() for o in RE_OPT_SPLIT.findall(stem)]
    if len(inline_stem) >= 2:                      # 'Q1】 … ① 예 ② 아니오'
        options += [o for o in inline_stem if o]
        cut = min(stem.index(ch) for ch in CIRCLED if ch in stem)
        stem = stem[:cut].strip()
    boxes = [strip_box(o) for o in RE_BOX_SPLIT.findall(stem)]
    if len(boxes) >= 2:                            # '□ 예  □ 아니오' 가 문항에 붙은 경우
        options += [b for b in boxes if b]
        stem = stem[: stem.index("□")].strip() if "□" in stem else stem
    elif style == "prefixed":                      # 한 줄에 몰린 코드 보기 분리
        codes = code_options(stem)
        if len(codes) >= 2:
            options += codes
            stem = stem[: stem.index(codes[0])].strip()
    if inline_only:
        return finish_question(label, stem, tag, scale, options, notes, progs,
                               add_matrix_hint, add_alone_prog, datas), i

    while i < len(flat):
        kind, payload = flat[i]
        if kind != "p":
            break
        text = payload.strip()
        if not text or set(text) <= {"-", "―", "—", "="}:
            i += 1
            continue
        if RE_LABEL.match(text) or is_sub_numbered(text) \
                or (style == "bare" and RE_NUM_Q.match(text)):
            break            # 다음 문항 시작 ('문1.' 표기 문서에서 '1.'은 보기다)
        if RE_PART_LINE.match(text):
            break                                  # 구역이 바뀌었다
        if text.startswith("※"):
            note_scale = scale_from_note(text, SCALE_5)
            if note_scale:
                scale = note_scale                 # ※ 1=…, 3=…, 5=…
                i += 1
                continue

        prog = RE_PROG_SRC.match(text)
        if prog:
            progs.append(prog.group(1))
            i += 1
            continue

        data_note = RE_DP_SRC.match(text)
        if data_note:
            datas.append(data_note.group(1))       # [DP: …] 는 데이터 지시문
            i += 1
            continue
        if RE_NUM_INPUT.match(text):
            tag = "수치 입력"                        # '(   )명' 기입란
            i += 1
            continue
        if RE_FIELDWORK.match(text):               # 조사원 지시문
            progs.append(f"조사원 - {RE_FIELDWORK.match(text).group(1).strip()}")
            i += 1
            continue
        if RE_FOOTNOTE.match(text) or RE_LEADIN.match(text):
            notes.append(text)
            i += 1
            continue

        inline = [o.strip() for o in RE_OPT_SPLIT.findall(text)]
        boxes = [o.strip() for o in RE_BOX_SPLIT.findall(text)]
        if len(inline) >= 2:                       # '① 예   ② 아니오' 처럼 한 줄에 여럿
            options += [strip_circle(o) for o in inline if strip_circle(o)]
            i += 1
            continue
        if len(boxes) >= 2:                        # '□ 예  □ 아니오' 형태
            options += [strip_box(o) for o in boxes if strip_box(o)]
            i += 1
            continue
        code = RE_OPT_CODE.match(text)
        if code and style == "prefixed":            # 응답 코드는 그대로 살린다
            options.append(f"{code.group(1)}. {code.group(2).strip()}")
            i += 1
            continue
        opt = RE_OPT_LINE.match(text)
        if opt:
            options.append(opt.group(1).strip())
            i += 1
            continue

        paren = RE_PAREN_ONLY.match(text)
        if paren and RE_SKIP_NOTE.search(paren.group(1)):
            progs.append(paren.group(1).strip())    # '② 아니오 응답 시 설문 종료'
            i += 1
            continue
        if paren is None and len(text) < 60 and not RE_PROG_SRC.match(text):
            notes.append(text)                     # ': 신문 PDF판' 같은 짧은 부속 줄
            i += 1
            continue
        if paren:
            inner = paren.group(1)
            if "점 척도" in inner:
                scale = scale_labels(inner)
                tag = MATRIX_TAG
            elif "출생" in inner or "년도" in inner:
                tag = "출생년도 입력"
            elif "동일" in inner or "지도" in inner:
                tag = "지도에서 선택"
                notes.append(inner)
            else:
                notes.append(inner)
            i += 1
            continue
        break                                   # 다음 문단은 문항 바깥(상자글 등)

    return finish_question(label, stem, tag, scale, options, notes, progs,
                           add_matrix_hint, add_alone_prog, datas), i


def finish_question(label, stem, tag, scale, options, notes, progs,
                    add_matrix_hint, add_alone_prog, datas=None) -> list[str]:
    """문항 한 덩어리를 DP DSL 줄로 마무리한다."""
    if tag is None:
        if re.search(r"복수\s*응답|모두", stem):
            tag = MULTI_TAG
        elif len(options) >= 3 and any(o.strip().startswith(h)
                                       for o in options for h in MULTI_SIGNS):
            tag = MULTI_TAG          # 단독선택용 배타 보기를 둔 문항
        elif scale:
            tag = MATRIX_TAG
        else:
            tag = SINGLE_TAG
    stem = re.sub(r"[(（]\s*복수\s*응답\s*[)）]", "", stem).strip()

    label_out = f"{label}." if label != "?" else "?."
    out = []
    if tag == MATRIX_TAG:
        if add_matrix_hint and not re.search(r"체크|표시", stem):
            stem = f"{stem} {MATRIX_HINT}".replace("  ", " ")
        out.append(f"{label_out} {stem} [{tag}]")
        out.append(f"@행별: {','.join(scale or SCALE_5)}")
    else:
        out.append(f"{label_out} {stem} [{tag}]")

    for n, opt in enumerate(options, 1):           # '아니오(설문 종료)'
        m = RE_STOP.search(opt) or re.search(r"\[\s*종료\s*\]", opt)
        if m:
            options[n - 1] = opt[: m.start()].strip()
            progs.append(f"{n}번 선택자 설문 종료")

    out += [f"! {n}" for n in notes]
    out += [f"- {o}" for o in options]

    if add_alone_prog and tag == MULTI_TAG:
        for n, opt in enumerate(options, 1):
            if any(opt.strip().startswith(h) or opt.strip() == h for h in ALONE_HINTS):
                text = f"{n}번 보기는 단독선택만 가능"
                if not any(text in p for p in progs):
                    progs.insert(0, text)
    if tag == "지도에서 선택" and not progs:
        progs.append("17개시도 지도제시")
    out += [f"%PROG: {p}" for p in progs]
    out += [f"%검증: {d}" for d in (datas or [])]
    return out


def scale_labels(inner: str) -> list[str]:
    n = int(re.search(r"(\d+)\s*점", inner).group(1))
    if n == 5:
        return SCALE_5
    if n == 7:
        return ["전혀 그렇지 않다", "그렇지 않다", "약간 그렇지 않다", "보통이다",
                "약간 그렇다", "그렇다", "매우 그렇다"]
    return [f"{i}" for i in range(1, n + 1)]


def collapse(lines):
    out = []
    for line in lines:
        if not line.strip() and (not out or not out[-1].strip()):
            continue
        out.append(line)
    return out


# =====================================================================
# 2) DP 중간 텍스트 -> 블록
# =====================================================================
def parse_dp(text: str) -> dict:
    doc = {"제목": "", "대상자": "", "샘플수": "", "쿼터": "", "쿼터표": [],
           "제외": "", "blocks": []}
    cur = None

    for raw in text.splitlines():
        s = raw.strip()
        if not s:
            continue

        field = RE_FIELD.match(s)
        if field:
            name, value = field.group(1), field.group(2).strip()
            if name == "쿼터표":
                doc["쿼터표"].append([c.strip() for c in value.split(",")])
            elif name == "행별":
                if cur:
                    cur["scale"] = [c.strip() for c in value.split(",") if c.strip()]
            else:
                doc[name] = value
            continue

        if s.startswith("@표:"):
            row = [c.strip() for c in s.split(":", 1)[1].split(",")]
            blocks = doc["blocks"]
            if blocks and blocks[-1]["kind"] == "grid":
                blocks[-1]["rows"].append(row)
            else:
                cur = None
                blocks.append({"kind": "grid", "rows": [row]})
            continue

        if s.startswith("##"):
            cur = None
            doc["blocks"].append({"kind": "section",
                                  "text": s.lstrip("#").strip()})
            continue

        if s.startswith("%PROG:") or s.startswith("%prog:"):
            doc["blocks"].append({"kind": "prog", "text": s.split(":", 1)[1].strip()})
            continue
        if s.startswith("%검증:"):
            doc["blocks"].append({"kind": "verify", "text": s.split(":", 1)[1].strip()})
            continue
        if s.startswith("~"):
            cur = None
            doc["blocks"].append({"kind": "box", "text": s.lstrip("~ ").strip()})
            continue
        if s.startswith("!"):
            doc["blocks"].append({"kind": "note", "text": s.lstrip("! ").strip()})
            continue
        if s.startswith("--"):
            target = cur or last_question(doc)
            if target:
                target["options"].append({"type": "group",
                                          "text": s.lstrip("- ").strip()})
            continue
        if s.startswith("-"):
            # 표(@표:)가 끼어들어 cur 가 끊긴 뒤에도 보기를 잃지 않는다
            target = cur or last_question(doc)
            if target:
                target["options"].append({"type": "row",
                                          "text": s.lstrip("- ").strip()})
            else:
                doc["blocks"].append({"kind": "note", "text": s.lstrip("- ").strip()})
            continue

        m = RE_LABEL.match(s)
        label, body = (m.group(1).upper(), m.group(2)) if m else ("", s)
        tag = SINGLE_TAG
        t = RE_TAG.search(body)
        if t:
            tag = t.group(1).strip()
            body = body[: t.start()].strip()
        cur = {"kind": "question", "label": label, "text": body.strip(),
               "tag": tag, "options": [], "scale": None}
        doc["blocks"].append(cur)

    return doc


def last_question(doc):
    for block in reversed(doc["blocks"]):
        if block["kind"] == "question":
            return block
    return None


def summarize_dp(doc) -> dict:
    qs = [b for b in doc["blocks"] if b["kind"] == "question"]
    return {
        "문항": len(qs),
        "일반 표": sum(1 for b in doc["blocks"] if b["kind"] == "grid"),
        "선정문항(SQ)": sum(1 for q in qs if q["label"].startswith("SQ")),
        "행별 표": sum(1 for q in qs if q["tag"] == MATRIX_TAG),
        "PROG 지시문": sum(1 for b in doc["blocks"] if b["kind"] == "prog"),
    }


# =====================================================================
# 3) 블록 -> DP 스크립트 워드 문서
# =====================================================================
class DPWriter:
    """리서치사 DP 스크립트 서식. 기본값은 납품 문서에서 실측한 값."""

    NAVY = "002060"
    BLUE = "0000FF"
    RED = "FF0000"

    def __init__(self, font="나눔고딕", base_pt=10.0, content_cm=18.6,
                 row_label_cm=9.85, spec_label_cm=2.5):
        self.font_name = font
        self.base = Pt(base_pt)
        self.content_cm = content_cm
        self.row_label_cm = row_label_cm
        self.spec_label_cm = spec_label_cm
        self.doc = self._new_document()

    # ------------------------------------------------------------ 기본
    def _new_document(self):
        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = self.font_name
        normal.font.size = self.base
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), self.font_name)
        pf = normal.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        sec = doc.sections[0]
        sec.top_margin = sec.bottom_margin = Cm(1.27)
        sec.left_margin = sec.right_margin = Cm(1.27)
        return doc

    def run(self, paragraph, text, bold=False, color=None, size=None,
            highlight=False):
        r = paragraph.add_run(text)
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            rFonts.set(qn(attr), self.font_name)
        r.font.size = size or self.base
        r.font.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        if highlight:
            hl = rPr.makeelement(qn("w:highlight"), {})
            hl.set(qn("w:val"), "yellow")
            rPr.append(hl)
        return r

    def para(self, text="", bold=False, color=None, align=None, after=0,
             size=None, highlight=False):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.0
        if align is not None:
            p.alignment = align
        if text:
            self.run(p, text, bold=bold, color=color, size=size,
                     highlight=highlight)
        return p

    @staticmethod
    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {})
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    @staticmethod
    def fix_layout(table, header_repeat=False):
        tblPr = table._tbl.tblPr
        layout = tblPr.makeelement(qn("w:tblLayout"), {})
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        for i, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(trPr.makeelement(qn("w:cantSplit"), {}))
            if header_repeat and i == 0:
                trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))

    def set_widths(self, table, widths):
        for i, col in enumerate(table.columns):
            col.width = widths[i]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths):
                    cell.width = widths[i]

    def cell_text(self, cell, text, bold=False, color=None, align=None,
                  highlight=False, size=None):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if align is not None:
            p.alignment = align
        for n, line in enumerate(str(text).split("\n")):
            target = p if n == 0 else cell.add_paragraph()
            if n:
                target.paragraph_format.space_after = Pt(0)
                target.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
            self.run(target, line, bold=bold, color=color, highlight=highlight,
                     size=size)
        return cell

    # ------------------------------------------------------------ 구성 요소
    def spec_table(self, doc):
        rows = [("대상자", doc["대상자"]), ("샘플수", doc["샘플수"])]
        if doc["쿼터"] or doc["쿼터표"]:
            rows.append(("쿼터", doc["쿼터"]))
        if doc["제외"]:
            rows.append(("", doc["제외"]))
        if not any(v for _, v in rows):
            return

        table = self.doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        widths = [Cm(self.spec_label_cm),
                  Cm(self.content_cm - self.spec_label_cm)]

        for r, (label, value) in enumerate(rows):
            head, body = table.rows[r].cells
            self.cell_text(head, label, bold=True, color="FFFFFF",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            head.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self.shade(head, self.NAVY)
            self.cell_text(body, value, highlight=(label == "" and bool(value)))
            if label == "쿼터" and doc["쿼터표"]:
                self.quota_table(body, doc["쿼터표"])

        self.set_widths(table, widths)
        self.fix_layout(table)
        self.para(after=6)

    def quota_table(self, cell, rows):
        cols = max(len(r) for r in rows)
        inner = cell.add_table(rows=0, cols=cols)
        inner.style = "Table Grid"
        width = Cm((self.content_cm - self.spec_label_cm - 0.4) / cols)
        for row in rows:
            cells = inner.add_row().cells
            for c in range(cols):
                value = row[c] if c < len(row) else ""
                self.cell_text(cells[c], value)
                cells[c].width = width
        for col in inner.columns:
            col.width = width
        self.fix_layout(inner)

    def box(self, texts):
        table = self.doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        cell.width = Cm(self.content_cm)
        for n, text in enumerate(texts):
            p = cell.paragraphs[0] if n == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            self.run(p, text)
        for col in table.columns:
            col.width = Cm(self.content_cm)
        self.fix_layout(table)
        self.para(after=6)

    def section(self, text):
        """'신문 이용' 같은 영역 배너: 진한 바탕의 한 칸 표."""
        self.para()
        table = self.doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        cell.width = Cm(self.content_cm)
        self.cell_text(cell, text, bold=True, color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        self.shade(cell, self.NAVY)
        for col in table.columns:
            col.width = Cm(self.content_cm)
        self.fix_layout(table)
        self.para()

    def grid(self, rows):
        """분류되지 않은 원본 표를 격자 그대로 옮긴다(빈도 표, 기입 표 등)."""
        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=0, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        width = Cm(self.content_cm / cols)
        for n, row in enumerate(rows):
            cells = table.add_row().cells
            for c in range(cols):
                self.cell_text(cells[c], row[c] if c < len(row) else "",
                               bold=(n == 0), align=WD_ALIGN_PARAGRAPH.CENTER,
                               size=Pt(9))
                cells[c].width = width
        for col in table.columns:
            col.width = width
        self.fix_layout(table)
        self.para()

    def matrix(self, block):
        scale = block["scale"] or SCALE_5
        rows = block["options"] or [{"type": "row", "text": "항목"}]
        first = Cm(self.row_label_cm)
        rest = Cm((self.content_cm - self.row_label_cm) / len(scale))

        table = self.doc.add_table(rows=1, cols=len(scale) + 1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        self.cell_text(table.rows[0].cells[0], "")
        for i, label in enumerate(scale, 1):
            wrapped = "\n".join(label.split())      # '전혀 그렇지 않다' -> 3줄
            cell = table.rows[0].cells[i]
            self.cell_text(cell, f"{wrapped}\n{i}", align=WD_ALIGN_PARAGRAPH.CENTER)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

        for item in rows:
            cells = table.add_row().cells
            if item["type"] == "group":                 # 표 안 소제목 행
                merged = cells[0].merge(cells[-1])
                self.cell_text(merged, item["text"], bold=True)
                self.shade(merged, "EAEFF7")
                continue
            self.cell_text(cells[0], item["text"])
            for i in range(1, len(scale) + 1):
                self.cell_text(cells[i], "", align=WD_ALIGN_PARAGRAPH.CENTER)

        self.set_widths(table, [first] + [rest] * len(scale))
        self.fix_layout(table, header_repeat=True)
        self.para()

    # ------------------------------------------------------------ 본문
    def write(self, doc):
        if doc["제목"]:
            self.para(doc["제목"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                      after=6, size=Pt(12))
        self.spec_table(doc)

        box_buffer: list[str] = []
        pending_input = None

        def flush_box():
            if box_buffer:
                self.box(list(box_buffer))
                box_buffer.clear()

        for b in doc["blocks"]:
            if b["kind"] == "note" and box_buffer and len(b["text"]) <= 30:
                box_buffer.append(f"({b['text']})")     # 상자 안 (예시) 줄
                continue
            if b["kind"] != "box":
                flush_box()

            if b["kind"] == "box":
                box_buffer.append(b["text"])
            elif b["kind"] == "section":
                self.section(b["text"])
                pending_input = None
            elif b["kind"] == "grid":
                self.grid(b["rows"])
                pending_input = None
            elif b["kind"] == "question":
                self.para()
                stem = self.para(f"{b['label']}. {b['text']} [{b['tag']}]".lstrip(". "))
                if b["tag"] == MATRIX_TAG:
                    stem.paragraph_format.keep_with_next = True
                    self.matrix(b)
                    pending_input = None
                elif "입력" in b["tag"]:
                    pending_input = self.para("________ 년 " if "년" in b["tag"]
                                              or "출생" in b["tag"] else "________ ")
                else:
                    pending_input = None
                    n = 0
                    for opt in b["options"]:
                        if opt["type"] == "group":
                            self.para(opt["text"], bold=True)
                            continue
                        if re.match(r"^\d{1,4}[.)]", opt["text"]):
                            self.para(opt["text"])       # '9997. 기타' 같은 코드
                            continue
                        n += 1
                        self.para(f"{n}) {opt['text']}")
            elif b["kind"] == "prog":
                if pending_input is not None:          # 입력형은 같은 줄에 붙인다
                    self.run(pending_input, f"[PROG : {b['text']}]", color=self.BLUE)
                    pending_input = None
                else:
                    self.para(f"[PROG: {b['text']}]", color=self.BLUE)
            elif b["kind"] == "verify":
                self.para()
                self.para(f"[데이터 검증: {b['text']}]", bold=True, color=self.RED)
            else:
                self.para(b["text"])

        flush_box()
        return self

    def save(self, path):
        self.doc.save(path)
        return path

    def to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


def build_dp_docx(doc, **opts) -> bytes:
    return DPWriter(**opts).write(doc).to_bytes()
